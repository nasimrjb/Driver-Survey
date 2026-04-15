"""Generate the complete Driver Survey Pipeline Guide as a .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level, size in [(1, 18), (2, 15), (3, 13)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)

# Helpers
def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    return p

def note(text):
    p = doc.add_paragraph()
    r1 = p.add_run('Note: ')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.italic = True

def steps(items):
    for i, s in enumerate(items, 1):
        doc.add_paragraph(f'{i}. {s}')

def bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def visual_spec(title, props):
    doc.add_heading(title, level=3)
    add_table(['Property', 'Value'], props)
    doc.add_paragraph()


# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Driver Survey\nComplete Pipeline Guide')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(
    'From SQL Server Setup to Power BI Dashboards & HTML Reports\n\n'
    'Server: 192.168.18.37 | Database: Cab_Studies | Schema: Cab\n'
    'Updated: April 2026'
)
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc = [
    'Phase 1: SQL Server Foundation',
    '  1.1 Create Base Tables & Load Data',
    '  1.2 Create SQL Views',
    '  1.3 YearWeek Formatting (ISO 8601)',
    'Phase 2: HTML Dashboards',
    '  2.1 Driver Survey Dashboard',
    '  2.2 Routine Analysis Dashboard',
    '  2.3 Web Server Setup',
    'Phase 3: Power BI Report',
    '  3.1 Connect to SQL Server',
    '  3.2 Data Model & Relationships',
    '  3.3 YearWeek Sort-by-Column Setup',
    '  3.4 DAX Measures',
    '  3.5 Page 1: Executive Overview',
    '  3.6 Page 2: NPS Deep-Dive',
    '  3.7 Page 3: Satisfaction Deep-Dive',
    '  3.8 Page 4: Incentive Analysis',
    '  3.9 Page 5: Operations & Demographics',
    'Phase 4: Publishing & Maintenance',
    '  4.1 Publish to Power BI Service',
    '  4.2 Scheduled Refresh',
    '  4.3 Adding New Survey Weeks',
    'Appendix A: File Inventory',
    'Appendix B: Color Palette',
    'Appendix C: Troubleshooting',
]
for item in toc:
    doc.add_paragraph(item)
doc.add_page_break()

# ============================================================
# PHASE 1: SQL SERVER FOUNDATION
# ============================================================
doc.add_heading('Phase 1: SQL Server Foundation', level=1)
doc.add_paragraph(
    'This phase sets up the database tables, loads the survey data, and creates '
    'the 20 SQL views that both the HTML dashboards and Power BI report consume.'
)

# --- 1.1 ---
doc.add_heading('1.1 Create Base Tables & Load Data', level=2)
doc.add_paragraph(
    'The pipeline uses 6 base tables in the Cab schema. The Python script load_to_database.py '
    'handles everything: creating tables with proper column types and loading data from CSVs.'
)
doc.add_paragraph()

doc.add_heading('load_to_database.py', level=3)
doc.add_paragraph('File: PowerBI/load_to_database.py')
doc.add_paragraph(
    'This is the single script for all table creation and data loading. It reads processed CSV files, '
    'determines column types from Sources/column_rename_mapping.json, creates properly-typed tables '
    '(INT, FLOAT, DATETIME, NVARCHAR), and bulk-inserts data. It supports two modes:'
)
doc.add_paragraph()

doc.add_heading('Incremental Mode (default) — for adding new survey weeks', level=3)
doc.add_paragraph(
    'Compares recordIDs in the CSV against those already in the database. '
    'Only rows with NEW recordIDs are inserted. Safe to re-run without duplicating data.'
)
steps([
    'Ensure Python 3.x is installed with pyodbc and pandas.',
    'Ensure VPN to Snapp data center is active.',
    'Process new survey data so CSVs in processed/ contain all weeks (old + new).',
    'Run: python PowerBI/load_to_database.py',
    'The script creates tables if they do not exist, then inserts only new rows.',
])
doc.add_paragraph()

doc.add_heading('Full Reload Mode — for rebuilding from scratch', level=3)
doc.add_paragraph(
    'Drops and recreates all 6 tables, then loads everything. '
    'Use when column types change or data needs to be completely replaced.'
)
steps([
    'Run: python PowerBI/load_to_database.py --full-reload',
])
doc.add_paragraph()
note(
    'The file create_driver_survey_table.sql is included as a schema reference — '
    'you can open it in SSMS to inspect column names and types. '
    'However, it is NOT required: load_to_database.py handles table creation automatically.'
)
doc.add_paragraph()

doc.add_heading('The 6 Tables', level=3)
add_table(
    ['Table', 'Columns', 'Source CSV', 'Description'],
    [
        ['DriverSurvey_ShortMain', '68', 'short_survey_main.csv', 'Core weekly fields (satisfaction, rides, demographics)'],
        ['DriverSurvey_ShortRare', '176', 'short_survey_rare.csv', 'Less-frequent fields (referral, registration, CS)'],
        ['DriverSurvey_WideMain', '67', 'wide_survey_main.csv', 'Binary-encoded incentive type columns'],
        ['DriverSurvey_WideRare', '190', 'wide_survey_rare.csv', 'Binary-encoded unsatisfaction/refusal reason columns'],
        ['DriverSurvey_LongMain', '28', 'long_survey_main.csv', 'Question-answer pairs for main survey'],
        ['DriverSurvey_LongRare', '9', 'long_survey_rare.csv', 'Question-answer pairs for rare survey'],
    ]
)
doc.add_paragraph()
note(
    'Both approaches produce identical table structures with proper column types derived from '
    'Sources/column_rename_mapping.json (dtype field). Columns are typed as INT, FLOAT, DATETIME, '
    'or NVARCHAR(500) based on their data type. Computed/derived columns not in the JSON '
    '(e.g., snapp_ride, wheel, edu) are auto-detected as FLOAT; binary 0/1 wide-format columns '
    'are typed as INT. recordID is always INT NOT NULL.'
)

# --- 1.2 ---
doc.add_heading('1.2 Create SQL Views', level=2)
doc.add_paragraph('File: PowerBI/create_views.sql')
doc.add_paragraph(
    'This script creates 20 views plus 1 base view (vw_ShortBase) in the Cab schema. '
    'vw_ShortBase joins ShortMain and ShortRare, adds computed columns (yearweek, driver_type), '
    'and serves as the foundation for all other views.'
)
doc.add_paragraph()

doc.add_heading('Complete View List', level=3)
views_desc = [
    ['vw_ShortBase', 'Base view: joins ShortMain + ShortRare, adds yearweek & driver_type'],
    ['vw_KPISummary', 'Single-row summary: total responses, weeks, cities, overall satisfaction'],
    ['vw_WeeklySatisfaction', 'Weekly averages of all satisfaction metrics, rides, incentive'],
    ['vw_WeeklyNPS', 'Weekly NPS scores: promoter/passive/detractor percentages'],
    ['vw_SatisfactionByCity', 'Per-city satisfaction averages and NPS'],
    ['vw_SatisfactionByCityWeek', 'City x Week satisfaction matrix (for heatmaps)'],
    ['vw_SatisfactionByDemographics', 'Satisfaction by age group, gender, cooperation type, driver type'],
    ['vw_HoneymoonEffect', 'Satisfaction by driver tenure (Snapp age)'],
    ['vw_IncentiveByWeek', 'Weekly incentive amounts, satisfaction, message/participation rates'],
    ['vw_IncentiveByCity', 'Per-city incentive metrics'],
    ['vw_IncentiveAmountByCity', 'Incentive amount bracket distribution by city and platform'],
    ['vw_RideShareByCityWeek', 'Snapp vs Tapsi ride share percentages by city and week'],
    ['vw_RideShare', 'Aggregate ride share metrics'],
    ['vw_NavigationUsage', 'Navigation app preference counts'],
    ['vw_NavigationByWeek', 'Navigation app usage percentage by week'],
    ['vw_Demographics', 'Demographic distributions (age, gender, city, cooperation type, etc.)'],
    ['vw_PersonaByCity', 'Driver persona breakdown by city'],
    ['vw_WideIncentiveTypes', 'Incentive type binary counts and percentages'],
    ['vw_WideUnsatisfactionReasons', 'Unsatisfaction reason counts and percentages by platform'],
    ['vw_LongSurveyAnswers', 'Main survey question-answer distributions'],
    ['vw_LongRareSurveyAnswers', 'Rare survey question-answer distributions'],
]
add_table(['View Name', 'Description'], views_desc)
doc.add_paragraph()

doc.add_heading('How to Run', level=3)
steps([
    'Connect to SQL Server (192.168.18.37) using SSMS.',
    'Open PowerBI/create_views.sql.',
    'Execute the entire script. Views are created with IF EXISTS / DROP logic (safe to re-run).',
    'Verify with: SELECT * FROM Cab.vw_KPISummary (should return 1 row).',
])

# --- 1.3 ---
doc.add_heading('1.3 YearWeek Formatting (ISO 8601)', level=2)
doc.add_paragraph(
    'Week numbers use standard ISO 8601 (Monday-based, week 1 contains the first Thursday of the year). '
    'Six views contain a yearweek column formatted as text "YY-WW" (e.g., "25-01") '
    'for clean axis labels, with a companion integer yearweek_sort column (e.g., 2501) '
    'for proper chronological ordering.'
)
doc.add_paragraph()
doc.add_paragraph(
    'The yearweek integer handles ISO year boundaries: week 52+ in January maps to the previous year, '
    'and week 1 in December maps to the next year. This avoids overlap issues at year boundaries.'
)
doc.add_paragraph()
doc.add_heading('SQL Formula Used in Views', level=3)
code("CASE WHEN weeknumber >= 52 AND MONTH(datetime) = 1\n"
     "     THEN ((YEAR(datetime) - 1) % 100) * 100 + weeknumber\n"
     "     WHEN weeknumber = 1 AND MONTH(datetime) = 12\n"
     "     THEN ((YEAR(datetime) + 1) % 100) * 100 + weeknumber\n"
     "     ELSE (YEAR(datetime) % 100) * 100 + weeknumber END AS yearweek")
doc.add_paragraph()
doc.add_heading('Views with yearweek', level=3)
bullets([
    'vw_WeeklySatisfaction', 'vw_WeeklyNPS', 'vw_SatisfactionByCityWeek',
    'vw_IncentiveByWeek', 'vw_RideShareByCityWeek', 'vw_NavigationByWeek'
])

doc.add_page_break()

# ============================================================
# PHASE 2: HTML DASHBOARDS
# ============================================================
doc.add_heading('Phase 2: HTML Dashboards', level=1)
doc.add_paragraph(
    'Two interactive HTML dashboards are generated from SQL Server data and served via a Python web server.'
)

# --- 2.1 ---
doc.add_heading('2.1 Driver Survey Dashboard', level=2)
doc.add_paragraph('File: PowerBI/build_dashboard.py')
doc.add_paragraph('Output: PowerBI/DriverSurvey_Dashboard.html')
doc.add_paragraph(
    'Generates an interactive 6-page dashboard with Plotly charts covering satisfaction trends, NPS, '
    'incentive analysis, demographics, navigation usage, per-question explorers, and per-city ride share. '
    'Charts are organized by topic tabs with Persian (Jalali) week labels on all time-series axes.'
)
doc.add_heading('How to Generate', level=3)
steps([
    'Ensure VPN connection to Snapp data center is active.',
    'Run: python PowerBI/build_dashboard.py',
    'The script queries all SQL views and generates DriverSurvey_Dashboard.html.',
    'Open the HTML file in a browser to view.',
])
doc.add_paragraph()
doc.add_heading('Dashboard Pages', level=3)
add_table(
    ['Page', 'Tab Name', 'Key Charts'],
    [
        ['1', 'Executive Overview', 'KPI cards, weekly satisfaction lines, NPS trend, NPS decomposition, response count'],
        ['2', 'Satisfaction Deep-Dive', 'City dot plot, demographics dumbbell, honeymoon effect, city×week heatmap'],
        ['3', 'Incentive Analysis', 'ROI combo (bars+lines), funnel (% of all drivers), lollipop types, butterfly unsatisfaction (%), city dot plot'],
        ['4', 'Operations & Demographics', 'Navigation donuts, nav market share area, demographic lollipops, gender bar, ride share line'],
        ['5', 'Survey Explorer', 'Per-question horizontal bars with gradient coloring (main + rare)'],
        ['6', 'Ride Share by City', 'One Snapp vs Tapsi line chart per city (21 cities)'],
    ]
)
doc.add_paragraph()
note(
    'All time-series x-axes show Persian (Jalali) week labels (e.g., "25-W14: 18-24 Far"). '
    'Legends are positioned above charts to avoid overlap with x-axis labels. '
    'Charts with background sample-size bars use semi-transparent fills so lines remain visually dominant.'
)
doc.add_paragraph()

# --- 2.2 ---
doc.add_heading('2.2 Routine Analysis Dashboard', level=2)
doc.add_paragraph('File: PowerBI/build_routine_dashboard.py')
doc.add_paragraph('Output: PowerBI/RoutineAnalysis_Dashboard.html')
doc.add_paragraph(
    'Generates a heatmap-style dashboard with 45 sheets across 8 tabs (Incentive, Satisfaction, '
    'Market Share, Operations, Commission Free, Support & NPS, Registration, Decline & Demand). '
    'Tables use color scales matching the Excel routine analysis output: white-to-green for percentages, '
    'white-to-red for dissatisfaction, red-yellow-green for 1-5 satisfaction scores.'
)
doc.add_heading('How to Generate', level=3)
steps([
    'Ensure VPN connection is active.',
    'Run: python PowerBI/build_routine_dashboard.py',
    'The script generates RoutineAnalysis_Dashboard.html.',
])

# --- 2.3 ---
doc.add_heading('2.3 Web Server Setup', level=2)
doc.add_paragraph('File: PowerBI/serve_dashboards.py')
doc.add_paragraph('Launcher: PowerBI/start_server.bat')
doc.add_paragraph(
    'A Python HTTP server on port 8800 serves both dashboards with live SQL Server refresh capability.'
)
doc.add_paragraph()
doc.add_heading('Endpoints', level=3)
add_table(
    ['URL', 'Description'],
    [
        ['http://localhost:8800/', 'Index page with links to both dashboards'],
        ['http://localhost:8800/driver-survey', 'Driver Survey Dashboard'],
        ['http://localhost:8800/routine-analysis', 'Routine Analysis Dashboard'],
        ['http://localhost:8800/refresh/driver-survey', 'Regenerate Driver Survey from SQL Server'],
        ['http://localhost:8800/refresh/routine', 'Regenerate Routine Analysis from SQL Server'],
        ['http://localhost:8800/api/status', 'Server status and last refresh timestamps'],
    ]
)
doc.add_paragraph()
doc.add_heading('How to Start', level=3)
steps([
    'Double-click start_server.bat, or run: python PowerBI/serve_dashboards.py',
    'The server starts on http://localhost:8800.',
    'To refresh data from SQL Server, click the refresh buttons on the index page or visit /refresh/* endpoints.',
    'The server uses subprocess to call build scripts, avoiding import/reload issues.',
])

doc.add_page_break()

# ============================================================
# PHASE 3: POWER BI REPORT
# ============================================================
doc.add_heading('Phase 3: Power BI Report', level=1)
doc.add_paragraph('File: PowerBI/Dashboard.pbix')
doc.add_paragraph(
    'A Power BI report connected to the same SQL views, published to Power BI Service '
    'for interactive exploration with slicers, cross-filtering, and drill-through.'
)

# --- 3.1 ---
doc.add_heading('3.1 Connect to SQL Server', level=2)
steps([
    'Open Power BI Desktop.',
    'Click "Get Data" > "SQL Server".',
    'Enter connection details (see table below).',
    'In Navigator, expand Cab_Studies > Cab schema > select all 20 vw_* views.',
    'Click "Load" (Import mode).',
])
doc.add_paragraph()
add_table(
    ['Parameter', 'Value'],
    [
        ['Server', '192.168.18.37'],
        ['Database', 'Cab_Studies'],
        ['Data Connectivity', 'Import'],
        ['Authentication', 'Database'],
        ['User name', 'nasim.rajabi'],
        ['Password', '(enter your password)'],
    ]
)

# --- 3.2 ---
doc.add_heading('3.2 Data Model & Relationships', level=2)
doc.add_paragraph(
    'Star schema: vw_WeeklySatisfaction is the hub. Other time-series views connect via yearweek.'
)
doc.add_paragraph()
add_table(
    ['From Table', 'Column', 'To Table', 'Column', 'Cardinality', 'Direction'],
    [
        ['vw_WeeklySatisfaction', 'yearweek', 'vw_WeeklyNPS', 'yearweek', '1:1', 'Both'],
        ['vw_IncentiveByWeek', 'yearweek', 'vw_WeeklySatisfaction', 'yearweek', '1:1', 'Both'],
        ['vw_SatisfactionByCityWeek', 'yearweek', 'vw_WeeklySatisfaction', 'yearweek', 'Many:1', 'Single'],
        ['vw_RideShareByCityWeek', 'yearweek', 'vw_WeeklySatisfaction', 'yearweek', 'Many:1', 'Single'],
        ['vw_NavigationByWeek', 'yearweek', 'vw_WeeklySatisfaction', 'yearweek', 'Many:1', 'Single'],
    ]
)
doc.add_paragraph()
note('Remaining views are standalone and do not require relationships.')

# --- 3.3 ---
doc.add_heading('3.3 YearWeek Sort-by-Column Setup (Required, One-Time)', level=2)
doc.add_paragraph(
    'The yearweek column is text ("25-01") for clean labels. You must tell Power BI to sort it '
    'by the integer yearweek_sort column for correct chronological order.'
)
doc.add_paragraph()
steps([
    'Go to Data view (table icon on left sidebar).',
    'Select a table (e.g., vw_WeeklySatisfaction).',
    'Click the "yearweek" column header.',
    'In the ribbon: Column tools > Sort by Column > select "yearweek_sort".',
    'Repeat for all 6 tables: vw_WeeklySatisfaction, vw_WeeklyNPS, vw_SatisfactionByCityWeek, '
    'vw_IncentiveByWeek, vw_RideShareByCityWeek, vw_NavigationByWeek.',
])
doc.add_paragraph()
note(
    'Without this, text sorts alphabetically. While "25-01" through "25-52" happen to sort correctly, '
    'data spanning years (e.g., "24-50" vs "25-01") would be misordered.'
)

# --- 3.4 ---
doc.add_heading('3.4 DAX Measures', level=2)
doc.add_paragraph(
    'Create a _measures table (Enter Data > empty table > rename to "_measures") to hold all DAX measures. '
    'All "latest week" KPIs use the following pattern to handle weeks with NULL values:'
)
doc.add_paragraph()
doc.add_heading('Latest-Week Pattern', level=3)
code('''Latest [Field] =
CALCULATE(
    AVERAGE(TableName[field]),
    FILTER(
        ALL(TableName),
        TableName[yearweek_sort] = MAXX(
            FILTER(ALL(TableName), NOT(ISBLANK(TableName[field]))),
            TableName[yearweek_sort]
        )
    )
)''')
doc.add_paragraph()
note(
    'This pattern is required because the most recent week may have NULLs for some columns. '
    'TOPN-based and VAR-based approaches do not work in Power BI Report Server (September 2025). '
    'The FILTER + ALL + MAXX + NOT(ISBLANK) pattern is confirmed working.'
)
doc.add_paragraph()

doc.add_heading('Page 1 KPI Measures (vw_KPISummary)', level=3)
kpi_measures = [
    ('Total Responses', 'FIRSTNONBLANK(vw_KPISummary[total_responses], 1)'),
    ('Survey Weeks', 'FIRSTNONBLANK(vw_KPISummary[survey_weeks], 1)'),
    ('Cities Covered', 'FIRSTNONBLANK(vw_KPISummary[cities], 1)'),
    ('Joint Driver Pct', 'FIRSTNONBLANK(vw_KPISummary[joint_driver_pct], 1)'),
]
for name, formula in kpi_measures:
    code(f'{name} = {formula}')
doc.add_paragraph()

doc.add_heading('Page 1 Latest-Week Measures (vw_WeeklySatisfaction)', level=3)
for name, field in [('Snapp Fare Sat', 'snapp_fare_sat_avg'), ('Tapsi Fare Sat', 'tapsi_fare_sat_avg')]:
    code(f'''{name} =
CALCULATE(
    AVERAGE(vw_WeeklySatisfaction[{field}]),
    FILTER(ALL(vw_WeeklySatisfaction),
        vw_WeeklySatisfaction[yearweek_sort] = MAXX(
            FILTER(ALL(vw_WeeklySatisfaction), NOT(ISBLANK(vw_WeeklySatisfaction[{field}]))),
            vw_WeeklySatisfaction[yearweek_sort])))''')
doc.add_paragraph()

doc.add_heading('Page 2 NPS Measures (vw_WeeklyNPS)', level=3)
nps_fields = [
    ('Latest Snapp NPS', 'snapp_nps'),
    ('Latest Tapsi NPS', 'tapsi_nps'),
    ('Latest Snapp Promoter Pct', 'snapp_promoter_pct'),
    ('Latest Tapsi Promoter Pct', 'tapsi_promoter_pct'),
    ('Latest Snapp Detractor Pct', 'snapp_detractor_pct'),
    ('Latest Tapsi Detractor Pct', 'tapsi_detractor_pct'),
]
for name, field in nps_fields:
    code(f'''{name} =
CALCULATE(
    AVERAGE(vw_WeeklyNPS[{field}]),
    FILTER(ALL(vw_WeeklyNPS),
        vw_WeeklyNPS[yearweek_sort] = MAXX(
            FILTER(ALL(vw_WeeklyNPS), NOT(ISBLANK(vw_WeeklyNPS[{field}]))),
            vw_WeeklyNPS[yearweek_sort])))''')
doc.add_paragraph()

doc.add_heading('Page 4 Incentive Measures (vw_IncentiveByWeek)', level=3)
inc_fields = [
    ('Latest Snapp Incentive MRial', 'snapp_incentive_avg_mrial'),
    ('Latest Got Message Pct', 'snapp_gotmsg_pct'),
    ('Latest Participation Pct', 'snapp_participation_pct'),
    ('Latest CommFree Pct', 'snapp_commfree_pct'),
]
for name, field in inc_fields:
    code(f'''{name} =
CALCULATE(
    AVERAGE(vw_IncentiveByWeek[{field}]),
    FILTER(ALL(vw_IncentiveByWeek),
        vw_IncentiveByWeek[yearweek_sort] = MAXX(
            FILTER(ALL(vw_IncentiveByWeek), NOT(ISBLANK(vw_IncentiveByWeek[{field}]))),
            vw_IncentiveByWeek[yearweek_sort])))''')

doc.add_page_break()

# --- 3.5 PAGE 1 ---
doc.add_heading('3.5 Page 1: Executive Overview', level=2)
doc.add_paragraph('Source: vw_KPISummary, vw_WeeklySatisfaction')

doc.add_heading('Row 1: KPI Cards (8 cards, horizontal)', level=3)
add_table(
    ['Card', 'Measure', 'Format'],
    [
        ['Total Responses', 'Total Responses', 'Whole number, thousands separator'],
        ['Survey Weeks', 'Survey Weeks', 'Whole number'],
        ['Cities Covered', 'Cities Covered', 'Whole number'],
        ['Joint Driver Pct', 'Joint Driver Pct', '1 decimal + "%"'],
        ['Snapp Fare Sat', 'Snapp Fare Sat', '2 decimals'],
        ['Tapsi Fare Sat', 'Tapsi Fare Sat', '2 decimals'],
        ['Snapp NPS', 'Latest Snapp NPS', 'Whole number'],
        ['Tapsi NPS', 'Latest Tapsi NPS', 'Whole number'],
    ]
)
doc.add_paragraph()

visual_spec('Row 2: Weekly Satisfaction Trends (Line Chart)', [
    ['Visual Type', 'Line Chart (full width)'],
    ['Table', 'vw_WeeklySatisfaction'],
    ['X-Axis', 'yearweek (Sort by yearweek_sort ascending)'],
    ['Y-Axis', 'snapp_fare_sat_avg, snapp_income_sat_avg, snapp_req_sat_avg, tapsi_fare_sat_avg'],
    ['Rename', '"Snapp Fare", "Snapp Income", "Snapp Request", "Tapsi Fare"'],
    ['Y-Axis Range', 'Min: 1, Max: 5'],
    ['Colors', 'Snapp shades: #00BFA5; Tapsi: #FF6F00'],
    ['Title', 'Weekly Satisfaction Trends (1-5 Scale)'],
])

visual_spec('Row 3: Weekly Response Count (Column Chart)', [
    ['Visual Type', 'Clustered Column Chart (full width)'],
    ['Table', 'vw_WeeklySatisfaction'],
    ['X-Axis', 'yearweek (Sort by yearweek_sort ascending)'],
    ['Y-Axis', 'response_count'],
    ['Color', '#2E75B6'],
    ['Title', 'Weekly Response Count'],
])

# --- 3.6 PAGE 2 ---
doc.add_heading('3.6 Page 2: NPS Deep-Dive', level=2)
doc.add_paragraph('Source: vw_WeeklyNPS')

doc.add_heading('Row 1: KPI Cards (6 cards)', level=3)
add_table(
    ['Card', 'Measure', 'Format'],
    [
        ['Snapp NPS', 'Latest Snapp NPS', '1 decimal'],
        ['Tapsi NPS', 'Latest Tapsi NPS', '1 decimal'],
        ['Snapp Promoters', 'Latest Snapp Promoter Pct', '1 decimal + "%"'],
        ['Tapsi Promoters', 'Latest Tapsi Promoter Pct', '1 decimal + "%"'],
        ['Snapp Detractors', 'Latest Snapp Detractor Pct', '1 decimal + "%"'],
        ['Tapsi Detractors', 'Latest Tapsi Detractor Pct', '1 decimal + "%"'],
    ]
)
doc.add_paragraph()

visual_spec('Row 2: Weekly NPS Trend (Line Chart)', [
    ['Visual Type', 'Line Chart (full width)'],
    ['Table', 'vw_WeeklyNPS'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Y-Axis', 'snapp_nps, tapsi_nps'],
    ['Rename', '"Snapp NPS", "Tapsi NPS"'],
    ['Colors', 'Snapp: #00BFA5, Tapsi: #FF6F00'],
    ['Reference Line', 'Constant at Y=0, dashed gray (Analytics pane)'],
    ['Title', 'Weekly NPS Trend'],
])

visual_spec('Row 3 Left: Snapp NPS Decomposition', [
    ['Visual Type', 'Stacked Column Chart (half width)'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Y-Axis', 'snapp_promoter_pct, snapp_passive_pct, snapp_detractor_pct'],
    ['Rename', '"Promoters", "Passives", "Detractors"'],
    ['Colors', 'Promoters: #4CAF50, Passives: #FFC107, Detractors: #F44336'],
    ['Title', 'Snapp NPS Decomposition'],
])

visual_spec('Row 3 Right: Tapsi NPS Decomposition', [
    ['Visual Type', 'Stacked Column Chart (half width)'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Y-Axis', 'tapsi_promoter_pct, tapsi_passive_pct, tapsi_detractor_pct'],
    ['Rename', '"Promoters", "Passives", "Detractors"'],
    ['Colors', 'Same as Snapp decomposition'],
    ['Title', 'Tapsi NPS Decomposition'],
])

# --- 3.7 PAGE 3 ---
doc.add_heading('3.7 Page 3: Satisfaction Deep-Dive', level=2)
doc.add_paragraph('Source: vw_WeeklySatisfaction, vw_SatisfactionByCityWeek')

visual_spec('Row 1: Satisfaction Gap (Line Chart)', [
    ['Visual Type', 'Line Chart (full width)'],
    ['Table', 'vw_WeeklySatisfaction'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Y-Axis', 'fare_sat_gap, income_sat_gap, req_sat_gap'],
    ['Rename', '"Fare Gap", "Income Gap", "Request Gap"'],
    ['Reference Line', 'Constant at Y=0, dashed gray'],
    ['Title', 'Snapp - Tapsi Satisfaction Gap (positive = Snapp ahead)'],
])

visual_spec('Row 2 Left: Snapp Satisfaction Trends', [
    ['Visual Type', 'Line Chart (half width)'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Y-Axis', 'snapp_fare_sat_avg, snapp_income_sat_avg, snapp_req_sat_avg, snapp_overall_sat_avg'],
    ['Rename', '"Fare", "Income", "Request Count", "Overall"'],
    ['Y-Axis Range', 'Min: 1, Max: 5'],
    ['Title', 'Snapp Satisfaction Trends'],
])

visual_spec('Row 2 Right: Tapsi Satisfaction Trends', [
    ['Visual Type', 'Line Chart (half width)'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Y-Axis', 'tapsi_fare_sat_avg, tapsi_income_sat_avg, tapsi_req_sat_avg, tapsi_overall_sat_avg'],
    ['Rename', '"Fare", "Income", "Request Count", "Overall"'],
    ['Y-Axis Range', 'Min: 1, Max: 5'],
    ['Title', 'Tapsi Satisfaction Trends'],
])

visual_spec('Row 3: City x Week Heatmap (Matrix)', [
    ['Visual Type', 'Matrix'],
    ['Table', 'vw_SatisfactionByCityWeek'],
    ['Rows', 'city'],
    ['Columns', 'yearweek (sorted by yearweek_sort)'],
    ['Values', 'snapp_overall_sat (Average)'],
    ['Conditional Formatting', 'Background color scale: Red (#E53935) at 1 > Yellow (#FDD835) at 3 > Green (#43A047) at 5'],
    ['How To', 'Right-click value > Conditional formatting > Background color > Color scale'],
    ['Title', 'Snapp Overall Satisfaction by City x Week'],
])

# --- 3.8 PAGE 4 ---
doc.add_heading('3.8 Page 4: Incentive Analysis', level=2)
doc.add_paragraph('Source: vw_IncentiveByWeek')

doc.add_heading('Row 1: KPI Cards (4 cards)', level=3)
add_table(
    ['Card', 'Measure', 'Format'],
    [
        ['Snapp Incentive (M Rial)', 'Latest Snapp Incentive MRial', '2 decimals'],
        ['Got Message %', 'Latest Got Message Pct', '1 decimal + "%"'],
        ['Participation %', 'Latest Participation Pct', '1 decimal + "%"'],
        ['Commission-Free %', 'Latest CommFree Pct', '1 decimal + "%"'],
    ]
)
doc.add_paragraph()

visual_spec('Row 2: Incentive vs Satisfaction (Combo Chart)', [
    ['Visual Type', 'Line and Clustered Column Chart (full width)'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Column', 'snapp_incentive_avg_mrial > rename "Snapp Incentive (M Rial)"'],
    ['Line', 'snapp_inc_sat_avg > rename "Incentive Satisfaction" (secondary Y-axis, 1-5)'],
    ['Bar Color', '#2E75B6'],
    ['Line Color', '#00BFA5'],
    ['Title', 'Incentive Amount vs Satisfaction'],
])

visual_spec('Row 3 Left: Engagement Funnel (Line Chart)', [
    ['Visual Type', 'Line Chart (half width)'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Y-Axis', 'snapp_gotmsg_pct, snapp_participation_pct'],
    ['Rename', '"Got Message %", "Participated %"'],
    ['Colors', '#42A5F5 (light blue), #1565C0 (dark blue)'],
    ['Note', 'Both metrics are % of ALL drivers (same denominator). Gap = drop-off.'],
    ['Title', 'Incentive Engagement Funnel'],
])

visual_spec('Row 3 Right: Commission-Free (Area Chart)', [
    ['Visual Type', 'Area Chart (half width)'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Y-Axis', 'snapp_commfree_pct > rename "Commission-Free %"'],
    ['Color', '#00BFA5'],
    ['Title', 'Commission-Free Ride Share'],
])

# --- 3.9 PAGE 5 ---
doc.add_heading('3.9 Page 5: Operations & Demographics', level=2)
doc.add_paragraph('Source: vw_NavigationByWeek, vw_Demographics, vw_RideShareByCityWeek')

visual_spec('Row 1: Navigation App Trend (Stacked Area)', [
    ['Visual Type', 'Stacked Area Chart (full width)'],
    ['Table', 'vw_NavigationByWeek'],
    ['X-Axis', 'yearweek (sorted by yearweek_sort)'],
    ['Y-Axis', 'pct'],
    ['Legend', 'nav_app'],
    ['Title', 'Navigation App Usage Over Time'],
])

doc.add_heading('Row 2: Demographics (3 Donut Charts)', level=3)
add_table(
    ['Position', 'Filter (dimension =)', 'Legend', 'Values', 'Title'],
    [
        ['Left', 'cooperation_type', 'category', 'n', 'Cooperation Type'],
        ['Center', 'age_group', 'category', 'n', 'Age Distribution'],
        ['Right', 'driver_type', 'category', 'n', 'Joint vs Exclusive'],
    ]
)
doc.add_paragraph()
note('For each donut: use vw_Demographics, add a visual-level filter where dimension = the specified value.')
doc.add_paragraph()

visual_spec('Row 3: Ride Share by City x Week (Matrix)', [
    ['Visual Type', 'Matrix'],
    ['Table', 'vw_RideShareByCityWeek'],
    ['Rows', 'city'],
    ['Columns', 'yearweek (sorted by yearweek_sort)'],
    ['Values', 'snapp_ride_share_pct (Average)'],
    ['Conditional Formatting', 'Background: White (#FFFFFF) at 50% > Green (#43A047) at 100%'],
    ['Title', 'Snapp Ride Share % by City x Week'],
])

doc.add_page_break()

# ============================================================
# PHASE 4: PUBLISHING & MAINTENANCE
# ============================================================
doc.add_heading('Phase 4: Publishing & Maintenance', level=1)

doc.add_heading('4.1 Publish to Power BI Service', level=2)
steps([
    'Save the .pbix file.',
    'Click File > Publish > Publish to Power BI.',
    'Sign in with your organizational account (nasim.rajabi@snapp.cab).',
    'Select the target workspace.',
    'Open the published report via the link provided.',
])

doc.add_heading('4.2 Scheduled Refresh (On-Premises Gateway)', level=2)
doc.add_paragraph(
    'The SQL Server is on-premises (192.168.18.37), so a Power BI Gateway is required.'
)
steps([
    'Install "On-premises data gateway" on a machine with access to the SQL Server.',
    'Configure the gateway with your Power BI account.',
    'In Power BI Service: Settings > Manage gateways > Add SQL Server data source.',
    'Dataset Settings > Scheduled refresh > Toggle ON > set frequency (e.g., Daily 08:00).',
])

doc.add_heading('4.3 Adding New Survey Weeks', level=2)
doc.add_paragraph('When new weekly survey data arrives:')
steps([
    'Process the raw survey data through the survey analysis pipeline (generates CSVs in processed/).',
    'Run: python PowerBI/load_to_database.py (incremental mode — only new recordIDs are inserted).',
    'The SQL views automatically reflect new data (no view changes needed).',
    'For HTML dashboards: click Refresh on http://localhost:8800, or re-run build scripts.',
    'For Power BI: click Refresh in Power BI Desktop, or let scheduled refresh pick it up.',
])
doc.add_paragraph()
note(
    'The script compares recordIDs in the CSV against those already in the database. '
    'Only rows with new recordIDs are inserted. This makes it safe to re-run without duplicating data.'
)

doc.add_page_break()

# ============================================================
# APPENDIX A: FILE INVENTORY
# ============================================================
doc.add_heading('Appendix A: File Inventory', level=1)

doc.add_heading('Sources/ folder', level=2)
add_table(
    ['File', 'Purpose'],
    [
        ['column_rename.xlsx', 'Column name mapping reference (original > English names)'],
        ['column_rename_mapping.json', 'JSON mapping used by load_to_database.py'],
    ]
)
doc.add_paragraph()

doc.add_heading('PowerBI/ folder', level=2)
add_table(
    ['File', 'Purpose'],
    [
        ['Dashboard.pbix', 'Power BI report file (5 pages, connected to SQL views)'],
        ['PowerBI_Dashboard_Guide.docx', 'This guide document'],
        ['create_driver_survey_table.sql', 'Schema reference: DDL for 6 base tables with column types (optional — load_to_database.py creates tables automatically)'],
        ['create_views.sql', 'DDL for 21 SQL views (1 base + 20 analytics)'],
        ['load_to_database.py', 'Creates 6 base tables and bulk-loads CSV data into SQL Server'],
        ['build_dashboard.py', 'Generates DriverSurvey HTML dashboard from SQL views'],
        ['build_routine_dashboard.py', 'Generates RoutineAnalysis HTML heatmap dashboard'],
        ['serve_dashboards.py', 'Python web server (port 8800) for both HTML dashboards'],
        ['start_server.bat', 'Batch launcher for the web server'],
        ['DriverSurvey_Dashboard.html', 'Generated HTML dashboard (Plotly charts)'],
        ['RoutineAnalysis_Dashboard.html', 'Generated HTML dashboard (heatmap tables)'],
        ['build_docx_guide.py', 'Python script to regenerate this guide document'],
    ]
)

doc.add_page_break()

# ============================================================
# APPENDIX B: COLOR PALETTE
# ============================================================
doc.add_heading('Appendix B: Color Palette', level=1)
add_table(
    ['Element', 'Color', 'Hex Code', 'Usage'],
    [
        ['Snapp Primary', 'Teal/Green', '#00BFA5', 'All Snapp metrics'],
        ['Tapsi Primary', 'Orange', '#FF6F00', 'All Tapsi metrics'],
        ['Neutral Blue', 'Blue', '#2E75B6', 'Non-platform visuals'],
        ['NPS Promoter', 'Green', '#4CAF50', 'NPS promoter segment'],
        ['NPS Passive', 'Amber', '#FFC107', 'NPS passive segment'],
        ['NPS Detractor', 'Red', '#F44336', 'NPS detractor segment'],
        ['Sat High (5)', 'Green', '#43A047', 'Heatmap max'],
        ['Sat Mid (3)', 'Yellow', '#FDD835', 'Heatmap midpoint'],
        ['Sat Low (1)', 'Red', '#E53935', 'Heatmap min'],
        ['Background', 'White', '#FFFFFF', 'Page background'],
        ['Text', 'Dark Gray', '#333333', 'Labels and titles'],
    ]
)

doc.add_paragraph()
doc.add_heading('Axis Formatting Rules', level=2)
bullets([
    'Satisfaction axes: Y min=1, max=5.',
    'NPS axes: Constant reference line at Y=0 (dashed gray).',
    'Percentage axes: 0-1 decimal places.',
    'YearWeek X-axis: Categorical type, sorted by yearweek_sort ascending.',
])

doc.add_page_break()

# ============================================================
# APPENDIX C: TROUBLESHOOTING
# ============================================================
doc.add_heading('Appendix C: Troubleshooting', level=1)

issues = [
    (
        'DAX KPI cards show (Blank)',
        'The latest week may have NULL values for that field. Use the FILTER + ALL + MAXX + '
        'NOT(ISBLANK) pattern (see Section 3.4). Do NOT use TOPN or VAR-based equality comparisons '
        '-- they fail in Power BI Report Server (September 2025).'
    ),
    (
        'YearWeek axis shows 2500, 2520, 2540... (continuous)',
        'The yearweek column is being treated as numeric. Ensure the SQL views cast yearweek as '
        'NVARCHAR/VARCHAR text. Re-run Sources/create_views.sql if needed.'
    ),
    (
        'YearWeek axis is out of order',
        'Set Sort-by-Column: Data view > click yearweek column > Column tools > Sort by Column > '
        'yearweek_sort. Repeat for all 6 tables.'
    ),
    (
        'Cannot connect to SQL Server',
        'Ensure VPN to Snapp data center is active. Server: 192.168.18.37, Database: Cab_Studies, '
        'Schema: Cab.'
    ),
    (
        'HTML dashboard refresh fails',
        'Check that the VPN is active and the SQL Server is reachable. The serve_dashboards.py uses '
        'subprocess to call build scripts -- check console output for errors.'
    ),
    (
        'Power BI scheduled refresh fails',
        'Requires an on-premises gateway with network access to 192.168.18.37. Check gateway status '
        'in Power BI Service > Settings > Manage gateways.'
    ),
    (
        'GotBonus chart labels cut off in HTML dashboard',
        'The min chart height was increased from 200px to 280px in build_dashboard.py. '
        'Regenerate the HTML: python PowerBI/build_dashboard.py'
    ),
    (
        'Routine Analysis shows col_N or numbers in headers',
        'The build_routine_dashboard.py uses display_map to separate internal unique names from '
        'display names. If new sheets have similar issues, check the column deduplication logic.'
    ),
]

for title, fix in issues:
    doc.add_heading(title, level=3)
    doc.add_paragraph(fix)
    doc.add_paragraph()

doc.add_paragraph('--- End of Guide ---')

# Save
output_path = 'D:/Work/Driver Survey/PowerBI/PowerBI_Dashboard_Guide.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
