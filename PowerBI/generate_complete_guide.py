import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── rendering helpers ────────────────────────────────────────────────────────
def set_font(run, size=10, bold=False, color=None, name='Arial'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Arial'
    return p

def body(text='', bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def code_line(text, is_name_line=False):
    """Render one line of DAX as a code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    if is_name_line:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 70, 127)
    else:
        run.font.color.rgb = RGBColor(30, 30, 120)
    return p

def dax_measure(dax_str):
    """Render a complete DAX measure string. First line (MeasureName =) is bolded."""
    # Strip RA<n> prefix from the measure name line
    lines = dax_str.split('\n')
    lines[0] = re.sub(r'^RA\d+\s+', '', lines[0])
    dax_str = '\n'.join(lines)
    doc.add_paragraph()  # spacer before
    lines = dax_str.split('\n')
    for i, line in enumerate(lines):
        if line.strip():
            code_line(line, is_name_line=(i == 0))
    doc.add_paragraph()  # spacer after

def bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + indent * 0.2)
    run = p.add_run(text)
    set_font(run, size=9.5)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run('ℹ  ' + text)
    set_font(run, size=9, color=(100, 100, 100))
    return p

def col_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.name = 'Arial'
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.cell(r + 1, c)
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)
            cell.paragraphs[0].runs[0].font.name = 'Courier New'
    return t

# ── DAX string builders ──────────────────────────────────────────────────────
def clean_name(name):
    """Strip 'RA<n> ' prefix so measure names in the doc match Power BI (e.g. 'Snapp NPS')."""
    return re.sub(r'^RA\d+\s+', '', name)

def yw_var(view):
    return (
        f'VAR SelYearWeek = IF(HASONEVALUE({view}[yearweek]),\n'
        f'    SELECTEDVALUE({view}[yearweek]),\n'
        f'    CALCULATE(MAX({view}[yearweek]), ALL({view})))'
    )

def metric_dax(name, view, metric_expr, n_col='n', extra_filters=None):
    """Standard metric with N-gate."""
    name = clean_name(name)
    filters = f', {view}[yearweek] = SelYearWeek'
    if extra_filters:
        filters += f',\n        {extra_filters}'
    n_filters = f'{view}[yearweek] = SelYearWeek'
    if extra_filters:
        n_filters += f',\n        {extra_filters}'
    return (
        f'{name} =\n'
        f'{yw_var(view)}\n'
        f'VAR MinN = [Min N Cutoff Value]\n'
        f'RETURN IF(\n'
        f'    CALCULATE(SUM({view}[{n_col}]), {n_filters}) >= MinN,\n'
        f'    CALCULATE({metric_expr}{filters}),\n'
        f'    BLANK())'
    )

def count_dax(name, view, count_col='n', extra_filters=None):
    """Simple count — no N-gate."""
    name = clean_name(name)
    n_filter = f'{view}[yearweek] = SelYearWeek'
    if extra_filters:
        n_filter += f', {extra_filters}'
    return (
        f'{name} =\n'
        f'{yw_var(view)}\n'
        f'RETURN CALCULATE(SUM({view}[{count_col}]), {n_filter})'
    )

def wow_dax(name, view, metric_expr, n_col='n', extra_filters=None):
    """WoW delta measure."""
    name = clean_name(name)
    base_filter = f'{view}[yearweek] = SelYearWeek'
    prev_filter = f'{view}[yearweek] = PrevYearWeek'
    if extra_filters:
        base_filter += f',\n        {extra_filters}'
        prev_filter += f',\n        {extra_filters}'
    n_filter = f'{view}[yearweek] = SelYearWeek'
    if extra_filters:
        n_filter += f', {extra_filters}'
    return (
        f'{name} =\n'
        f'{yw_var(view)}\n'
        f'VAR PrevYearWeek = CALCULATE(MAX({view}[yearweek]),\n'
        f'    ALL({view}), {view}[yearweek] < SelYearWeek)\n'
        f'VAR MinN = [Min N Cutoff Value]\n'
        f'VAR CurrVal = CALCULATE({metric_expr}, {base_filter})\n'
        f'VAR PrevVal = CALCULATE({metric_expr}, {prev_filter})\n'
        f'RETURN IF(\n'
        f'    CALCULATE(SUM({view}[{n_col}]), {n_filter}) >= MinN,\n'
        f'    CurrVal - PrevVal,\n'
        f'    BLANK())'
    )

# ════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════
heading('Driver Survey – Power BI Routine Analysis Guide v5', level=1)
heading('Complete DAX Reference for All 17 RA Views', level=2)
body('Database: Cab_Studies  •  Server: 192.168.18.37  •  Schema: [Cab]', color=(80, 80, 80))
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 1. SETUP
# ════════════════════════════════════════════════════════════════════════════
heading('1. Power BI Setup', level=1)
body('Connection', bold=True)
bullet('Home → Get Data → SQL Server')
bullet('Server: 192.168.18.37  |  Database: Cab_Studies')
bullet('Import mode (not DirectQuery)')
body('Import all 17 RA views', bold=True)
for v in [
    'vw_RA_SatReview', 'vw_RA_CitiesOverview', 'vw_RA_RideShare', 'vw_RA_PersonaPartTime',
    'vw_RA_IncentiveAmounts', 'vw_RA_IncentiveDuration', 'vw_RA_Persona',
    'vw_RA_CommFree', 'vw_RA_CSRare', 'vw_RA_NavReco',
    'vw_RA_IncentiveTypes', 'vw_RA_IncentiveUnsatCity', 'vw_RA_IncentiveUnsatNational',
    'vw_RA_Navigation', 'vw_RA_Referral', 'vw_RA_TapsiInactivity', 'vw_RA_LuckyWheel',
]:
    bullet(v, indent=1)
body('Do NOT create relationships between views – each is self-contained.', color=(160, 0, 0))
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 2. GLOBAL HELPERS
# ════════════════════════════════════════════════════════════════════════════
heading('2. Global Helper Measures & Parameter Table', level=1)

body('Min N Cutoff parameter table', bold=True)
note('Create a blank table via Enter Data named "Min N Cutoff" with one column "Min N Cutoff" and one row: 30. '
     'Users can edit the value in the table to change the global threshold.')

body('Helper measures (create in a "Measures" blank table):', bold=True)
dax_measure(
    'Min N Cutoff Value =\n'
    'SELECTEDVALUE(\'Min N Cutoff\'[Min N Cutoff], 0)'
)
note('All metric measures reference [Min N Cutoff Value]. Changing the table value updates all pages at once.')

body('Year-Week sorting', bold=True)
note('In the Model view, select yearweek (TEXT column) → Sort by Column → yearweek_sort (INT). '
     'Apply this to every imported view.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# RA-1  vw_RA_SatReview
# ════════════════════════════════════════════════════════════════════════════
heading('3. RA-1 – Satisfaction & Participation Review', level=1)
body('View: vw_RA_SatReview  |  Excel Page: #3', bold=True)
body('Incentive participation rates and satisfaction scores (1–5) for Snapp & Tapsi. '
     'Filter cooperation_type: All Drivers / Part-Time / Full-Time.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek', 'TEXT', '"26-01" style'),
        ('yearweek_sort', 'INT', 'for axis ordering'),
        ('weeknumber', 'INT', ''),
        ('city', 'TEXT', ''),
        ('cooperation_type', 'TEXT', 'All Drivers / Part-Time / Full-Time'),
        ('n', 'INT', 'all respondents in city'),
        ('n_joint', 'INT', 'joint drivers'),
        ('Part_pct_Snapp', 'FLOAT', '% participated in Snapp incentive'),
        ('Part_pct_Jnt_Snapp', 'FLOAT', '% joint participated in Snapp'),
        ('Part_pct_Jnt_Tapsi', 'FLOAT', '% joint participated in Tapsi'),
        ('Part_GotMsg_pct_Snapp', 'FLOAT', '% participated among who got Snapp msg'),
        ('Part_GotMsg_pct_Jnt_Snapp', 'FLOAT', ''),
        ('Part_GotMsg_pct_Jnt_Tapsi', 'FLOAT', ''),
        ('Incentive_Sat_Snapp', 'FLOAT', 'avg 1-5'),
        ('Incentive_Sat_Jnt_Snapp', 'FLOAT', ''),
        ('Incentive_Sat_Jnt_Tapsi', 'FLOAT', ''),
        ('Fare_Sat_Snapp / Jnt_Snapp / Jnt_Tapsi', 'FLOAT', 'avg 1-5'),
        ('Request_Sat_Snapp / Jnt_Snapp / Jnt_Tapsi', 'FLOAT', 'avg 1-5'),
        ('Income_Sat_Snapp / Jnt_Snapp / Jnt_Tapsi', 'FLOAT', 'avg 1-5'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)
note('Use a page-level filter on cooperation_type rather than a slicer to avoid cross-filtering other visuals.')

V = 'vw_RA_SatReview'
dax_measure(count_dax('RA1 N All', V))
dax_measure(count_dax('RA1 N Joint', V, count_col='n_joint'))
dax_measure(metric_dax('RA1 Part% Snapp', V, f'AVERAGE({V}[Part_pct_Snapp])'))
dax_measure(metric_dax('RA1 Part% Jnt Snapp', V, f'AVERAGE({V}[Part_pct_Jnt_Snapp])', n_col='n_joint'))
dax_measure(metric_dax('RA1 Part% Jnt Tapsi', V, f'AVERAGE({V}[Part_pct_Jnt_Tapsi])', n_col='n_joint'))
dax_measure(metric_dax('RA1 GotMsg Part% Snapp', V, f'AVERAGE({V}[Part_GotMsg_pct_Snapp])'))
dax_measure(metric_dax('RA1 Incentive Sat Snapp', V, f'AVERAGE({V}[Incentive_Sat_Snapp])'))
dax_measure(metric_dax('RA1 Incentive Sat Jnt Snapp', V, f'AVERAGE({V}[Incentive_Sat_Jnt_Snapp])', n_col='n_joint'))
dax_measure(metric_dax('RA1 Incentive Sat Jnt Tapsi', V, f'AVERAGE({V}[Incentive_Sat_Jnt_Tapsi])', n_col='n_joint'))
dax_measure(metric_dax('RA1 Fare Sat Snapp', V, f'AVERAGE({V}[Fare_Sat_Snapp])'))
dax_measure(metric_dax('RA1 Request Sat Snapp', V, f'AVERAGE({V}[Request_Sat_Snapp])'))
dax_measure(metric_dax('RA1 Income Sat Snapp', V, f'AVERAGE({V}[Income_Sat_Snapp])'))
dax_measure(wow_dax('RA1 WoW Incentive Sat Snapp', V, f'AVERAGE({V}[Incentive_Sat_Snapp])'))
note('Duplicate Incentive/Fare/Request/Income Sat measures for Jnt_Snapp and Jnt_Tapsi columns (use n_joint as MinN gate).')
note('Visual: Matrix – city × satisfaction column, week slicer. Cards for national averages. Line for WoW.')

# ════════════════════════════════════════════════════════════════════════════
# RA-2  vw_RA_CitiesOverview
# ════════════════════════════════════════════════════════════════════════════
heading('4. RA-2 – Cities Overview', level=1)
body('View: vw_RA_CitiesOverview  |  Excel Page: #12', bold=True)
body('Three independent groups: E (all), F (joint), G (tapsi_LOC > 0). No cooperation_type split.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('E_n', 'INT', 'all respondents'),
        ('F_n', 'INT', 'joint drivers'),
        ('G_n', 'INT', 'drivers with tapsi LOC > 0'),
        ('pct_Joint', 'FLOAT', '% who are joint (base E_n)'),
        ('pct_Dual_SU', 'FLOAT', '% with tapsi LOC > 0 (base E_n)'),
        ('AvgLOC_All_Snapp', 'FLOAT', 'avg Snapp LOC (base E_n)'),
        ('GotMsg_All_Snapp', 'FLOAT', '% got Snapp msg (base E_n)'),
        ('AvgLOC_Joint_Snapp', 'FLOAT', 'avg Snapp LOC (base F_n)'),
        ('GotMsg_Joint_Snapp', 'FLOAT', '% joint got Snapp msg (base F_n)'),
        ('GotMsg_Joint_Cmpt', 'FLOAT', '% joint got Tapsi msg (base F_n)'),
        ('AvgLOC_Joint_Cmpt', 'FLOAT', 'avg Tapsi LOC (base F_n)'),
        ('AvgLOC_Joint_Cmpt_SU', 'FLOAT', 'avg Tapsi LOC (base G_n)'),
        ('GotMsg_Joint_Cmpt_SU', 'FLOAT', '% G-group got Tapsi msg (base G_n)'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_CitiesOverview'
dax_measure(count_dax('RA2 E_n', V, count_col='E_n'))
dax_measure(count_dax('RA2 F_n', V, count_col='F_n'))
dax_measure(count_dax('RA2 G_n', V, count_col='G_n'))
dax_measure(metric_dax('RA2 pct Joint', V, f'AVERAGE({V}[pct_Joint])'))
dax_measure(metric_dax('RA2 pct Dual SU', V, f'AVERAGE({V}[pct_Dual_SU])'))
dax_measure(metric_dax('RA2 AvgLOC All Snapp', V, f'AVERAGE({V}[AvgLOC_All_Snapp])'))
dax_measure(metric_dax('RA2 GotMsg All Snapp', V, f'AVERAGE({V}[GotMsg_All_Snapp])'))
dax_measure(metric_dax('RA2 GotMsg Joint Snapp', V, f'AVERAGE({V}[GotMsg_Joint_Snapp])', n_col='F_n'))
dax_measure(metric_dax('RA2 GotMsg Joint Cmpt', V, f'AVERAGE({V}[GotMsg_Joint_Cmpt])', n_col='F_n'))
dax_measure(metric_dax('RA2 GotMsg Joint Cmpt SU', V, f'AVERAGE({V}[GotMsg_Joint_Cmpt_SU])', n_col='G_n'))
dax_measure(wow_dax('RA2 WoW pct Joint', V, f'AVERAGE({V}[pct_Joint])'))
note('Use E_n / F_n / G_n as the MinN gate for corresponding metric groups.')

# ════════════════════════════════════════════════════════════════════════════
# RA-3  vw_RA_RideShare
# ════════════════════════════════════════════════════════════════════════════
heading('5. RA-3 – Ride Share', level=1)
body('View: vw_RA_RideShare  |  Excel Page: #13', bold=True)
body('Total ride counts and Snapp/Tapsi share split by driver segment.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('total_Res', 'INT', 'all respondents'),
        ('Joint_Res', 'INT', 'joint drivers'),
        ('Ex_drivers', 'INT', 'exclusive Snapp drivers'),
        ('Total_Ride', 'FLOAT', 'Snapp + Tapsi combined'),
        ('Total_Ride_Snapp', 'FLOAT', 'all Snapp rides'),
        ('Ex_Ride_Snapp', 'FLOAT', 'exclusive driver Snapp rides'),
        ('Jnt_Snapp_Ride', 'FLOAT', 'joint driver Snapp rides'),
        ('Jnt_Tapsi_Ride', 'FLOAT', 'joint driver Tapsi rides'),
        ('All_Snapp_pct', 'FLOAT', 'Snapp share of total rides'),
        ('Ex_Drivers_Snapp_pct', 'FLOAT', 'exclusive driver share'),
        ('Jnt_at_Snapp_pct', 'FLOAT', 'joint Snapp share'),
        ('Jnt_at_Tapsi_pct', 'FLOAT', 'joint Tapsi share'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_RideShare'
dax_measure(count_dax('RA3 Total Respondents', V, count_col='total_Res'))
dax_measure(count_dax('RA3 Joint Respondents', V, count_col='Joint_Res'))
dax_measure(metric_dax('RA3 All Snapp pct', V, f'AVERAGE({V}[All_Snapp_pct])', n_col='total_Res'))
dax_measure(metric_dax('RA3 Jnt at Snapp pct', V, f'AVERAGE({V}[Jnt_at_Snapp_pct])', n_col='Joint_Res'))
dax_measure(metric_dax('RA3 Jnt at Tapsi pct', V, f'AVERAGE({V}[Jnt_at_Tapsi_pct])', n_col='Joint_Res'))
dax_measure(wow_dax('RA3 WoW All Snapp pct', V, f'AVERAGE({V}[All_Snapp_pct])', n_col='total_Res'))
note('Visual: Clustered bar for pct columns by city. Line chart for WoW trend.')

# ════════════════════════════════════════════════════════════════════════════
# RA-4  vw_RA_PersonaPartTime
# ════════════════════════════════════════════════════════════════════════════
heading('6. RA-4 – Persona Part-Time', level=1)
body('View: vw_RA_PersonaPartTime  |  Excel Page: #15 (Part-Time sub-table)', bold=True)
body('Part-Time rate and average rides per boarded driver, by city.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('total_Res', 'INT', ''),
        ('Joint_Res', 'INT', ''),
        ('Ex_drivers', 'INT', ''),
        ('PT_pct_Joint', 'FLOAT', '% Part-Time among joint'),
        ('PT_pct_Exclusive', 'FLOAT', '% Part-Time among exclusive'),
        ('RidePerBoarded_Snapp', 'FLOAT', 'avg Snapp rides per joint driver'),
        ('RidePerBoarded_Tapsi', 'FLOAT', 'avg Tapsi rides per joint driver'),
        ('AvgAllRides', 'FLOAT', 'avg Snapp rides across all respondents'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_PersonaPartTime'
dax_measure(metric_dax('RA4 PT pct Joint', V, f'AVERAGE({V}[PT_pct_Joint])', n_col='Joint_Res'))
dax_measure(metric_dax('RA4 PT pct Exclusive', V, f'AVERAGE({V}[PT_pct_Exclusive])', n_col='Ex_drivers'))
dax_measure(metric_dax('RA4 RidePerBoarded Snapp', V, f'AVERAGE({V}[RidePerBoarded_Snapp])', n_col='Joint_Res'))
dax_measure(metric_dax('RA4 RidePerBoarded Tapsi', V, f'AVERAGE({V}[RidePerBoarded_Tapsi])', n_col='Joint_Res'))
dax_measure(metric_dax('RA4 AvgAllRides', V, f'AVERAGE({V}[AvgAllRides])', n_col='total_Res'))

# ════════════════════════════════════════════════════════════════════════════
# RA-5  vw_RA_IncentiveAmounts
# ════════════════════════════════════════════════════════════════════════════
heading('7. RA-5 – Incentive Amounts (Long Format)', level=1)
body('View: vw_RA_IncentiveAmounts  |  Excel Pages: #1 (Snapp), #2 (Tapsi)', bold=True)
body('Long format: one row per city × incentive_range × week. '
     'Tapsi rows include joint drivers only. Add a page-level filter on platform.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('platform', 'TEXT', 'Snapp or Tapsi'),
        ('incentive_range', 'TEXT', 'bucket label e.g. "<20k"'),
        ('incentive_range_sort', 'INT', 'for column ordering'),
        ('n_range', 'INT', 'count in this bucket'),
        ('n_total', 'INT', 'total respondents for city × week'),
        ('pct', 'FLOAT', 'n_range / n_total × 100'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)
note('Sort incentive_range by incentive_range_sort in the Model view.')

V = 'vw_RA_IncentiveAmounts'

dax_measure(
    f'RA5 n Total (Snapp) =\n'
    f'{yw_var(V)}\n'
    f'RETURN CALCULATE(MAX({V}[n_total]),\n'
    f'    {V}[yearweek] = SelYearWeek,\n'
    f'    {V}[platform] = "Snapp")'
)
dax_measure(
    f'RA5 pct Range (Snapp) =\n'
    f'// Place on a Matrix: rows = city, columns = incentive_range, values = this measure\n'
    f'// Page filter: platform = "Snapp"\n'
    f'{yw_var(V)}\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'RETURN IF(\n'
    f'    CALCULATE(MAX({V}[n_total]),\n'
    f'        {V}[yearweek] = SelYearWeek,\n'
    f'        {V}[platform] = "Snapp") >= MinN,\n'
    f'    CALCULATE(SUM({V}[pct]),\n'
    f'        {V}[yearweek] = SelYearWeek,\n'
    f'        {V}[platform] = "Snapp"),\n'
    f'    BLANK())'
)
dax_measure(
    f'RA5 pct Range (Tapsi) =\n'
    f'// Page filter: platform = "Tapsi"\n'
    f'{yw_var(V)}\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'RETURN IF(\n'
    f'    CALCULATE(MAX({V}[n_total]),\n'
    f'        {V}[yearweek] = SelYearWeek,\n'
    f'        {V}[platform] = "Tapsi") >= MinN,\n'
    f'    CALCULATE(SUM({V}[pct]),\n'
    f'        {V}[yearweek] = SelYearWeek,\n'
    f'        {V}[platform] = "Tapsi"),\n'
    f'    BLANK())'
)
note('Visual: Matrix – city on rows, incentive_range on columns. Separate pages for Snapp and Tapsi.')

# ════════════════════════════════════════════════════════════════════════════
# RA-6  vw_RA_IncentiveDuration
# ════════════════════════════════════════════════════════════════════════════
heading('8. RA-6 – Incentive Duration (Long Format)', level=1)
body('View: vw_RA_IncentiveDuration  |  Excel Page: #4', bold=True)
body('How long drivers have had an active incentive. Snapp (all) and Tapsi (joint only) rows.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('platform', 'TEXT', 'Snapp or Tapsi'),
        ('duration_bucket', 'TEXT', '"Few Hours","1 Day","1_6 Days","7 Days",">7 Days"'),
        ('duration_bucket_sort', 'INT', '1–5 + 99'),
        ('n_range', 'INT', ''),
        ('n_total', 'INT', ''),
        ('pct', 'FLOAT', ''),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)
note('Sort duration_bucket by duration_bucket_sort in the Model view.')

V = 'vw_RA_IncentiveDuration'
dax_measure(
    f'RA6 n Total =\n'
    f'{yw_var(V)}\n'
    f'RETURN CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek)'
)
dax_measure(
    f'RA6 pct Bucket =\n'
    f'// Matrix: rows = city, columns = duration_bucket; add platform slicer or page filter\n'
    f'{yw_var(V)}\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'RETURN IF(\n'
    f'    CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek) >= MinN,\n'
    f'    CALCULATE(SUM({V}[pct]), {V}[yearweek] = SelYearWeek),\n'
    f'    BLANK())'
)

# ════════════════════════════════════════════════════════════════════════════
# RA-7  vw_RA_Persona
# ════════════════════════════════════════════════════════════════════════════
heading('9. RA-7 – Persona (Long Format)', level=1)
body('View: vw_RA_Persona  |  Excel Page: #15 (all demographic sub-tables)', bold=True)
body('Dimensions: Activity Type, Age Group, Education, Marital Status, Gender, Cooperation Type. '
     'Use a slicer on dimension to switch between them.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('dimension', 'TEXT', '"Activity Type", "Age Group", etc.'),
        ('category', 'TEXT', 'bucket value within the dimension'),
        ('category_sort', 'INT', 'for axis ordering'),
        ('n', 'INT', 'count in this category'),
        ('n_total', 'INT', 'total respondents for city × week × dimension'),
        ('pct', 'FLOAT', 'n / n_total × 100'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)
note('Sort category by category_sort in the Model view.')

V = 'vw_RA_Persona'
dax_measure(
    f'RA7 n Total =\n'
    f'{yw_var(V)}\n'
    f'RETURN CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek)'
)
dax_measure(
    f'RA7 pct Category =\n'
    f'// Matrix: rows = city, columns = category; dimension slicer selects which breakdown\n'
    f'{yw_var(V)}\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'RETURN IF(\n'
    f'    CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek) >= MinN,\n'
    f'    CALCULATE(SUM({V}[pct]), {V}[yearweek] = SelYearWeek),\n'
    f'    BLANK())'
)

# ════════════════════════════════════════════════════════════════════════════
# RA-8  vw_RA_CommFree
# ════════════════════════════════════════════════════════════════════════════
heading('10. RA-8 – Commission-Free Incentive', level=1)
body('View: vw_RA_CommFree  |  Excel Page: #18 (Snapp + Tapsi)', bold=True)
body('Mirrors Excel sheet #18 exactly. UNION ALL of two row sets per (yearweek, city):')
bullet('SNAPP block (platform="Snapp", platform_sort=1) — denominator E = ALL drivers in city.')
bullet('TAPSI block (platform="Tapsi", platform_sort=2) — denominator E = active_joint=1 drivers in city.')
bullet('Filter [platform] in every measure (Excel-style hardcoding) rather than relying on a slicer, '
       'so the same Matrix can show both side-by-side.')
bullet('IMPORTANT: incentive category (Money / Free-Commission / Money & Free-commission) is RE-DERIVED '
       'in the view from the WideMain binary type columns to match Excel exactly. '
       'ShortMain.snapp_incentive_category / tapsi_incentive_category in the database are NOT used here — '
       'they omit "Pay After Income" from the Money group, which would cause ~10 drivers/city '
       'to be miscategorized vs Excel.')

body('Excel cell → SQL column map:', bold=True)
col_table(
    ['Excel cell', 'Excel formula (Tapsi top / Snapp bottom)', 'SQL column'],
    [
        ('D', 'hardcoded population (not in survey)', '— (use DAX constant)'),
        ('E',  'COUNTIFS(BK=city, BU=1) / COUNTIFS(BK=city)',  'n'),
        ('F',  '+AJ=Yes / +I=Yes',                              'Who_Got_Message'),
        ('G',  '+CP=Money / +CO=Money',                         'GotMsg_Money'),
        ('H',  '+CP=Free-Commission / +CO=Free-Commission',     'GotMsg_FreeComm'),
        ('I',  '+CP=Money & Free-commission',                   'GotMsg_Money_FreeComm'),
        ('L',  'COUNTIFS(BK, BY>0) / COUNTIFS(BK, CA>0)',       'Free_Comm_Drivers'),
        ('M',  'F/E',                                           'pct_Got_Message'),
        ('N',  '(H+I)/F',                                       'pct_FreeComm_Message'),
        ('O',  '(BU=1, BY>0) / E   |   (CA>0)/E',               'pct_Free_Comm_Ride'),
        ('P',  '+AY=Yes (or +X=Yes for Snapp), CP=Money / G',   'pct_Part_Money'),
        ('Q',  'same with cat=Free-Commission / H',             'pct_Part_FreeComm'),
        ('R',  'same with cat=Money & Free-commission / I',     'pct_Part_Money_FreeComm'),
        ('S',  'participated AND TYPE=Income Guarantee / GotMsg_IncGuar',  'pct_Part_IncGuar'),
        ('T',  'participated AND TYPE=Pay After Ride / GotMsg_PayRide',    'pct_Part_PayRide'),
        ('U',  'participated AND TYPE=Pay After Income / GotMsg_PayInc',   'pct_Part_PayInc'),
        ('V',  '+CP="*", AY=Yes / F   (any category, participated)', 'pct_Participated'),
        ('AA', 'SUMIF(BK, BW) / SUMIF(BK, BV)',                 'Ride_Total'),
        ('AB', 'SUMIFS(BW, BK, BY>0) / SUMIFS(BV, BK, CA>0)',   'Ride_AmongFC'),
        ('AC', 'SUMIF(BK, BY) / SUMIF(BK, CA)',                 'Ride_FreeComm'),
        ('AD', 'AC/AA',                                         'pct_FCRide_AllShare'),
        ('AE', 'AC/AB',                                         'pct_FCRide_FCShare'),
        ('W,X,Y,Z,AF,AG', 'depend on hardcoded constants D, X, Y, Z',
         'compute in DAX (see Cost Measures below)'),
    ]
)
doc.add_paragraph()

body('Columns in vw_RA_CommFree:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('platform / platform_sort', 'TEXT/INT', '"Snapp" (1) or "Tapsi" (2)'),
        ('n', 'INT', 'Excel E — Snapp: all drivers; Tapsi: active_joint=1'),
        ('Who_Got_Message', 'INT', 'Excel F — gotmessage=Yes count'),
        ('GotMsg_Money / FreeComm / Money_FreeComm', 'INT', 'Excel G/H/I — by incentive_category'),
        ('GotMsg_PayRide / EarnCF / RideCF / IncGuar / PayInc / CFSome', 'INT',
         'multi-select TYPE breakdown (not in #18 but useful)'),
        ('Free_Comm_Drivers', 'INT', 'Excel L — commfree>0 count (no joint filter)'),
        ('Part_Money / Part_FreeComm / Part_Money_FreeComm', 'INT',
         'numerators for P/Q/R — participated by category'),
        ('Part_IncGuar / Part_PayRide / Part_PayInc', 'INT',
         'numerators for S/T/U — participated by incentive TYPE (multi-select)'),
        ('Participated', 'INT', 'Excel V numerator — got msg + any category + participated'),
        ('pct_Got_Message', 'FLOAT', 'Excel M = F/E × 100'),
        ('pct_FreeComm_Message', 'FLOAT', 'Excel N = (H+I)/F × 100'),
        ('pct_Free_Comm_Ride', 'FLOAT', 'Excel O = (commfree>0)/E × 100'),
        ('pct_Part_Money / FreeComm / Money_FreeComm', 'FLOAT', 'Excel P/Q/R — by category'),
        ('pct_Part_IncGuar / PayRide / PayInc', 'FLOAT', 'Excel S/T/U — by incentive TYPE'),
        ('pct_Participated', 'FLOAT', 'Excel V'),
        ('Ride_Total', 'FLOAT', 'Excel AA = SUM(ride)'),
        ('Ride_AmongFC', 'FLOAT', 'Excel AB = SUM(ride WHERE commfree>0)'),
        ('Ride_FreeComm', 'FLOAT', 'Excel AC = SUM(commfree)'),
        ('pct_FCRide_AllShare', 'FLOAT', 'Excel AD = AC/AA × 100'),
        ('pct_FCRide_FCShare', 'FLOAT', 'Excel AE = AC/AB × 100'),
        ('Avg_CF_Rides / Avg_Total_Rides / Avg_pct_CF_RideShare', 'FLOAT',
         'legacy averages (kept for back-compat)'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)
note('All measures hard-filter [platform]. Build twin matrices (one per platform) or '
     'use a Calculation Group if you want a single matrix that switches by platform.')

V = 'vw_RA_CommFree'

# n / Who Got Message — count measures per platform
for plat in ('Snapp', 'Tapsi'):
    dax_measure(
        f'RA8 n ({plat}) =\n'
        f'{yw_var(V)}\n'
        f'RETURN CALCULATE(SUM({V}[n]),\n'
        f'    {V}[yearweek] = SelYearWeek,\n'
        f'    {V}[platform] = "{plat}")'
    )
    dax_measure(
        f'RA8 Who Got Message ({plat}) =\n'
        f'{yw_var(V)}\n'
        f'RETURN CALCULATE(SUM({V}[Who_Got_Message]),\n'
        f'    {V}[yearweek] = SelYearWeek,\n'
        f'    {V}[platform] = "{plat}")'
    )

# Helper: build a [platform]-filtered pct measure with MinN gate on n
def pct_with_n_gate(name, plat, src_col):
    return (
        f'RA8 {name} ({plat}) =\n'
        f'{yw_var(V)}\n'
        f'VAR MinN = [Min N Cutoff Value]\n'
        f'RETURN IF(\n'
        f'    CALCULATE(SUM({V}[n]),\n'
        f'        {V}[yearweek] = SelYearWeek,\n'
        f'        {V}[platform] = "{plat}") >= MinN,\n'
        f'    CALCULATE(AVERAGE({V}[{src_col}]),\n'
        f'        {V}[yearweek] = SelYearWeek,\n'
        f'        {V}[platform] = "{plat}"),\n'
        f'    BLANK())'
    )

# Helper: pct measure with MinN gate on Who_Got_Message (for participation by-category)
def pct_with_whogot_gate(name, plat, src_col, gate_col='Who_Got_Message'):
    return (
        f'RA8 {name} ({plat}) =\n'
        f'{yw_var(V)}\n'
        f'VAR MinN = [Min N Cutoff Value]\n'
        f'RETURN IF(\n'
        f'    CALCULATE(SUM({V}[{gate_col}]),\n'
        f'        {V}[yearweek] = SelYearWeek,\n'
        f'        {V}[platform] = "{plat}") >= MinN,\n'
        f'    CALCULATE(AVERAGE({V}[{src_col}]),\n'
        f'        {V}[yearweek] = SelYearWeek,\n'
        f'        {V}[platform] = "{plat}"),\n'
        f'    BLANK())'
    )

# Excel M, N, O, V — gated on n
for plat in ('Snapp', 'Tapsi'):
    dax_measure(pct_with_n_gate('pct Got Message',     plat, 'pct_Got_Message'))
    dax_measure(pct_with_n_gate('pct FreeComm Message', plat, 'pct_FreeComm_Message'))
    dax_measure(pct_with_n_gate('pct Free Comm Ride',  plat, 'pct_Free_Comm_Ride'))
    dax_measure(pct_with_n_gate('pct Participated',    plat, 'pct_Participated'))

# Excel P/Q/R — gated on GotMsg_<cat> (denominator)
for plat in ('Snapp', 'Tapsi'):
    dax_measure(pct_with_whogot_gate('pct Part Money',           plat, 'pct_Part_Money',           'GotMsg_Money'))
    dax_measure(pct_with_whogot_gate('pct Part FreeComm',        plat, 'pct_Part_FreeComm',        'GotMsg_FreeComm'))
    dax_measure(pct_with_whogot_gate('pct Part Money FreeComm',  plat, 'pct_Part_Money_FreeComm',  'GotMsg_Money_FreeComm'))

# Excel S/T/U — gated on GotMsg_<type> (multi-select TYPE breakdown)
for plat in ('Snapp', 'Tapsi'):
    dax_measure(pct_with_whogot_gate('pct Part IncGuar', plat, 'pct_Part_IncGuar', 'GotMsg_IncGuar'))
    dax_measure(pct_with_whogot_gate('pct Part PayRide', plat, 'pct_Part_PayRide', 'GotMsg_PayRide'))
    dax_measure(pct_with_whogot_gate('pct Part PayInc',  plat, 'pct_Part_PayInc',  'GotMsg_PayInc'))

# Ride sums (no MinN gate — these are population sums, not survey ratios)
for plat in ('Snapp', 'Tapsi'):
    for short, col in [('Ride Total', 'Ride_Total'),
                       ('Ride AmongFC', 'Ride_AmongFC'),
                       ('Ride FreeComm', 'Ride_FreeComm')]:
        dax_measure(
            f'RA8 {short} ({plat}) =\n'
            f'{yw_var(V)}\n'
            f'RETURN CALCULATE(SUM({V}[{col}]),\n'
            f'    {V}[yearweek] = SelYearWeek,\n'
            f'    {V}[platform] = "{plat}")'
        )
    dax_measure(pct_with_n_gate('pct FCRide AllShare', plat, 'pct_FCRide_AllShare'))
    dax_measure(pct_with_n_gate('pct FCRide FCShare',  plat, 'pct_FCRide_FCShare'))

# WoW examples
dax_measure(
    f'RA8 WoW pct Got Message (Snapp) =\n'
    f'{yw_var(V)}\n'
    f'VAR PrevYearWeek = CALCULATE(MAX({V}[yearweek]),\n'
    f'    ALL({V}), {V}[yearweek] < SelYearWeek)\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'VAR CurrVal = CALCULATE(AVERAGE({V}[pct_Got_Message]),\n'
    f'    {V}[yearweek] = SelYearWeek, {V}[platform] = "Snapp")\n'
    f'VAR PrevVal = CALCULATE(AVERAGE({V}[pct_Got_Message]),\n'
    f'    {V}[yearweek] = PrevYearWeek, {V}[platform] = "Snapp")\n'
    f'RETURN IF(\n'
    f'    CALCULATE(SUM({V}[n]),\n'
    f'        {V}[yearweek] = SelYearWeek,\n'
    f'        {V}[platform] = "Snapp") >= MinN,\n'
    f'    CurrVal - PrevVal,\n'
    f'    BLANK())'
)
note('Apply the same WoW pattern to any other pct_* measure by swapping the source column. '
     'Apply the same pattern for Tapsi by swapping "Snapp" → "Tapsi".')

doc.add_paragraph()
body('Cost Measures (Excel AF / AG / W) — require hardcoded constants:', bold=True)
note('Excel AF/AG/W use four constants per (city, platform) that are NOT in the survey:\n'
     '  D = total active driver population (e.g. 107,326 for Tapsi Tehran)\n'
     '  X = rides per board (e.g. 11.85 for Tapsi Tehran, 21.66 for Snapp Tehran)\n'
     '  Y = commission % (typically 0.15)\n'
     '  Z = avg fare in Rial (e.g. 1,117,997 for Tapsi Tehran)\n'
     'Add these as a separate Power BI table (Constants) keyed on (city, platform), '
     'or as DAX SWITCH() against city. Below is the formula skeleton — replace VAR D / X / Y / Z '
     'with your constants table lookups.')

dax_measure(
    'RA8 final drivers (Tapsi) =\n'
    '// Excel W = O × D — population estimate × % who had a Free-Commission ride\n'
    f'{yw_var(V)}\n'
    'VAR D = [Total Driver Pop (Tapsi)]   // your constants lookup\n'
    'VAR PctO = CALCULATE(AVERAGE(vw_RA_CommFree[pct_Free_Comm_Ride]),\n'
    '    vw_RA_CommFree[yearweek] = SelYearWeek,\n'
    '    vw_RA_CommFree[platform] = "Tapsi")\n'
    'RETURN D * PctO / 100'
)
dax_measure(
    'RA8 Commission Free Cost #1 (Tapsi) =\n'
    '// Excel AF = D × M × N × AVG(Q,R) × X × AE × Y × Z\n'
    f'{yw_var(V)}\n'
    'VAR D = [Total Driver Pop (Tapsi)]\n'
    'VAR X = [Rides per Brd (Tapsi)]\n'
    'VAR Y = [Commission Pct (Tapsi)]   // e.g. 0.15\n'
    'VAR Z = [Avg Fare (Tapsi)]\n'
    'VAR M = CALCULATE(AVERAGE(vw_RA_CommFree[pct_Got_Message]),\n'
    '    vw_RA_CommFree[yearweek] = SelYearWeek, vw_RA_CommFree[platform] = "Tapsi") / 100\n'
    'VAR N = CALCULATE(AVERAGE(vw_RA_CommFree[pct_FreeComm_Message]),\n'
    '    vw_RA_CommFree[yearweek] = SelYearWeek, vw_RA_CommFree[platform] = "Tapsi") / 100\n'
    'VAR Q = CALCULATE(AVERAGE(vw_RA_CommFree[pct_Part_FreeComm]),\n'
    '    vw_RA_CommFree[yearweek] = SelYearWeek, vw_RA_CommFree[platform] = "Tapsi") / 100\n'
    'VAR R = CALCULATE(AVERAGE(vw_RA_CommFree[pct_Part_Money_FreeComm]),\n'
    '    vw_RA_CommFree[yearweek] = SelYearWeek, vw_RA_CommFree[platform] = "Tapsi") / 100\n'
    'VAR AE = CALCULATE(AVERAGE(vw_RA_CommFree[pct_FCRide_FCShare]),\n'
    '    vw_RA_CommFree[yearweek] = SelYearWeek, vw_RA_CommFree[platform] = "Tapsi") / 100\n'
    'VAR PartAvgQR = DIVIDE(Q + R, 2)\n'
    'RETURN D * M * N * PartAvgQR * X * AE * Y * Z'
)
dax_measure(
    'RA8 Commission Free Cost #2 (Tapsi) =\n'
    '// Excel AG = D × O × X × AE × Y × Z (simpler ride-based estimate)\n'
    f'{yw_var(V)}\n'
    'VAR D = [Total Driver Pop (Tapsi)]\n'
    'VAR X = [Rides per Brd (Tapsi)]\n'
    'VAR Y = [Commission Pct (Tapsi)]\n'
    'VAR Z = [Avg Fare (Tapsi)]\n'
    'VAR O = CALCULATE(AVERAGE(vw_RA_CommFree[pct_Free_Comm_Ride]),\n'
    '    vw_RA_CommFree[yearweek] = SelYearWeek, vw_RA_CommFree[platform] = "Tapsi") / 100\n'
    'VAR AE = CALCULATE(AVERAGE(vw_RA_CommFree[pct_FCRide_FCShare]),\n'
    '    vw_RA_CommFree[yearweek] = SelYearWeek, vw_RA_CommFree[platform] = "Tapsi") / 100\n'
    'RETURN D * O * X * AE * Y * Z'
)
note('Duplicate the three cost measures for Snapp by replacing "Tapsi" → "Snapp" and pointing '
     'to the corresponding Snapp constants. Excel calibrates AH = AF × (AJ45/AF45) where AJ45 '
     'is a ground-truth total cost from accounting (e.g. 227,140,884,338 Rial in week 52). '
     'If you have that ground truth, build it as a separate parameter and multiply.')
note('Visual: Two side-by-side matrices (Snapp & Tapsi) with city rows + cards for national totals. '
     'For the cost columns, build a separate "Constants" table with one row per (city, platform).')

# ════════════════════════════════════════════════════════════════════════════
# RA-9  vw_RA_CSRare
# ════════════════════════════════════════════════════════════════════════════
heading('11. RA-9 – Customer Support Satisfaction', level=1)
body('View: vw_RA_CSRare  |  Excel Pages: CS_Sat_Snapp / CS_Sat_Tapsi', bold=True)
body('ShortRare joined to ShortMain. All scores are 1–5 averages.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('n', 'INT', 'respondents with ShortRare record'),
        ('Snapp_CS_Overall', 'FLOAT', 'avg overall CS satisfaction 1-5'),
        ('Snapp_CS_WaitTime', 'FLOAT', ''),
        ('Snapp_CS_Solution', 'FLOAT', ''),
        ('Snapp_CS_Behaviour', 'FLOAT', ''),
        ('Snapp_CS_Relevance', 'FLOAT', ''),
        ('Snapp_CS_Solved_pct', 'FLOAT', '% whose issue was solved'),
        ('Tapsi_CS_Overall', 'FLOAT', ''),
        ('Tapsi_CS_WaitTime', 'FLOAT', ''),
        ('Tapsi_CS_Solution', 'FLOAT', ''),
        ('Tapsi_CS_Behaviour', 'FLOAT', ''),
        ('Tapsi_CS_Relevance', 'FLOAT', ''),
        ('Tapsi_CS_Solved_pct', 'FLOAT', ''),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_CSRare'
dax_measure(count_dax('RA9 n', V))
for col in ['Snapp_CS_Overall', 'Snapp_CS_WaitTime', 'Snapp_CS_Solution',
            'Snapp_CS_Behaviour', 'Snapp_CS_Relevance']:
    short = col.replace('Snapp_CS_', '').replace('_', ' ')
    dax_measure(metric_dax(f'RA9 Snapp CS {short}', V, f'AVERAGE({V}[{col}])'))
dax_measure(metric_dax('RA9 Snapp CS Solved pct', V, f'AVERAGE({V}[Snapp_CS_Solved_pct])'))
dax_measure(metric_dax('RA9 Tapsi CS Overall', V, f'AVERAGE({V}[Tapsi_CS_Overall])'))
dax_measure(wow_dax('RA9 WoW Snapp CS Overall', V, f'AVERAGE({V}[Snapp_CS_Overall])'))
note('Duplicate all five Snapp score measures + Solved_pct for Tapsi columns.')
note('Visual: Matrix city × score column. Line chart for WoW trend.')

# ════════════════════════════════════════════════════════════════════════════
# RA-10  vw_RA_NavReco
# ════════════════════════════════════════════════════════════════════════════
heading('12. RA-10 – Navigation & NPS Recommendation Scores', level=1)
body('View: vw_RA_NavReco  |  Excel Pages: NavReco_Scores / Reco_NPS', bold=True)
body('From ShortRare. NPS scores and navigation-app recommendation scores (0–10).')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('n', 'INT', ''),
        ('Snapp_NPS', 'FLOAT', 'avg recommend-Snapp score'),
        ('Tapsi_NPS_SnapDriver', 'FLOAT', 'avg recommend-Tapsi (Snapp driver respondent)'),
        ('Tapsi_NPS_TapsiDriver', 'FLOAT', 'avg recommend-Tapsi (Tapsi driver respondent)'),
        ('Reco_GoogleMap', 'FLOAT', 'avg recommendation score'),
        ('Reco_Waze', 'FLOAT', ''),
        ('Reco_Neshan', 'FLOAT', ''),
        ('Reco_Balad', 'FLOAT', ''),
        ('Snapp_Nav_Sat', 'FLOAT', 'avg Snapp nav app satisfaction'),
        ('Tapsi_Nav_Sat', 'FLOAT', 'avg Tapsi in-app nav satisfaction'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_NavReco'
dax_measure(count_dax('RA10 n', V))
for col, label in [
    ('Snapp_NPS', 'Snapp NPS'),
    ('Tapsi_NPS_SnapDriver', 'Tapsi NPS (Snapp Driver)'),
    ('Tapsi_NPS_TapsiDriver', 'Tapsi NPS (Tapsi Driver)'),
    ('Reco_Neshan', 'Reco Neshan'),
    ('Reco_Balad', 'Reco Balad'),
    ('Reco_GoogleMap', 'Reco GoogleMap'),
    ('Reco_Waze', 'Reco Waze'),
    ('Snapp_Nav_Sat', 'Snapp Nav Sat'),
    ('Tapsi_Nav_Sat', 'Tapsi Nav Sat'),
]:
    dax_measure(metric_dax(f'RA10 {label}', V, f'AVERAGE({V}[{col}])'))
dax_measure(wow_dax('RA10 WoW Snapp NPS', V, f'AVERAGE({V}[Snapp_NPS])'))

# ════════════════════════════════════════════════════════════════════════════
# RA-11  vw_RA_IncentiveTypes
# ════════════════════════════════════════════════════════════════════════════
heading('13. RA-11 – Incentive Type Distribution', level=1)
body('View: vw_RA_IncentiveTypes  |  Excel Pages: #5 (Snapp Excl), #6 (Joint)', bold=True)
body('Multi-select incentive type flags (6 types × 3 segments). '
     'Base for type %: n_excl_gotmsg (Excl), n_joint_gotmsg_sn (JntSn), n_joint_gotmsg_tp (JntTp) — '
     'only drivers who received a gotmessage incentive, matching Excel denominator.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('n / n_joint / n_excl', 'INT', 'all respondents in segment'),
        ('n_excl_gotmsg', 'INT', 'excl drivers with Snapp gotmessage=Yes — type % denominator'),
        ('n_joint_gotmsg_sn', 'INT', 'joint drivers with Snapp gotmessage=Yes'),
        ('n_joint_gotmsg_tp', 'INT', 'joint drivers with Tapsi gotmessage=Yes'),
        ('pct_GotMsg_Excl_Snapp', 'FLOAT', '% excl drivers who got Snapp msg (base = n_excl)'),
        ('pct_GotMsg_Jnt_Snapp', 'FLOAT', '% joint who got Snapp msg'),
        ('pct_GotMsg_Jnt_Tapsi', 'FLOAT', '% joint who got Tapsi msg'),
        ('pct_GotMsg_Both', 'FLOAT', '% joint who got both msgs'),
        ('pct_GotMsg_Diff', 'FLOAT', '% joint who got only one msg'),
        ('pct_PayRide_Excl / JntSn / JntTp', 'FLOAT', 'Pay After Ride — % of gotmsg drivers'),
        ('pct_EarnCF_Excl / JntSn / JntTp', 'FLOAT', 'Earning-Based CF type %'),
        ('pct_RideCF_Excl / JntSn / JntTp', 'FLOAT', 'Ride-Based CF type %'),
        ('pct_IncGuar_Excl / JntSn / JntTp', 'FLOAT', 'Income Guarantee type %'),
        ('pct_PayInc_Excl / JntSn / JntTp', 'FLOAT', 'Pay After Income type %'),
        ('pct_CFSome_Excl / JntSn / JntTp', 'FLOAT', 'CF on Some Trips type %'),
        ('Avg_CF_Rides_Snapp', 'FLOAT', 'avg CF rides (CF drivers only)'),
        ('Avg_CF_Rides_Tapsi', 'FLOAT', ''),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_IncentiveTypes'
note('n_excl_gotmsg / n_joint_gotmsg_sn / n_joint_gotmsg_tp are the gotmessage=Yes counts used as '
     'MinN gates for type % measures — matching the Excel denominator (only drivers who received the incentive message).')
dax_measure(count_dax('RA11 n Excl', V, count_col='n_excl'))
dax_measure(count_dax('RA11 n Joint', V, count_col='n_joint'))
dax_measure(count_dax('RA11 n Excl GotMsg', V, count_col='n_excl_gotmsg'))
dax_measure(count_dax('RA11 n Joint GotMsg Sn', V, count_col='n_joint_gotmsg_sn'))
dax_measure(count_dax('RA11 n Joint GotMsg Tp', V, count_col='n_joint_gotmsg_tp'))
dax_measure(metric_dax('RA11 pct GotMsg Excl Snapp', V, f'AVERAGE({V}[pct_GotMsg_Excl_Snapp])', n_col='n_excl'))
dax_measure(metric_dax('RA11 pct GotMsg Jnt Snapp', V, f'AVERAGE({V}[pct_GotMsg_Jnt_Snapp])', n_col='n_joint'))
dax_measure(metric_dax('RA11 pct GotMsg Jnt Tapsi', V, f'AVERAGE({V}[pct_GotMsg_Jnt_Tapsi])', n_col='n_joint'))
dax_measure(metric_dax('RA11 pct GotMsg Both', V, f'AVERAGE({V}[pct_GotMsg_Both])', n_col='n_joint'))
dax_measure(metric_dax('RA11 pct GotMsg Diff', V, f'AVERAGE({V}[pct_GotMsg_Diff])', n_col='n_joint'))

for seg, base in [('Excl', 'n_excl_gotmsg'), ('JntSn', 'n_joint_gotmsg_sn'), ('JntTp', 'n_joint_gotmsg_tp')]:
    for t in ['PayRide', 'EarnCF', 'RideCF', 'IncGuar', 'PayInc', 'CFSome']:
        col = f'pct_{t}_{seg}'
        dax_measure(metric_dax(f'RA11 {col}', V, f'AVERAGE({V}[{col}])', n_col=base))

dax_measure(metric_dax('RA11 Avg CF Rides Snapp', V, f'AVERAGE({V}[Avg_CF_Rides_Snapp])'))
dax_measure(metric_dax('RA11 Avg CF Rides Tapsi', V, f'AVERAGE({V}[Avg_CF_Rides_Tapsi])', n_col='n_joint'))

# ════════════════════════════════════════════════════════════════════════════
# RA-12  vw_RA_IncentiveUnsatCity
# ════════════════════════════════════════════════════════════════════════════
heading('14. RA-12 – Incentive Dissatisfaction by City', level=1)
body('View: vw_RA_IncentiveUnsatCity  |  Excel Page: #8', bold=True)
body('"Low rate" = driver gave incentive satisfaction rating < 4 (i.e. 1, 2, or 3). MATCHES EXCEL #8.')
bullet('Snapp denominator = n_sn_lowrate = drivers with snapp_overall_incentive_satisfaction < 4')
bullet('Tapsi denominator = n_tp_lowrate = JOINT drivers with tapsi_overall_incentive_satisfaction < 4')
bullet('Numerator = drivers with rating < 4 AND cited specific reason (multi-select WideMain flag = 1)')
bullet('IMPORTANT: this is NOT the same as "drivers who cited any reason". Some drivers rated < 4 '
       'but didn\'t pick any reason; some picked a reason but rated >= 4. Excel uses the rating filter, '
       'so we match it.')
note('Back-compat: aliases n_sn_low_sat / n_tp_low_sat are kept (= n_sn_lowrate / n_tp_lowrate) so '
     'existing measures referencing the old names continue to work.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('n_all', 'INT', 'all respondents in city'),
        ('n_joint', 'INT', 'joint drivers'),
        ('n_sn_lowrate', 'INT', 'Excel "sn<4" — drivers with Snapp rating < 4'),
        ('n_tp_lowrate', 'INT', 'Excel "t30<4" — joint drivers with Tapsi rating < 4'),
        ('n_sn_low_sat / n_tp_low_sat', 'INT', 'aliases of the above (back-compat)'),
        ('pct_Sn_NoTime', 'FLOAT', '% of n_sn_lowrate citing "Not Available"'),
        ('pct_Sn_ImpAmt', 'FLOAT', '% citing "Improper Amount"'),
        ('pct_Sn_LowTime', 'FLOAT', '% citing "No Time todo"'),
        ('pct_Sn_HardToDo', 'FLOAT', '% citing "difficult"'),
        ('pct_Sn_NonPay', 'FLOAT', '% citing "Non Payment"'),
        ('pct_Tp_NoTime', 'FLOAT', '% of n_tp_lowrate citing each Tapsi reason'),
        ('pct_Tp_ImpAmt / pct_Tp_LowTime / pct_Tp_HardToDo / pct_Tp_NonPay', 'FLOAT', ''),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_IncentiveUnsatCity'
dax_measure(count_dax('RA12 n All', V, count_col='n_all'))
dax_measure(count_dax('RA12 n Sn Low Sat', V, count_col='n_sn_low_sat'))
dax_measure(count_dax('RA12 n Tp Low Sat', V, count_col='n_tp_low_sat'))
for col, label in [
    ('pct_Sn_NoTime', 'pct Sn NoTime'),
    ('pct_Sn_ImpAmt', 'pct Sn ImpAmt'),
    ('pct_Sn_LowTime', 'pct Sn LowTime'),
    ('pct_Sn_HardToDo', 'pct Sn HardToDo'),
    ('pct_Sn_NonPay', 'pct Sn NonPay'),
]:
    dax_measure(metric_dax(f'RA12 {label}', V, f'AVERAGE({V}[{col}])', n_col='n_sn_low_sat'))
for col, label in [
    ('pct_Tp_NoTime', 'pct Tp NoTime'),
    ('pct_Tp_ImpAmt', 'pct Tp ImpAmt'),
    ('pct_Tp_LowTime', 'pct Tp LowTime'),
    ('pct_Tp_HardToDo', 'pct Tp HardToDo'),
    ('pct_Tp_NonPay', 'pct Tp NonPay'),
]:
    dax_measure(metric_dax(f'RA12 {label}', V, f'AVERAGE({V}[{col}])', n_col='n_tp_low_sat'))

# ════════════════════════════════════════════════════════════════════════════
# RA-13  vw_RA_IncentiveUnsatNational
# ════════════════════════════════════════════════════════════════════════════
heading('15. RA-13 – Incentive Dissatisfaction (National)', level=1)
body('View: vw_RA_IncentiveUnsatNational  |  Excel Page: #9', bold=True)
body('Long format – no city column. Segments: All Snapp / Joint Snapp / Joint Tapsi. '
     'MATCHES EXCEL #9 — denominator is rating-based (rating < 4), NOT "cited any reason":')
bullet('All Snapp:   n_low_sat = COUNT(snapp_overall_incentive_satisfaction < 4)')
bullet('Joint Snapp: n_low_sat = COUNT(active_joint=1 AND snapp rating < 4)')
bullet('Joint Tapsi: n_low_sat = COUNT(active_joint=1 AND tapsi rating < 4)')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber', '', ''),
        ('segment', 'TEXT', '"All Snapp", "Joint Snapp", "Joint Tapsi"'),
        ('segment_sort', 'INT', '1, 2, 3'),
        ('n', 'INT', 'total respondents in segment'),
        ('n_low_sat', 'INT', 'drivers with rating < 4 in segment (Excel denominator)'),
        ('pct_NoTime', 'FLOAT', '% of n_low_sat citing "Not Available"'),
        ('pct_ImpAmt', 'FLOAT', '% citing "Improper Amount"'),
        ('pct_LowTime', 'FLOAT', '% citing "No Time todo"'),
        ('pct_HardToDo', 'FLOAT', '% citing "difficult"'),
        ('pct_NonPay', 'FLOAT', '% citing "Non Payment"'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)
note('Create all five pct measures × three segments = 15 measures. Template below; replace segment name.')

V = 'vw_RA_IncentiveUnsatNational'
for seg in ['All Snapp', 'Joint Snapp', 'Joint Tapsi']:
    seg_var = seg.replace(' ', '')
    for col, label in [
        ('pct_NoTime', 'NoTime'),
        ('pct_ImpAmt', 'ImpAmt'),
        ('pct_LowTime', 'LowTime'),
        ('pct_HardToDo', 'HardToDo'),
        ('pct_NonPay', 'NonPay'),
    ]:
        dax_measure(
            f'RA13 {label} ({seg}) =\n'
            f'{yw_var(V)}\n'
            f'VAR MinN = [Min N Cutoff Value]\n'
            f'RETURN IF(\n'
            f'    CALCULATE(SUM({V}[n_low_sat]),\n'
            f'        {V}[yearweek] = SelYearWeek,\n'
            f'        {V}[segment] = "{seg}") >= MinN,\n'
            f'    CALCULATE(AVERAGE({V}[{col}]),\n'
            f'        {V}[yearweek] = SelYearWeek,\n'
            f'        {V}[segment] = "{seg}"),\n'
            f'    BLANK())'
        )

dax_measure(
    f'RA13 WoW pct NoTime (All Snapp) =\n'
    f'{yw_var(V)}\n'
    f'VAR PrevYearWeek = CALCULATE(MAX({V}[yearweek]),\n'
    f'    ALL({V}), {V}[yearweek] < SelYearWeek)\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'VAR CurrVal = CALCULATE(AVERAGE({V}[pct_NoTime]),\n'
    f'    {V}[yearweek] = SelYearWeek, {V}[segment] = "All Snapp")\n'
    f'VAR PrevVal = CALCULATE(AVERAGE({V}[pct_NoTime]),\n'
    f'    {V}[yearweek] = PrevYearWeek, {V}[segment] = "All Snapp")\n'
    f'RETURN IF(\n'
    f'    CALCULATE(SUM({V}[n_low_sat]),\n'
    f'        {V}[yearweek] = SelYearWeek,\n'
    f'        {V}[segment] = "All Snapp") >= MinN,\n'
    f'    CurrVal - PrevVal,\n'
    f'    BLANK())'
)
note('Visual: Clustered bar – reason on x-axis, segment as legend. Line for WoW trend per reason.')

# ════════════════════════════════════════════════════════════════════════════
# RA-14  vw_RA_Navigation
# ════════════════════════════════════════════════════════════════════════════
heading('16. RA-14 – Navigation App Usage by City', level=1)
body('View: vw_RA_Navigation  |  Excel Page: #14', bold=True)
body('UNION ALL Snapp (all) + Tapsi (joint only). '
     'Snapp has GoogleMap & Waze; Tapsi has InAppNav (NULL for the other platform).')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('platform', 'TEXT', '"Snapp" or "Tapsi"'),
        ('platform_sort', 'INT', '1=Snapp, 2=Tapsi — use to sort platform in Matrix Columns'),
        ('n', 'INT', 'respondents with non-null navigation answer'),
        ('pct_Neshan', 'FLOAT', ''),
        ('pct_Balad', 'FLOAT', ''),
        ('pct_None', 'FLOAT', '% No Navigation App'),
        ('pct_GoogleMap', 'FLOAT', 'Snapp only (NULL for Tapsi)'),
        ('pct_Waze', 'FLOAT', 'Snapp only (NULL for Tapsi)'),
        ('pct_InAppNav', 'FLOAT', 'Tapsi only (NULL for Snapp)'),
        ('pct_Other', 'FLOAT', ''),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_Navigation'
for plat, cols in [
    ('Snapp', ['pct_Neshan', 'pct_Balad', 'pct_None', 'pct_GoogleMap', 'pct_Waze', 'pct_Other']),
    ('Tapsi', ['pct_Neshan', 'pct_Balad', 'pct_None', 'pct_InAppNav', 'pct_Other']),
]:
    dax_measure(
        f'RA14 n ({plat}) =\n'
        f'{yw_var(V)}\n'
        f'RETURN CALCULATE(SUM({V}[n]),\n'
        f'    {V}[yearweek] = SelYearWeek, {V}[platform] = "{plat}")'
    )
    for col in cols:
        label = col.replace('pct_', '')
        dax_measure(
            f'RA14 {label} ({plat}) =\n'
            f'{yw_var(V)}\n'
            f'VAR MinN = [Min N Cutoff Value]\n'
            f'RETURN IF(\n'
            f'    CALCULATE(SUM({V}[n]),\n'
            f'        {V}[yearweek] = SelYearWeek,\n'
            f'        {V}[platform] = "{plat}") >= MinN,\n'
            f'    CALCULATE(AVERAGE({V}[{col}]),\n'
            f'        {V}[yearweek] = SelYearWeek,\n'
            f'        {V}[platform] = "{plat}"),\n'
            f'    BLANK())'
        )

dax_measure(
    f'RA14 WoW pct Neshan (Snapp) =\n'
    f'{yw_var(V)}\n'
    f'VAR PrevYearWeek = CALCULATE(MAX({V}[yearweek]),\n'
    f'    ALL({V}), {V}[yearweek] < SelYearWeek)\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'VAR CurrVal = CALCULATE(AVERAGE({V}[pct_Neshan]),\n'
    f'    {V}[yearweek] = SelYearWeek, {V}[platform] = "Snapp")\n'
    f'VAR PrevVal = CALCULATE(AVERAGE({V}[pct_Neshan]),\n'
    f'    {V}[yearweek] = PrevYearWeek, {V}[platform] = "Snapp")\n'
    f'RETURN IF(\n'
    f'    CALCULATE(SUM({V}[n]),\n'
    f'        {V}[yearweek] = SelYearWeek,\n'
    f'        {V}[platform] = "Snapp") >= MinN,\n'
    f'    CurrVal - PrevVal,\n'
    f'    BLANK())'
)
note('Visual: Stacked bar – navigation app as legend, city on axis. Separate Snapp / Tapsi pages.')

# ════════════════════════════════════════════════════════════════════════════
# RA-15  vw_RA_Referral
# ════════════════════════════════════════════════════════════════════════════
heading('17. RA-15 – Referral / Joining Bonus', level=1)
body('View: vw_RA_Referral  |  Excel Page: #16', bold=True)
body('Snapp joining bonus: all drivers. Tapsi joining bonus: joint drivers only.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('n_Snapp', 'INT', 'respondents with non-null Snapp joining_bonus'),
        ('joining_Snapp', 'INT', 'count who got Snapp bonus'),
        ('pct_Joining_Snapp', 'FLOAT', ''),
        ('n_Tapsi', 'INT', 'joint respondents with non-null Tapsi joining_bonus'),
        ('joining_Tapsi', 'INT', 'count who got Tapsi bonus'),
        ('pct_Joining_Tapsi', 'FLOAT', ''),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_Referral'
dax_measure(count_dax('RA15 n Snapp', V, count_col='n_Snapp'))
dax_measure(count_dax('RA15 n Tapsi', V, count_col='n_Tapsi'))
dax_measure(count_dax('RA15 joining Snapp', V, count_col='joining_Snapp'))
dax_measure(count_dax('RA15 joining Tapsi', V, count_col='joining_Tapsi'))
dax_measure(metric_dax('RA15 pct Joining Snapp', V, f'AVERAGE({V}[pct_Joining_Snapp])', n_col='n_Snapp'))
dax_measure(metric_dax('RA15 pct Joining Tapsi', V, f'AVERAGE({V}[pct_Joining_Tapsi])', n_col='n_Tapsi'))
dax_measure(wow_dax('RA15 WoW pct Joining Snapp', V, f'AVERAGE({V}[pct_Joining_Snapp])', n_col='n_Snapp'))
dax_measure(wow_dax('RA15 WoW pct Joining Tapsi', V, f'AVERAGE({V}[pct_Joining_Tapsi])', n_col='n_Tapsi'))
note('Visual: Matrix with city rows. Cards for national joining pct. Line chart for WoW.')

# ════════════════════════════════════════════════════════════════════════════
# RA-16  vw_RA_TapsiInactivity
# ════════════════════════════════════════════════════════════════════════════
heading('18. RA-16 – Tapsi Inactivity Before Incentive', level=1)
body('View: vw_RA_TapsiInactivity  |  Excel Page: #17', bold=True)
body('Long format with inactivity time bucket. Matches Excel #17 exactly:')
bullet('Filter (numerator): tapsi_gotmessage_text_incentive=\'Yes\' AND bucket = X')
bullet('n_total / "res" (denominator): joint_by_signup=1 AND tapsi_gotmessage_text_incentive=\'Yes\'')
bullet('NOTE: uses joint_by_signup, NOT active_joint — the latter is stricter and returns ~30% fewer drivers.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('inactivity_bucket', 'TEXT',
         '"Same Day", "1_3 Day Before", "3_7 Days Before", '
         '"8_14 Days Before", "15_30 Days_Before", "1_2 Month Before", '
         '"2_3 Month Before", "3_6Month Before", ">6 Month Before"'),
        ('bucket_sort', 'INT', '1–9 + 99'),
        ('n', 'INT', 'count in this bucket'),
        ('n_total', 'INT', 'total joint drivers in city × week'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)
note('Sort inactivity_bucket by bucket_sort in the Model view (so columns appear in chronological order).')
note('PREFERRED Matrix layout (matches Excel #17): Rows = city, Columns = inactivity_bucket, '
     'Values = single [pct Bucket] measure → produces all 9 bucket columns auto-summing to 100% per row. '
     'The 9 explicit per-bucket measures below are provided for cases where you want fixed columns '
     '(e.g. cards or a flat table) instead of using inactivity_bucket on the columns axis.')

V = 'vw_RA_TapsiInactivity'

BUCKETS = [
    ('Same Day',          '<1 day'),
    ('1_3 Day Before',    '1-3 days'),
    ('3_7 Days Before',   '4-7 days'),
    ('8_14 Days Before',  '8-14 days'),
    ('15_30 Days_Before', '15-30 days'),
    ('1_2 Month Before',  '1-2 months'),
    ('2_3 Month Before',  '2-3 months'),
    ('3_6Month Before',   '3-6 months'),
    ('>6 Month Before',   '>6 months'),
]
LT_1MONTH = [b[0] for b in BUCKETS[:5]]   # Same Day .. 15_30 Days_Before
GT_1MONTH = [b[0] for b in BUCKETS[5:]]   # 1_2 Month .. >6 Month

dax_measure(
    f'RA16 n Total =\n'
    f'{yw_var(V)}\n'
    f'RETURN CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek)'
)
dax_measure(
    f'RA16 pct Bucket =\n'
    f'// Use this single measure with inactivity_bucket on the Matrix Columns well\n'
    f'// (rows = city, sorted by bucket_sort). Each cell = % of city drivers in that bucket.\n'
    f'{yw_var(V)}\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'RETURN IF(\n'
    f'    CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek) >= MinN,\n'
    f'    DIVIDE(\n'
    f'        CALCULATE(SUM({V}[n]), {V}[yearweek] = SelYearWeek),\n'
    f'        CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek)) * 100,\n'
    f'    BLANK())'
)

# Explicit per-bucket measures so the user can build a flat table matching Excel #17
for bucket, label in BUCKETS:
    dax_measure(
        f'RA16 pct {label} =\n'
        f'{yw_var(V)}\n'
        f'VAR MinN = [Min N Cutoff Value]\n'
        f'RETURN IF(\n'
        f'    CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek) >= MinN,\n'
        f'    DIVIDE(\n'
        f'        CALCULATE(SUM({V}[n]),\n'
        f'            {V}[yearweek] = SelYearWeek,\n'
        f'            {V}[inactivity_bucket] = "{bucket}"),\n'
        f'        CALCULATE(MAX({V}[n_total]),\n'
        f'            {V}[yearweek] = SelYearWeek)) * 100,\n'
        f'    BLANK())'
    )

# <1 Month / >1 Months summary (matches the small bottom table in Excel #17)
def _bucket_list_dax(buckets):
    return ', '.join(f'"{b}"' for b in buckets)

dax_measure(
    f'RA16 pct <1 Month =\n'
    f'// Sum of "<1 day" through "15-30 days" — matches Excel "<1 Month" column\n'
    f'{yw_var(V)}\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'RETURN IF(\n'
    f'    CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek) >= MinN,\n'
    f'    DIVIDE(\n'
    f'        CALCULATE(SUM({V}[n]),\n'
    f'            {V}[yearweek] = SelYearWeek,\n'
    f'            {V}[inactivity_bucket] IN {{{_bucket_list_dax(LT_1MONTH)}}}),\n'
    f'        CALCULATE(MAX({V}[n_total]),\n'
    f'            {V}[yearweek] = SelYearWeek)) * 100,\n'
    f'    BLANK())'
)
dax_measure(
    f'RA16 pct >1 Months =\n'
    f'// Sum of "1-2 months" through ">6 months" — matches Excel ">1 Months" column\n'
    f'{yw_var(V)}\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'RETURN IF(\n'
    f'    CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek) >= MinN,\n'
    f'    DIVIDE(\n'
    f'        CALCULATE(SUM({V}[n]),\n'
    f'            {V}[yearweek] = SelYearWeek,\n'
    f'            {V}[inactivity_bucket] IN {{{_bucket_list_dax(GT_1MONTH)}}}),\n'
    f'        CALCULATE(MAX({V}[n_total]),\n'
    f'            {V}[yearweek] = SelYearWeek)) * 100,\n'
    f'    BLANK())'
)

dax_measure(
    f'RA16 WoW pct <1 day =\n'
    f'{yw_var(V)}\n'
    f'VAR PrevYearWeek = CALCULATE(MAX({V}[yearweek]),\n'
    f'    ALL({V}), {V}[yearweek] < SelYearWeek)\n'
    f'VAR MinN = [Min N Cutoff Value]\n'
    f'VAR CurrVal = DIVIDE(\n'
    f'    CALCULATE(SUM({V}[n]),\n'
    f'        {V}[yearweek] = SelYearWeek, {V}[inactivity_bucket] = "Same Day"),\n'
    f'    CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek)) * 100\n'
    f'VAR PrevVal = DIVIDE(\n'
    f'    CALCULATE(SUM({V}[n]),\n'
    f'        {V}[yearweek] = PrevYearWeek, {V}[inactivity_bucket] = "Same Day"),\n'
    f'    CALCULATE(MAX({V}[n_total]), {V}[yearweek] = PrevYearWeek)) * 100\n'
    f'RETURN IF(\n'
    f'    CALCULATE(MAX({V}[n_total]), {V}[yearweek] = SelYearWeek) >= MinN,\n'
    f'    CurrVal - PrevVal,\n'
    f'    BLANK())'
)
note('Apply the same WoW pattern to any other bucket by swapping "Same Day" for the target bucket value.')
note('To match Excel #17 exactly: Matrix with city rows + inactivity_bucket columns + [pct Bucket] in Values. '
     'Add a small companion table/cards for [pct <1 Month] and [pct >1 Months] below.')

# ════════════════════════════════════════════════════════════════════════════
# RA-17  vw_RA_LuckyWheel
# ════════════════════════════════════════════════════════════════════════════
heading('19. RA-17 – Lucky Wheel Usage', level=1)
body('View: vw_RA_LuckyWheel  |  Excel Page: #19', bold=True)
body('Matches Excel #19 exactly:')
bullet('n = COUNT(active_joint=1) per city — denominator for pct_usage')
bullet('n_users = COUNT(active_joint=1 AND tapsi_magical_window=\'Yes\') — Excel "Res Usage"')
bullet('pct_usage = n_users / n × 100 — Excel "Usage" column')
bullet('avg_wheel_amount = AVG(wheel) over active_joint drivers (NULLs skipped, 0s included) — Excel "Ave. Lucky wheel"')
bullet('IMPORTANT: usage = answered "Yes" to the lucky-wheel question, NOT wheel>0 (= won money). '
       'A driver can use the wheel and win 0.')
body('Columns:', bold=True)
col_table(
    ['Column', 'Type', 'Notes'],
    [
        ('yearweek / yearweek_sort / weeknumber / city', '', ''),
        ('n', 'INT', 'active_joint drivers in city — Usage % denominator'),
        ('n_users', 'INT', 'active_joint AND tapsi_magical_window=Yes — Excel Res Usage'),
        ('pct_usage', 'FLOAT', '% who used the wheel (Excel Usage)'),
        ('avg_wheel_amount', 'FLOAT', 'avg wheel amount over all active_joint (Excel Ave. Lucky wheel)'),
    ]
)
doc.add_paragraph()
body('DAX Measures:', bold=True)

V = 'vw_RA_LuckyWheel'
dax_measure(count_dax('RA17 n', V))
dax_measure(count_dax('RA17 n Users', V, count_col='n_users'))
dax_measure(metric_dax('RA17 pct Usage', V, f'AVERAGE({V}[pct_usage])'))
dax_measure(metric_dax('RA17 Avg Wheel Amount', V, f'AVERAGE({V}[avg_wheel_amount])', n_col='n_users'))
dax_measure(wow_dax('RA17 WoW pct Usage', V, f'AVERAGE({V}[pct_usage])'))
dax_measure(wow_dax('RA17 WoW Avg Wheel Amount', V, f'AVERAGE({V}[avg_wheel_amount])', n_col='n_users'))
note('Visual: Matrix with city rows. Cards for national pct_usage and avg_wheel_amount. Line for WoW.')

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX
# ════════════════════════════════════════════════════════════════════════════
heading('20. Appendix – General DAX Patterns', level=1)

body('Standard Metric Pattern', bold=True)
dax_measure(
    'Metric Name =\n'
    'VAR SelYearWeek = IF(HASONEVALUE(ViewName[yearweek]),\n'
    '    SELECTEDVALUE(ViewName[yearweek]),\n'
    '    CALCULATE(MAX(ViewName[yearweek]), ALL(ViewName)))\n'
    'VAR MinN = [Min N Cutoff Value]\n'
    'RETURN IF(\n'
    '    CALCULATE(SUM(ViewName[n]), ViewName[yearweek] = SelYearWeek) >= MinN,\n'
    '    CALCULATE(AVERAGE(ViewName[metric_col]), ViewName[yearweek] = SelYearWeek),\n'
    '    BLANK())'
)

body('WoW Delta Pattern', bold=True)
dax_measure(
    'WoW Metric Name =\n'
    'VAR SelYearWeek = IF(HASONEVALUE(ViewName[yearweek]),\n'
    '    SELECTEDVALUE(ViewName[yearweek]),\n'
    '    CALCULATE(MAX(ViewName[yearweek]), ALL(ViewName)))\n'
    'VAR PrevYearWeek = CALCULATE(MAX(ViewName[yearweek]),\n'
    '    ALL(ViewName), ViewName[yearweek] < SelYearWeek)\n'
    'VAR MinN    = [Min N Cutoff Value]\n'
    'VAR CurrVal = CALCULATE(AVERAGE(ViewName[metric_col]), ViewName[yearweek] = SelYearWeek)\n'
    'VAR PrevVal = CALCULATE(AVERAGE(ViewName[metric_col]), ViewName[yearweek] = PrevYearWeek)\n'
    'RETURN IF(\n'
    '    CALCULATE(SUM(ViewName[n]), ViewName[yearweek] = SelYearWeek) >= MinN,\n'
    '    CurrVal - PrevVal,\n'
    '    BLANK())'
)

body('Min N Cutoff Value (helper measure)', bold=True)
dax_measure(
    'Min N Cutoff Value =\n'
    "SELECTEDVALUE('Min N Cutoff'[Min N Cutoff], 0)"
)

body('Platform-filtered UNION ALL views (CommFree, Navigation)', bold=True)
note('Hardcode platform = "Snapp" or "Tapsi" inside each measure\'s CALCULATE filter. '
     'Do not rely on a slicer for platform — it bypasses the N-gate logic.')

body('Long-format Matrix views (IncentiveAmounts, Duration, Persona, TapsiInactivity)', bold=True)
note('Place the bucket/category column on the matrix column axis. '
     'Sort it by the _sort column via Sort by Column in the Model view. '
     'Put the pct measure as Values. city on Rows. yearweek slicer filters the whole matrix.')

body('Segment-filtered views (IncentiveUnsatNational)', bold=True)
note('Hardcode segment = "All Snapp" / "Joint Snapp" / "Joint Tapsi" in each measure. '
     'Do not use a slicer for segment filtering.')

body('Sorting yearweek on x-axis / slicer', bold=True)
note('In the Model view, select yearweek (TEXT) → Sort by Column → yearweek_sort (INT). '
     'Apply to every imported view before building any visual.')

doc.add_paragraph()
body('End of Document – Driver Survey Power BI Routine Analysis Guide v5 (Complete)', bold=True, color=(80, 80, 80))

out = r'D:\Work\Driver Survey\PowerBI\PowerBI_Routine_Analysis_Guide_v5_Complete.docx'
doc.save(out)
print(f'Saved: {out}')
