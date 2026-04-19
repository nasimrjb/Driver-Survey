"""
Generates PowerBI_Routine_Analysis_Guide.docx
Run with: python build_guide_docx.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = os.path.join(os.path.dirname(__file__), "PowerBI_Routine_Analysis_Guide.docx")

# ── Color palette ──────────────────────────────────────────────────────────────
C_SNAPP   = RGBColor(0x00, 0xC8, 0x53)   # Snapp green
C_DARK    = RGBColor(0x1F, 0x39, 0x64)   # dark navy
C_ACCENT  = RGBColor(0x2E, 0x75, 0xB6)   # Power BI blue
C_CODE_BG = "F2F2F2"                      # light gray for code blocks
C_HEAD_BG = "2E75B6"                      # table header blue
C_WHITE   = "FFFFFF"

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'CCCCCC')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = C_DARK
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        run.font.color.rgb = C_ACCENT
        run.font.size = Pt(13)
        run.bold = True
    elif level == 3:
        run.font.color.rgb = C_DARK
        run.font.size = Pt(11)
        run.bold = True
    return p


def add_para(doc, text="", bold_parts=None, size=10):
    """Add a normal paragraph. bold_parts = list of substrings to bold."""
    if bold_parts is None:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.size = Pt(size)
        return p
    p = doc.add_paragraph()
    remaining = text
    for bp in bold_parts:
        idx = remaining.find(bp)
        if idx == -1:
            continue
        if idx > 0:
            r = p.add_run(remaining[:idx])
            r.font.size = Pt(size)
        r = p.add_run(bp)
        r.bold = True
        r.font.size = Pt(size)
        remaining = remaining[idx + len(bp):]
    if remaining:
        r = p.add_run(remaining)
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_code_block(doc, code_text):
    """Add a code block with monospace font and light gray background."""
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        # Set background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), C_CODE_BG)
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    # Add spacing after code block
    doc.add_paragraph()


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"ℹ  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x80)


def add_summary_table(doc, headers, rows):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, C_HEAD_BG)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F7F7F7" if r_idx % 2 == 0 else C_WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)

    set_cell_borders(table)
    return table


# ── SQL content ────────────────────────────────────────────────────────────────

SQL_VIEWS = [
    {
        "name": "vw_RA_SatReview",
        "subtitle": "replaces #3_Sat sheets",
        "purpose": "Satisfaction & Participation Review by city, week, and driver type. The cooperation_type column lets you slice by All / Part-Time / Full-Time in Power BI.",
        "n_cutoff": "n (controls Snapp columns), n_joint (controls Jnt_ columns)",
        "sql": """\
-- TRY_CAST(active_joint AS INT) in base CTE avoids nvarchar implicit conversion errors
CREATE OR ALTER VIEW [Cab].[vw_RA_SatReview] AS
WITH base AS (
    SELECT weeknumber, city, cooperation_type,
        TRY_CAST(active_joint AS INT) AS is_joint,
        snapp_gotmessage_text_incentive, tapsi_gotmessage_text_incentive,
        snapp_incentive_participation,   tapsi_incentive_participation,
        TRY_CAST(snapp_overall_incentive_satisfaction AS FLOAT) AS snapp_inc_sat,
        TRY_CAST(tapsi_overall_incentive_satisfaction AS FLOAT) AS tapsi_inc_sat,
        TRY_CAST(snapp_fare_satisfaction      AS FLOAT) AS snapp_fare_sat,
        TRY_CAST(tapsi_fare_satisfaction      AS FLOAT) AS tapsi_fare_sat,
        TRY_CAST(snapp_req_count_satisfaction AS FLOAT) AS snapp_req_sat,
        TRY_CAST(tapsi_req_count_satisfaction AS FLOAT) AS tapsi_req_sat,
        TRY_CAST(snapp_income_satisfaction    AS FLOAT) AS snapp_income_sat,
        TRY_CAST(tapsi_income_satisfaction    AS FLOAT) AS tapsi_income_sat
    FROM [Cab].[DriverSurvey_ShortMain] WHERE city IS NOT NULL
)
SELECT weeknumber, city, cooperation_type,
    COUNT(*) AS n,
    SUM(CASE WHEN is_joint=1 THEN 1 ELSE 0 END) AS n_joint,
    100.0*AVG(CASE WHEN snapp_incentive_participation='Yes' THEN 1.0 ELSE 0.0 END) AS Part_pct_Snapp,
    100.0*SUM(CASE WHEN is_joint=1 AND snapp_incentive_participation='Yes' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=1 THEN 1.0 ELSE 0.0 END),0) AS Part_pct_Jnt_Snapp,
    100.0*SUM(CASE WHEN is_joint=1 AND tapsi_incentive_participation='Yes' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=1 THEN 1.0 ELSE 0.0 END),0) AS Part_pct_Jnt_Tapsi,
    100.0*SUM(CASE WHEN snapp_gotmessage_text_incentive='Yes' AND snapp_incentive_participation='Yes' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN snapp_gotmessage_text_incentive='Yes' THEN 1.0 ELSE 0.0 END),0) AS Part_GotMsg_pct_Snapp,
    100.0*SUM(CASE WHEN is_joint=1 AND snapp_gotmessage_text_incentive='Yes' AND snapp_incentive_participation='Yes' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=1 AND snapp_gotmessage_text_incentive='Yes' THEN 1.0 ELSE 0.0 END),0) AS Part_GotMsg_pct_Jnt_Snapp,
    100.0*SUM(CASE WHEN is_joint=1 AND tapsi_gotmessage_text_incentive='Yes' AND tapsi_incentive_participation='Yes' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=1 AND tapsi_gotmessage_text_incentive='Yes' THEN 1.0 ELSE 0.0 END),0) AS Part_GotMsg_pct_Jnt_Tapsi,
    AVG(snapp_inc_sat) AS Incentive_Sat_Snapp,
    AVG(CASE WHEN is_joint=1 THEN snapp_inc_sat ELSE NULL END) AS Incentive_Sat_Jnt_Snapp,
    AVG(CASE WHEN is_joint=1 THEN tapsi_inc_sat ELSE NULL END) AS Incentive_Sat_Jnt_Tapsi,
    AVG(snapp_fare_sat) AS Fare_Sat_Snapp,
    AVG(CASE WHEN is_joint=1 THEN snapp_fare_sat ELSE NULL END) AS Fare_Sat_Jnt_Snapp,
    AVG(CASE WHEN is_joint=1 THEN tapsi_fare_sat ELSE NULL END) AS Fare_Sat_Jnt_Tapsi,
    AVG(snapp_req_sat) AS Request_Sat_Snapp,
    AVG(CASE WHEN is_joint=1 THEN snapp_req_sat ELSE NULL END) AS Request_Sat_Jnt_Snapp,
    AVG(CASE WHEN is_joint=1 THEN tapsi_req_sat ELSE NULL END) AS Request_Sat_Jnt_Tapsi,
    AVG(snapp_income_sat) AS Income_Sat_Snapp,
    AVG(CASE WHEN is_joint=1 THEN snapp_income_sat ELSE NULL END) AS Income_Sat_Jnt_Snapp,
    AVG(CASE WHEN is_joint=1 THEN tapsi_income_sat ELSE NULL END) AS Income_Sat_Jnt_Tapsi
FROM base GROUP BY weeknumber, city, cooperation_type;"""
    },
    {
        "name": "vw_RA_CitiesOverview",
        "subtitle": "replaces #12_Cities_Overview",
        "purpose": "Three independent n-groups (E=all drivers, F=joint drivers, G=competitor-signup drivers) with LOC averages and incentive message rates.",
        "n_cutoff": "E_n (pct_Joint, pct_Dual_SU, AvgLOC_All_Snapp, GotMsg_All_Snapp), F_n (AvgLOC_Joint_Snapp, GotMsg_Joint_Snapp, GotMsg_Joint_Cmpt), G_n (AvgLOC_Joint_Cmpt, AvgLOC_Joint_Cmpt_SU, GotMsg_Joint_Cmpt_SU)",
        "sql": """\
CREATE OR ALTER VIEW [Cab].[vw_RA_CitiesOverview] AS
WITH src AS (
    SELECT weeknumber, city,
        TRY_CAST(active_joint AS INT)  AS is_joint,
        TRY_CAST(snapp_LOC   AS FLOAT) AS snapp_loc_f,
        TRY_CAST(tapsi_LOC   AS FLOAT) AS tapsi_loc_f,
        snapp_gotmessage_text_incentive, tapsi_gotmessage_text_incentive
    FROM [Cab].[DriverSurvey_ShortMain] WHERE city IS NOT NULL
)
SELECT weeknumber, city,
    COUNT(*) AS E_n,
    SUM(CASE WHEN is_joint=1    THEN 1 ELSE 0 END) AS F_n,
    SUM(CASE WHEN tapsi_loc_f>0 THEN 1 ELSE 0 END) AS G_n,
    100.0*AVG(CASE WHEN is_joint=1    THEN 1.0 ELSE 0.0 END) AS pct_Joint,
    100.0*AVG(CASE WHEN tapsi_loc_f>0 THEN 1.0 ELSE 0.0 END) AS pct_Dual_SU,
    AVG(snapp_loc_f) AS AvgLOC_All_Snapp,
    100.0*AVG(CASE WHEN snapp_gotmessage_text_incentive='Yes' THEN 1.0 ELSE 0.0 END) AS GotMsg_All_Snapp,
    AVG(CASE WHEN is_joint=1 THEN snapp_loc_f ELSE NULL END) AS AvgLOC_Joint_Snapp,
    100.0*SUM(CASE WHEN is_joint=1 AND snapp_gotmessage_text_incentive='Yes' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=1 THEN 1.0 ELSE 0.0 END),0) AS GotMsg_Joint_Snapp,
    100.0*SUM(CASE WHEN is_joint=1 AND tapsi_gotmessage_text_incentive='Yes' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=1 THEN 1.0 ELSE 0.0 END),0) AS GotMsg_Joint_Cmpt,
    AVG(CASE WHEN is_joint=1    THEN tapsi_loc_f ELSE NULL END) AS AvgLOC_Joint_Cmpt,
    AVG(CASE WHEN tapsi_loc_f>0 THEN tapsi_loc_f ELSE NULL END) AS AvgLOC_Joint_Cmpt_SU,
    100.0*SUM(CASE WHEN tapsi_loc_f>0 AND tapsi_gotmessage_text_incentive='Yes' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN tapsi_loc_f>0 THEN 1.0 ELSE 0.0 END),0) AS GotMsg_Joint_Cmpt_SU
FROM src GROUP BY weeknumber, city;"""
    },
    {
        "name": "vw_RA_RideShare",
        "subtitle": "replaces #13_RideShare",
        "purpose": "Ride volume counts and share percentages per city and week. WoW delta is computed in DAX using a weeknumber offset.",
        "n_cutoff": "total_Res",
        "sql": """\
CREATE OR ALTER VIEW [Cab].[vw_RA_RideShare] AS
WITH src AS (
    SELECT weeknumber, city,
        TRY_CAST(active_joint AS INT)   AS is_joint,
        TRY_CAST(snapp_ride   AS FLOAT) AS snapp_f,
        TRY_CAST(tapsi_ride   AS FLOAT) AS tapsi_f
    FROM [Cab].[DriverSurvey_ShortMain] WHERE city IS NOT NULL
)
SELECT weeknumber, city,
    COUNT(*) AS total_Res,
    SUM(CASE WHEN is_joint=1 THEN 1 ELSE 0 END) AS Joint_Res,
    SUM(CASE WHEN is_joint=0 THEN 1 ELSE 0 END) AS Ex_drivers,
    ISNULL(SUM(snapp_f),0)+ISNULL(SUM(tapsi_f),0) AS Total_Ride,
    ISNULL(SUM(snapp_f),0) AS Total_Ride_Snapp,
    ISNULL(SUM(CASE WHEN is_joint=0 THEN snapp_f ELSE 0 END),0) AS Ex_Ride_Snapp,
    ISNULL(SUM(CASE WHEN is_joint=1 THEN snapp_f ELSE 0 END),0) AS Jnt_Snapp_Ride,
    ISNULL(SUM(CASE WHEN is_joint=1 THEN tapsi_f ELSE 0 END),0) AS Jnt_Tapsi_Ride,
    100.0*ISNULL(SUM(snapp_f),0)/NULLIF(ISNULL(SUM(snapp_f),0)+ISNULL(SUM(tapsi_f),0),0) AS All_Snapp_pct,
    100.0*ISNULL(SUM(CASE WHEN is_joint=0 THEN snapp_f ELSE 0 END),0)/NULLIF(ISNULL(SUM(snapp_f),0)+ISNULL(SUM(tapsi_f),0),0) AS Ex_Drivers_Snapp_pct,
    100.0*ISNULL(SUM(CASE WHEN is_joint=1 THEN snapp_f ELSE 0 END),0)/NULLIF(ISNULL(SUM(snapp_f),0)+ISNULL(SUM(tapsi_f),0),0) AS Jnt_at_Snapp_pct,
    100.0*ISNULL(SUM(CASE WHEN is_joint=1 THEN tapsi_f ELSE 0 END),0)/NULLIF(ISNULL(SUM(snapp_f),0)+ISNULL(SUM(tapsi_f),0),0) AS Jnt_at_Tapsi_pct
FROM src GROUP BY weeknumber, city;"""
    },
    {
        "name": "vw_RA_PersonaPartTime",
        "subtitle": "replaces #15_Persona_PartTime",
        "purpose": "Part-time driver percentages and ride-per-boarded metrics by city and week.",
        "n_cutoff": "total_Res (use Joint_Res for PT_pct_Joint, Ex_drivers for PT_pct_Exclusive)",
        "sql": """\
CREATE OR ALTER VIEW [Cab].[vw_RA_PersonaPartTime] AS
WITH src AS (
    SELECT weeknumber, city, cooperation_type,
        TRY_CAST(active_joint AS INT)   AS is_joint,
        TRY_CAST(snapp_ride   AS FLOAT) AS snapp_f,
        TRY_CAST(tapsi_ride   AS FLOAT) AS tapsi_f
    FROM [Cab].[DriverSurvey_ShortMain] WHERE city IS NOT NULL
)
SELECT weeknumber, city,
    COUNT(*) AS total_Res,
    SUM(CASE WHEN is_joint=1 THEN 1 ELSE 0 END) AS Joint_Res,
    SUM(CASE WHEN is_joint=0 THEN 1 ELSE 0 END) AS Ex_drivers,
    100.0*SUM(CASE WHEN is_joint=1 AND cooperation_type='Part-Time' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=1 THEN 1.0 ELSE 0.0 END),0) AS PT_pct_Joint,
    100.0*SUM(CASE WHEN is_joint=0 AND cooperation_type='Part-Time' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=0 THEN 1.0 ELSE 0.0 END),0) AS PT_pct_Exclusive,
    ISNULL(SUM(CASE WHEN is_joint=1 THEN snapp_f ELSE 0 END),0)
        /NULLIF(SUM(CASE WHEN is_joint=1 THEN 1.0 ELSE 0.0 END),0) AS RidePerBoarded_Snapp,
    ISNULL(SUM(CASE WHEN is_joint=1 THEN tapsi_f ELSE 0 END),0)
        /NULLIF(SUM(CASE WHEN is_joint=1 THEN 1.0 ELSE 0.0 END),0) AS RidePerBoarded_Tapsi,
    ISNULL(SUM(snapp_f),0)/NULLIF(COUNT(*),0) AS AvgAllRides
FROM src GROUP BY weeknumber, city;"""
    },
    {
        "name": "vw_RA_IncentiveAmounts",
        "subtitle": "replaces #1 and #2 sheets (long format)",
        "purpose": "Incentive amount bucket (rial range) distribution per city, week, and platform. Long format: use incentive_range as Matrix columns in Power BI. For #2, filter platform = 'Tapsi'.",
        "n_cutoff": "n_total",
        "sql": """\
CREATE OR ALTER VIEW [Cab].[vw_RA_IncentiveAmounts] AS
-- Snapp incentive amounts (all drivers)
SELECT
    weeknumber, city, 'Snapp' AS platform,
    snapp_incentive_rial_details AS incentive_range,
    COUNT(*) AS n_range,
    SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total,
    100.0 * COUNT(*) / NULLIF(SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city),0) AS pct
FROM [Cab].[DriverSurvey_ShortMain]
WHERE city IS NOT NULL AND snapp_incentive_rial_details IS NOT NULL
GROUP BY weeknumber, city, snapp_incentive_rial_details
UNION ALL
-- Tapsi incentive amounts (joint drivers only)
SELECT
    weeknumber, city, 'Tapsi' AS platform,
    tapsi_incentive_rial_details AS incentive_range,
    COUNT(*) AS n_range,
    SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total,
    100.0 * COUNT(*) / NULLIF(SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city),0) AS pct
FROM [Cab].[DriverSurvey_ShortMain]
WHERE city IS NOT NULL AND tapsi_incentive_rial_details IS NOT NULL
  AND TRY_CAST(active_joint AS INT) = 1
GROUP BY weeknumber, city, tapsi_incentive_rial_details;"""
    },
    {
        "name": "vw_RA_IncentiveDuration",
        "subtitle": "replaces #4_Incentive_Duration (long format)",
        "purpose": "Incentive active-duration bucket distribution per city, week, and platform. Long format: use duration_bucket as Matrix columns.",
        "n_cutoff": "n_total",
        "sql": """\
CREATE OR ALTER VIEW [Cab].[vw_RA_IncentiveDuration] AS
SELECT weeknumber, city, 'Snapp' AS platform,
    snapp_incentive_active_duration AS duration_bucket,
    COUNT(*) AS n_range,
    SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total,
    100.0 * COUNT(*) / NULLIF(SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city),0) AS pct
FROM [Cab].[DriverSurvey_ShortMain]
WHERE city IS NOT NULL AND snapp_incentive_active_duration IS NOT NULL
GROUP BY weeknumber, city, snapp_incentive_active_duration
UNION ALL
SELECT weeknumber, city, 'Tapsi' AS platform,
    tapsi_incentive_active_duration AS duration_bucket,
    COUNT(*) AS n_range,
    SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total,
    100.0 * COUNT(*) / NULLIF(SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city),0) AS pct
FROM [Cab].[DriverSurvey_ShortMain]
WHERE city IS NOT NULL AND tapsi_incentive_active_duration IS NOT NULL
GROUP BY weeknumber, city, tapsi_incentive_active_duration;"""
    },
    {
        "name": "vw_RA_Persona",
        "subtitle": "replaces all #15_Persona sub-sheets (long format)",
        "purpose": "All demographic crosstabs in one long-format view. Use the 'dimension' column as a slicer (Activity Type, Age Group, Education, Gender, etc.) and 'category' as Matrix columns.",
        "n_cutoff": "n_total",
        "sql": """\
CREATE OR ALTER VIEW [Cab].[vw_RA_Persona] AS
-- CAST all category columns to NVARCHAR to prevent UNION ALL type-precedence error
-- when edu/marr_stat are stored as numeric in ShortMain
WITH activity AS (
    SELECT weeknumber, city, 'Activity Type' AS dimension,
        CAST(active_time      AS NVARCHAR(100)) AS category,
        COUNT(*) AS n, SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total
    FROM [Cab].[DriverSurvey_ShortMain]
    WHERE city IS NOT NULL AND active_time IS NOT NULL
    GROUP BY weeknumber, city, active_time),
age_grp AS (
    SELECT weeknumber, city, 'Age Group' AS dimension,
        CAST(age_group        AS NVARCHAR(100)) AS category,
        COUNT(*) AS n, SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total
    FROM [Cab].[DriverSurvey_ShortMain]
    WHERE city IS NOT NULL AND age_group IS NOT NULL
    GROUP BY weeknumber, city, age_group),
edu AS (
    SELECT weeknumber, city, 'Education' AS dimension,
        CAST(edu              AS NVARCHAR(100)) AS category,
        COUNT(*) AS n, SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total
    FROM [Cab].[DriverSurvey_ShortMain]
    WHERE city IS NOT NULL AND edu IS NOT NULL
    GROUP BY weeknumber, city, edu),
marr AS (
    SELECT weeknumber, city, 'Marital Status' AS dimension,
        CAST(marr_stat        AS NVARCHAR(100)) AS category,
        COUNT(*) AS n, SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total
    FROM [Cab].[DriverSurvey_ShortMain]
    WHERE city IS NOT NULL AND marr_stat IS NOT NULL
    GROUP BY weeknumber, city, marr_stat),
gen AS (
    SELECT weeknumber, city, 'Gender' AS dimension,
        CAST(gender           AS NVARCHAR(100)) AS category,
        COUNT(*) AS n, SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total
    FROM [Cab].[DriverSurvey_ShortMain]
    WHERE city IS NOT NULL AND gender IS NOT NULL
    GROUP BY weeknumber, city, gender),
coop AS (
    SELECT weeknumber, city, 'Cooperation Type' AS dimension,
        CAST(cooperation_type AS NVARCHAR(100)) AS category,
        COUNT(*) AS n, SUM(COUNT(*)) OVER (PARTITION BY weeknumber, city) AS n_total
    FROM [Cab].[DriverSurvey_ShortMain]
    WHERE city IS NOT NULL AND cooperation_type IS NOT NULL
    GROUP BY weeknumber, city, cooperation_type)
SELECT *, 100.0*n/NULLIF(n_total,0) AS pct FROM activity
UNION ALL SELECT *, 100.0*n/NULLIF(n_total,0) AS pct FROM age_grp
UNION ALL SELECT *, 100.0*n/NULLIF(n_total,0) AS pct FROM edu
UNION ALL SELECT *, 100.0*n/NULLIF(n_total,0) AS pct FROM marr
UNION ALL SELECT *, 100.0*n/NULLIF(n_total,0) AS pct FROM gen
UNION ALL SELECT *, 100.0*n/NULLIF(n_total,0) AS pct FROM coop;"""
    },
    {
        "name": "vw_RA_CommFree",
        "subtitle": "replaces #18_CommFree sheets",
        "purpose": "Commission-free incentive analysis: message receipt counts, incentive category breakdown, and free-commission ride share per city, week, and platform.",
        "n_cutoff": "n",
        "sql": """\
CREATE OR ALTER VIEW [Cab].[vw_RA_CommFree] AS
WITH src AS (
    SELECT weeknumber, city,
        TRY_CAST(active_joint   AS INT)    AS is_joint,
        TRY_CAST(snapp_commfree AS FLOAT)  AS snapp_cf,
        TRY_CAST(tapsi_commfree AS FLOAT)  AS tapsi_cf,
        snapp_gotmessage_text_incentive, tapsi_gotmessage_text_incentive,
        CAST(snapp_incentive_category AS NVARCHAR(100)) AS snapp_inc_cat,
        CAST(tapsi_incentive_category AS NVARCHAR(100)) AS tapsi_inc_cat
    FROM [Cab].[DriverSurvey_ShortMain] WHERE city IS NOT NULL
)
SELECT weeknumber, city, 'Snapp' AS platform,
    COUNT(*) AS n,
    SUM(CASE WHEN snapp_gotmessage_text_incentive='Yes' THEN 1 ELSE 0 END) AS Who_Got_Message,
    SUM(CASE WHEN snapp_gotmessage_text_incentive='Yes' AND snapp_inc_cat='Money' THEN 1 ELSE 0 END) AS GotMsg_Money,
    SUM(CASE WHEN snapp_gotmessage_text_incentive='Yes' AND snapp_inc_cat='Free-Commission' THEN 1 ELSE 0 END) AS GotMsg_FreeComm,
    SUM(CASE WHEN snapp_gotmessage_text_incentive='Yes' AND snapp_inc_cat='Money & Free-commission' THEN 1 ELSE 0 END) AS GotMsg_Money_FreeComm,
    SUM(CASE WHEN snapp_cf>0 THEN 1 ELSE 0 END) AS Free_Comm_Drivers,
    100.0*SUM(CASE WHEN snapp_gotmessage_text_incentive='Yes' THEN 1.0 ELSE 0.0 END)/NULLIF(COUNT(*),0) AS pct_Got_Message,
    100.0*SUM(CASE WHEN snapp_cf>0 THEN 1.0 ELSE 0.0 END)/NULLIF(COUNT(*),0) AS pct_Free_Comm_Ride
FROM src GROUP BY weeknumber, city
UNION ALL
SELECT weeknumber, city, 'Tapsi' AS platform,
    SUM(CASE WHEN is_joint=1 THEN 1 ELSE 0 END) AS n,
    SUM(CASE WHEN is_joint=1 AND tapsi_gotmessage_text_incentive='Yes' THEN 1 ELSE 0 END) AS Who_Got_Message,
    SUM(CASE WHEN is_joint=1 AND tapsi_gotmessage_text_incentive='Yes' AND tapsi_inc_cat='Money' THEN 1 ELSE 0 END) AS GotMsg_Money,
    SUM(CASE WHEN is_joint=1 AND tapsi_gotmessage_text_incentive='Yes' AND tapsi_inc_cat='Free-Commission' THEN 1 ELSE 0 END) AS GotMsg_FreeComm,
    SUM(CASE WHEN is_joint=1 AND tapsi_gotmessage_text_incentive='Yes' AND tapsi_inc_cat='Money & Free-commission' THEN 1 ELSE 0 END) AS GotMsg_Money_FreeComm,
    SUM(CASE WHEN is_joint=1 AND tapsi_cf>0 THEN 1 ELSE 0 END) AS Free_Comm_Drivers,
    100.0*SUM(CASE WHEN is_joint=1 AND tapsi_gotmessage_text_incentive='Yes' THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=1 THEN 1.0 ELSE 0.0 END),0) AS pct_Got_Message,
    100.0*SUM(CASE WHEN is_joint=1 AND tapsi_cf>0 THEN 1.0 ELSE 0.0 END)
        /NULLIF(SUM(CASE WHEN is_joint=1 THEN 1.0 ELSE 0.0 END),0) AS pct_Free_Comm_Ride
FROM src GROUP BY weeknumber, city;"""
    },
    {
        "name": "vw_RA_CSRare",
        "subtitle": "replaces #CS_Sat_Snapp and #CS_Sat_Tapsi sheets",
        "purpose": "Customer Support satisfaction scores (1-5 scale) and resolution rates from DriverSurvey_ShortRare, joined to ShortMain for city and weeknumber.",
        "n_cutoff": "n",
        "sql": """\
CREATE OR ALTER VIEW [Cab].[vw_RA_CSRare] AS
SELECT
    sm.weeknumber, sm.city,
    COUNT(*) AS n,
    AVG(TRY_CAST(sr.snapp_CS_satisfaction_overall   AS FLOAT)) AS Snapp_CS_Overall,
    AVG(TRY_CAST(sr.snapp_CS_satisfaction_waittime  AS FLOAT)) AS Snapp_CS_WaitTime,
    AVG(TRY_CAST(sr.snapp_CS_satisfaction_solution  AS FLOAT)) AS Snapp_CS_Solution,
    AVG(TRY_CAST(sr.snapp_CS_satisfaction_behaviour AS FLOAT)) AS Snapp_CS_Behaviour,
    AVG(TRY_CAST(sr.snapp_CS_satisfaction_relevance AS FLOAT)) AS Snapp_CS_Relevance,
    100.0 * AVG(CASE WHEN sr.snapp_CS_solved='Yes' THEN 1.0 ELSE 0.0 END) AS Snapp_CS_Solved_pct,
    AVG(TRY_CAST(sr.tapsi_CS_satisfaction_overall   AS FLOAT)) AS Tapsi_CS_Overall,
    AVG(TRY_CAST(sr.tapsi_CS_satisfaction_waittime  AS FLOAT)) AS Tapsi_CS_WaitTime,
    AVG(TRY_CAST(sr.tapsi_CS_satisfaction_solution  AS FLOAT)) AS Tapsi_CS_Solution,
    AVG(TRY_CAST(sr.tapsi_CS_satisfaction_behaviour AS FLOAT)) AS Tapsi_CS_Behaviour,
    AVG(TRY_CAST(sr.tapsi_CS_satisfaction_relevance AS FLOAT)) AS Tapsi_CS_Relevance,
    100.0 * AVG(CASE WHEN sr.tapsi_CS_solved='Yes' THEN 1.0 ELSE 0.0 END) AS Tapsi_CS_Solved_pct
FROM [Cab].[DriverSurvey_ShortRare]  sr
JOIN [Cab].[DriverSurvey_ShortMain]  sm ON sm.recordID = sr.recordID
WHERE sm.city IS NOT NULL
GROUP BY sm.weeknumber, sm.city;"""
    },
    {
        "name": "vw_RA_NavReco",
        "subtitle": "replaces #NavReco_Scores and #Reco_NPS sheets",
        "purpose": "Navigation app recommendation scores (1-10) and NPS recommend scores from DriverSurvey_ShortRare joined to ShortMain.",
        "n_cutoff": "n",
        "sql": """\
CREATE OR ALTER VIEW [Cab].[vw_RA_NavReco] AS
SELECT
    sm.weeknumber, sm.city,
    COUNT(*) AS n,
    AVG(TRY_CAST(sr.snapp_recommend               AS FLOAT)) AS Snapp_NPS,
    AVG(TRY_CAST(sr.snappdriver_tapsi_recommend    AS FLOAT)) AS Tapsi_NPS_SnapDriver,
    AVG(TRY_CAST(sr.tapsidriver_tapsi_recommend    AS FLOAT)) AS Tapsi_NPS_TapsiDriver,
    AVG(TRY_CAST(sr.recommendation_googlemap       AS FLOAT)) AS Reco_GoogleMap,
    AVG(TRY_CAST(sr.recommendation_waze            AS FLOAT)) AS Reco_Waze,
    AVG(TRY_CAST(sr.recommendation_neshan          AS FLOAT)) AS Reco_Neshan,
    AVG(TRY_CAST(sr.recommendation_balad           AS FLOAT)) AS Reco_Balad,
    AVG(TRY_CAST(sr.snapp_navigation_app_satisfaction    AS FLOAT)) AS Snapp_Nav_Sat,
    AVG(TRY_CAST(sr.tapsi_in_app_navigation_satisfaction AS FLOAT)) AS Tapsi_Nav_Sat
FROM [Cab].[DriverSurvey_ShortRare]  sr
JOIN [Cab].[DriverSurvey_ShortMain]  sm ON sm.recordID = sr.recordID
WHERE sm.city IS NOT NULL
GROUP BY sm.weeknumber, sm.city;"""
    },
]

DAX_SAT_REVIEW = """\
-- vw_RA_SatReview — 18 measures
-- n guards Snapp columns; n_joint guards Jnt_ columns
-- Percentages stored 0-100 → divide by 100. Satisfaction scores 1-5 → AVERAGE directly.

Part Pct Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n]) >= MinN, AVERAGE(vw_RA_SatReview[Part_pct_Snapp])/100, BLANK())

Part Pct Jnt Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Part_pct_Jnt_Snapp])/100, BLANK())

Part Pct Jnt Tapsi =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Part_pct_Jnt_Tapsi])/100, BLANK())

Part GotMsg Pct Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n]) >= MinN, AVERAGE(vw_RA_SatReview[Part_GotMsg_pct_Snapp])/100, BLANK())

Part GotMsg Pct Jnt Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Part_GotMsg_pct_Jnt_Snapp])/100, BLANK())

Part GotMsg Pct Jnt Tapsi =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Part_GotMsg_pct_Jnt_Tapsi])/100, BLANK())

Incentive Sat Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n]) >= MinN, AVERAGE(vw_RA_SatReview[Incentive_Sat_Snapp]), BLANK())

Incentive Sat Jnt Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Incentive_Sat_Jnt_Snapp]), BLANK())

Incentive Sat Jnt Tapsi =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Incentive_Sat_Jnt_Tapsi]), BLANK())

Fare Sat Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n]) >= MinN, AVERAGE(vw_RA_SatReview[Fare_Sat_Snapp]), BLANK())

Fare Sat Jnt Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Fare_Sat_Jnt_Snapp]), BLANK())

Fare Sat Jnt Tapsi =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Fare_Sat_Jnt_Tapsi]), BLANK())

Request Sat Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n]) >= MinN, AVERAGE(vw_RA_SatReview[Request_Sat_Snapp]), BLANK())

Request Sat Jnt Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Request_Sat_Jnt_Snapp]), BLANK())

Request Sat Jnt Tapsi =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Request_Sat_Jnt_Tapsi]), BLANK())

Income Sat Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n]) >= MinN, AVERAGE(vw_RA_SatReview[Income_Sat_Snapp]), BLANK())

Income Sat Jnt Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Income_Sat_Jnt_Snapp]), BLANK())

Income Sat Jnt Tapsi =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_SatReview[n_joint]) >= MinN, AVERAGE(vw_RA_SatReview[Income_Sat_Jnt_Tapsi]), BLANK())"""

DAX_CITIES = """\
-- vw_RA_CitiesOverview — 10 measures
-- Three independent n guards: E_n (all drivers), F_n (joint), G_n (competitor-signup)
-- Percentages stored 0-100 → divide by 100. LOC values → AVERAGE directly.

-- E-group (guarded by E_n)
Pct Joint =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[E_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[pct_Joint])/100, BLANK())

Pct Dual SU =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[E_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[pct_Dual_SU])/100, BLANK())

AvgLOC All Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[E_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[AvgLOC_All_Snapp]), BLANK())

GotMsg All Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[E_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[GotMsg_All_Snapp])/100, BLANK())

-- F-group (guarded by F_n)
AvgLOC Joint Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[F_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[AvgLOC_Joint_Snapp]), BLANK())

GotMsg Joint Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[F_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[GotMsg_Joint_Snapp])/100, BLANK())

GotMsg Joint Cmpt =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[F_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[GotMsg_Joint_Cmpt])/100, BLANK())

-- G-group (guarded by G_n)
AvgLOC Joint Cmpt =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[G_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[AvgLOC_Joint_Cmpt]), BLANK())

AvgLOC Joint Cmpt SU =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[G_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[AvgLOC_Joint_Cmpt_SU]), BLANK())

GotMsg Joint Cmpt SU =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CitiesOverview[G_n]) >= MinN, AVERAGE(vw_RA_CitiesOverview[GotMsg_Joint_Cmpt_SU])/100, BLANK())"""

DAX_RIDESHARE = """\
-- vw_RA_RideShare — 11 measures
-- n-cutoff column: total_Res
-- Count columns → SUM with cutoff. Percentage columns → AVERAGE / 100.

Joint Res =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, SUM(vw_RA_RideShare[Joint_Res]), BLANK())

Ex Drivers =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, SUM(vw_RA_RideShare[Ex_drivers]), BLANK())

Total Ride =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, SUM(vw_RA_RideShare[Total_Ride]), BLANK())

Total Ride Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, SUM(vw_RA_RideShare[Total_Ride_Snapp]), BLANK())

Ex Ride Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, SUM(vw_RA_RideShare[Ex_Ride_Snapp]), BLANK())

Jnt Snapp Ride =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, SUM(vw_RA_RideShare[Jnt_Snapp_Ride]), BLANK())

Jnt Tapsi Ride =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, SUM(vw_RA_RideShare[Jnt_Tapsi_Ride]), BLANK())

All Snapp Pct =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, AVERAGE(vw_RA_RideShare[All_Snapp_pct])/100, BLANK())

Ex Drivers Snapp Pct =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, AVERAGE(vw_RA_RideShare[Ex_Drivers_Snapp_pct])/100, BLANK())

Jnt at Snapp Pct =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, AVERAGE(vw_RA_RideShare[Jnt_at_Snapp_pct])/100, BLANK())

Jnt at Tapsi Pct =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_RideShare[total_Res]) >= MinN, AVERAGE(vw_RA_RideShare[Jnt_at_Tapsi_pct])/100, BLANK())"""

DAX_PERSONAPT = """\
-- vw_RA_PersonaPartTime — 5 measures
-- n-cutoffs: Joint_Res (PT_pct_Joint, RidePerBoarded_*), Ex_drivers (PT_pct_Exclusive), total_Res (AvgAllRides)

PT Pct Joint =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_PersonaPartTime[Joint_Res]) >= MinN, AVERAGE(vw_RA_PersonaPartTime[PT_pct_Joint])/100, BLANK())

PT Pct Exclusive =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_PersonaPartTime[Ex_drivers]) >= MinN, AVERAGE(vw_RA_PersonaPartTime[PT_pct_Exclusive])/100, BLANK())

RidePerBoarded Snapp =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_PersonaPartTime[Joint_Res]) >= MinN, AVERAGE(vw_RA_PersonaPartTime[RidePerBoarded_Snapp]), BLANK())

RidePerBoarded Tapsi =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_PersonaPartTime[Joint_Res]) >= MinN, AVERAGE(vw_RA_PersonaPartTime[RidePerBoarded_Tapsi]), BLANK())

AvgAllRides =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_PersonaPartTime[total_Res]) >= MinN, AVERAGE(vw_RA_PersonaPartTime[AvgAllRides]), BLANK())"""

DAX_INCAMT = """\
-- vw_RA_IncentiveAmounts — 1 measure
-- Matrix setup: city = Rows, incentive_range = Columns, platform = Slicer

Incentive Pct =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_IncentiveAmounts[n_total]) >= MinN, AVERAGE(vw_RA_IncentiveAmounts[pct])/100, BLANK())"""

DAX_INCDUR = """\
-- vw_RA_IncentiveDuration — 1 measure
-- Matrix setup: city = Rows, duration_bucket = Columns, platform = Slicer

Duration Pct =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_IncentiveDuration[n_total]) >= MinN, AVERAGE(vw_RA_IncentiveDuration[pct])/100, BLANK())"""

DAX_PERSONA = """\
-- vw_RA_Persona — 1 measure
-- Matrix setup: city = Rows, category = Columns, dimension = Slicer

Persona Pct =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_Persona[n_total]) >= MinN, AVERAGE(vw_RA_Persona[pct])/100, BLANK())"""

DAX_COMMFREE = """\
-- vw_RA_CommFree — 7 measures
-- n-cutoff column: n  |  platform = Slicer (Snapp / Tapsi)
-- Count columns → SUM with cutoff. Percentage columns → AVERAGE / 100.

Who Got Message =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CommFree[n]) >= MinN, SUM(vw_RA_CommFree[Who_Got_Message]), BLANK())

GotMsg Money =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CommFree[n]) >= MinN, SUM(vw_RA_CommFree[GotMsg_Money]), BLANK())

GotMsg FreeComm =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CommFree[n]) >= MinN, SUM(vw_RA_CommFree[GotMsg_FreeComm]), BLANK())

GotMsg Money FreeComm =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CommFree[n]) >= MinN, SUM(vw_RA_CommFree[GotMsg_Money_FreeComm]), BLANK())

Free Comm Drivers =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CommFree[n]) >= MinN, SUM(vw_RA_CommFree[Free_Comm_Drivers]), BLANK())

Pct Got Message =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CommFree[n]) >= MinN, AVERAGE(vw_RA_CommFree[pct_Got_Message])/100, BLANK())

Pct Free Comm Ride =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CommFree[n]) >= MinN, AVERAGE(vw_RA_CommFree[pct_Free_Comm_Ride])/100, BLANK())"""

DAX_CSRARE = """\
-- vw_RA_CSRare — 12 measures
-- n-cutoff column: n
-- Satisfaction scores 1-5 → AVERAGE directly. Solved % stored 0-100 → divide by 100.

Snapp CS Overall =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Snapp_CS_Overall]), BLANK())

Snapp CS WaitTime =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Snapp_CS_WaitTime]), BLANK())

Snapp CS Solution =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Snapp_CS_Solution]), BLANK())

Snapp CS Behaviour =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Snapp_CS_Behaviour]), BLANK())

Snapp CS Relevance =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Snapp_CS_Relevance]), BLANK())

Snapp CS Solved Pct =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Snapp_CS_Solved_pct])/100, BLANK())

Tapsi CS Overall =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Tapsi_CS_Overall]), BLANK())

Tapsi CS WaitTime =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Tapsi_CS_WaitTime]), BLANK())

Tapsi CS Solution =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Tapsi_CS_Solution]), BLANK())

Tapsi CS Behaviour =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Tapsi_CS_Behaviour]), BLANK())

Tapsi CS Relevance =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Tapsi_CS_Relevance]), BLANK())

Tapsi CS Solved Pct =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_CSRare[n]) >= MinN, AVERAGE(vw_RA_CSRare[Tapsi_CS_Solved_pct])/100, BLANK())"""

DAX_NAVRECO = """\
-- vw_RA_NavReco — 9 measures
-- n-cutoff column: n
-- All scores are 0-10 → AVERAGE directly (do NOT divide by 100)

Snapp NPS =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_NavReco[n]) >= MinN, AVERAGE(vw_RA_NavReco[Snapp_NPS]), BLANK())

Tapsi NPS SnapDriver =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_NavReco[n]) >= MinN, AVERAGE(vw_RA_NavReco[Tapsi_NPS_SnapDriver]), BLANK())

Tapsi NPS TapsiDriver =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_NavReco[n]) >= MinN, AVERAGE(vw_RA_NavReco[Tapsi_NPS_TapsiDriver]), BLANK())

Reco GoogleMap =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_NavReco[n]) >= MinN, AVERAGE(vw_RA_NavReco[Reco_GoogleMap]), BLANK())

Reco Waze =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_NavReco[n]) >= MinN, AVERAGE(vw_RA_NavReco[Reco_Waze]), BLANK())

Reco Neshan =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_NavReco[n]) >= MinN, AVERAGE(vw_RA_NavReco[Reco_Neshan]), BLANK())

Reco Balad =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_NavReco[n]) >= MinN, AVERAGE(vw_RA_NavReco[Reco_Balad]), BLANK())

Snapp Nav Sat =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_NavReco[n]) >= MinN, AVERAGE(vw_RA_NavReco[Snapp_Nav_Sat]), BLANK())

Tapsi Nav Sat =
VAR MinN = [Min N Cutoff Value]
RETURN IF(SUM(vw_RA_NavReco[n]) >= MinN, AVERAGE(vw_RA_NavReco[Tapsi_Nav_Sat]), BLANK())"""

DAX_WOW = """\
-- Week-over-Week pattern — apply to any measure by referencing it inside CALCULATE.
-- Works because the measure already encodes the n-cutoff.

WoW Fare Sat Snapp =
VAR SelectedWeek = SELECTEDVALUE(vw_RA_SatReview[weeknumber])
VAR CurrentVal   = CALCULATE([Fare Sat Snapp], vw_RA_SatReview[weeknumber] = SelectedWeek)
VAR PrevVal      = CALCULATE([Fare Sat Snapp], vw_RA_SatReview[weeknumber] = SelectedWeek - 1)
RETURN IF(NOT ISBLANK(CurrentVal) && NOT ISBLANK(PrevVal), CurrentVal - PrevVal, BLANK())

WoW Pct Joint =
VAR SelectedWeek = SELECTEDVALUE(vw_RA_CitiesOverview[weeknumber])
VAR CurrentVal   = CALCULATE([Pct Joint], vw_RA_CitiesOverview[weeknumber] = SelectedWeek)
VAR PrevVal      = CALCULATE([Pct Joint], vw_RA_CitiesOverview[weeknumber] = SelectedWeek - 1)
RETURN IF(NOT ISBLANK(CurrentVal) && NOT ISBLANK(PrevVal), CurrentVal - PrevVal, BLANK())

WoW Jnt at Tapsi Pct =
VAR SelectedWeek = SELECTEDVALUE(vw_RA_RideShare[weeknumber])
VAR CurrentVal   = CALCULATE([Jnt at Tapsi Pct], vw_RA_RideShare[weeknumber] = SelectedWeek)
VAR PrevVal      = CALCULATE([Jnt at Tapsi Pct], vw_RA_RideShare[weeknumber] = SelectedWeek - 1)
RETURN IF(NOT ISBLANK(CurrentVal) && NOT ISBLANK(PrevVal), CurrentVal - PrevVal, BLANK())"""

SUMMARY_ROWS = [
    ["vw_RA_SatReview",      "#3_Sat (All, Part-Time, Full-Time)", "n, n_joint"],
    ["vw_RA_CitiesOverview", "#12_Cities_Overview",                "E_n, F_n, G_n"],
    ["vw_RA_RideShare",      "#13_RideShare",                      "total_Res"],
    ["vw_RA_PersonaPartTime","#15_Persona_PartTime",               "total_Res"],
    ["vw_RA_IncentiveAmounts","#1_Snapp_Incentive_Amt, #2_Tapsi",  "n_total"],
    ["vw_RA_IncentiveDuration","#4_Incentive_Duration",            "n_total"],
    ["vw_RA_Persona",        "#15_Persona (all sub-sheets)",       "n_total"],
    ["vw_RA_CommFree",       "#18_CommFree_Snapp, #18_CommFree_Tapsi", "n"],
    ["vw_RA_CSRare",         "#CS_Sat_Snapp, #CS_Sat_Tapsi",      "n"],
    ["vw_RA_NavReco",        "#NavReco_Scores, #Reco_NPS",         "n"],
]

PAGE_ROWS = [
    ["Satisfaction Review",    "#3_Sat",    "vw_RA_SatReview",       "weeknumber, cooperation_type", "city", "18 DAX measures"],
    ["All Cities Overview",    "#12",       "vw_RA_CitiesOverview",   "weeknumber",                  "city", "E_n, F_n, G_n + 10 metrics"],
    ["Ride Share",             "#13",       "vw_RA_RideShare",        "weeknumber",                  "city", "8 counts + 4 % measures"],
    ["Incentive Amounts",      "#1/#2",     "vw_RA_IncentiveAmounts", "weeknumber, platform",        "city", "incentive_range → columns"],
    ["Incentive Duration",     "#4",        "vw_RA_IncentiveDuration","weeknumber, platform",        "city", "duration_bucket → columns"],
    ["Driver Persona",         "#15",       "vw_RA_Persona",          "weeknumber, dimension",       "city", "category → columns"],
    ["Part-Time & Rides",      "#15_PT",    "vw_RA_PersonaPartTime",  "weeknumber",                  "city", "7 measures"],
    ["Commission Free",        "#18",       "vw_RA_CommFree",         "weeknumber, platform",        "city", "counts + 2 % measures"],
    ["Customer Support",       "#CS_Sat",   "vw_RA_CSRare",           "weeknumber",                  "city", "12 sat/pct measures"],
    ["Navigation & NPS",       "#NavReco",  "vw_RA_NavReco",          "weeknumber",                  "city", "9 score measures"],
]


# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page setup: landscape A4
section = doc.sections[0]
section.page_width  = Inches(11.69)   # A4 landscape width
section.page_height = Inches(8.27)
section.left_margin   = Inches(0.8)
section.right_margin  = Inches(0.8)
section.top_margin    = Inches(0.75)
section.bottom_margin = Inches(0.75)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── Title page ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Power BI Dashboard")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = C_DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Driver Survey — Routine Analysis")
run2.font.size = Pt(22)
run2.font.bold = True
run2.font.color.rgb = C_ACCENT

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Replication Guide: From Python/Excel to Power BI")
run3.font.size = Pt(13)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# Server info box
info_table = doc.add_table(rows=1, cols=3)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = info_table.rows[0].cells
labels = ["SQL Server", "Database", "Schema"]
values = ["192.168.18.37", "Cab_Studies", "[Cab]"]
for i, (lbl, val) in enumerate(zip(labels, values)):
    cell = cells[i]
    set_cell_bg(cell, "E8F0FE")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{lbl}\n")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    r2 = p.add_run(val)
    r2.font.size = Pt(11)
    r2.font.bold = True
    r2.font.color.rgb = C_DARK
set_cell_borders(info_table)

doc.add_paragraph()

# ── Introduction ───────────────────────────────────────────────────────────────
add_heading(doc, "Introduction", 1)
add_para(doc, (
    "The Python script survey_routine_analysis.py produces weekly Excel reports "
    "where each sheet is a City × Metric matrix — cities as rows, computed metrics "
    "(percentages, satisfaction averages, counts) as columns, filtered to a selected week. "
    "Every sheet includes one or more n-cutoff columns (n, n_joint, Snapp_n, Tapsi_n, E_n, F_n, G_n) "
    "that control data visibility: rows where n falls below a threshold are blanked."
))
doc.add_paragraph()
add_para(doc, "This guide covers four steps:")
add_bullet(doc, "Step 1 — Create SQL Server views that pre-aggregate the raw survey tables")
add_bullet(doc, "Step 2 — Connect Power BI Desktop to those views (Import mode)")
add_bullet(doc, "Step 3 — Build matrix report pages with city rows and metric columns")
add_bullet(doc, "Step 4 — Implement the n-cutoff What-If Parameter with DAX measures")

# ── Part 1: SQL Views ──────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Part 1 — SQL Views to Create on the Server", 1)
add_para(doc, (
    "All 10 routine-analysis views are included in create_views.sql (same folder as this guide). "
    "Open that file in SSMS, connect to Cab_Studies on 192.168.18.37, and run the entire script — "
    "it creates all 30 views in one step (20 dashboard views + 10 vw_RA_* views). "
    "The script prints 'All 30 views created successfully' when done."
))
add_heading(doc, "How to run create_views.sql in SSMS", 3)
ssms_steps = [
    "Open SQL Server Management Studio and connect to 192.168.18.37.",
    "File → Open → File…  and select create_views.sql.",
    "From the database drop-down in the toolbar select Cab_Studies.",
    "Press F5 (or click Execute). The script will drop-and-recreate all existing views "
    "and create the 10 new vw_RA_* views.",
    "Check the Messages tab at the bottom — it should end with: "
    "All 30 views created successfully.",
]
for s in ssms_steps:
    add_numbered(doc, s)

add_note(doc, (
    "The vw_RA_* views use CREATE OR ALTER VIEW so they are safe to re-run at any time "
    "— they update in place without dropping first. "
    "Re-run the script after any schema change to the source tables."
))
doc.add_paragraph()

add_heading(doc, "Views created — reference", 2)
add_para(doc, "The table below lists the 10 routine-analysis views, what Python sheet each replaces, and which n-cutoff column(s) to use in DAX measures.")
add_summary_table(
    doc,
    ["View", "Python Sheet(s)", "N-Cutoff Column(s)", "Key columns"],
    [
        ["vw_RA_SatReview",       "#3_Sat (All / PT / FT)",              "n, n_joint",    "cooperation_type slicer; 18 metric columns"],
        ["vw_RA_CitiesOverview",  "#12_Cities_Overview",                  "E_n, F_n, G_n", "pct_Joint, AvgLOC_*, GotMsg_*"],
        ["vw_RA_RideShare",       "#13_RideShare",                        "total_Res",     "ride counts + 4 share percentages"],
        ["vw_RA_PersonaPartTime", "#15_Persona_PartTime",                 "total_Res",     "PT_pct_Joint, PT_pct_Exclusive, RidePerBoarded_*"],
        ["vw_RA_IncentiveAmounts","#1_Snapp_Incentive_Amt, #2_Tapsi",    "n_total",       "incentive_range → Matrix columns"],
        ["vw_RA_IncentiveDuration","#4_Incentive_Duration",               "n_total",       "duration_bucket → Matrix columns"],
        ["vw_RA_Persona",         "#15_Persona (all sub-sheets)",         "n_total",       "dimension slicer; category → Matrix columns"],
        ["vw_RA_CommFree",        "#18_CommFree_Snapp, #18_CommFree_Tapsi","n",            "platform filter; pct_Got_Message, pct_Free_Comm_Ride"],
        ["vw_RA_CSRare",          "#CS_Sat_Snapp, #CS_Sat_Tapsi",        "n",             "Snapp/Tapsi 5 CS scores + solved %"],
        ["vw_RA_NavReco",         "#NavReco_Scores, #Reco_NPS",           "n",             "NPS scores + 4 nav app recommendation scores"],
    ]
)
doc.add_paragraph()

# ── Part 2: Connect Power BI ───────────────────────────────────────────────────
add_heading(doc, "Part 2 — Connecting Power BI to SQL Server", 1)
steps = [
    "Open Power BI Desktop. Go to Home → Get Data → SQL Server.",
    "Enter Server: 192.168.18.37 and Database: Cab_Studies. Set mode to Import (recommended).",
    "Click OK and enter credentials: username nasim.rajabi with your password.",
    "In the Navigator, expand the Cab schema. Select all views whose names start with vw_RA_. "
    "Also select any existing vw_* views you want (vw_KPISummary, vw_WeeklySatisfaction, etc.).",
    "Click Load. Wait for the import to complete.",
]
for s in steps:
    add_numbered(doc, s)

add_note(doc, (
    "Use Import mode, not DirectQuery, for better performance on large survey datasets. "
    "Set up a scheduled refresh in Power BI Service to keep data current."
))

# ── Part 3: Data Model ─────────────────────────────────────────────────────────
add_heading(doc, "Part 3 — Data Model Setup", 1)
add_para(doc, (
    "All vw_RA_* views share weeknumber and city columns but have no formal relationships "
    "between them. Two approaches work:"
))
add_heading(doc, "Option A — Shared Slicer (Recommended)", 3)
add_para(doc, (
    "Create a standalone WeekList table from any single view, then sync the week slicer "
    "across all report pages via View → Sync Slicers."
))
add_code_block(doc, "WeekList = DISTINCT(SELECTCOLUMNS(vw_RA_SatReview, \"weeknumber\", vw_RA_SatReview[weeknumber]))")

add_heading(doc, "Option B — Star Schema with DimWeek", 3)
add_para(doc, (
    "Create a DimWeek dimension table with one row per weeknumber. "
    "Create relationships from DimWeek[weeknumber] to each vw_RA_*[weeknumber]. "
    "This enables cross-filtering between views on different pages."
))

add_heading(doc, "City Order Table", 3)
add_para(doc, (
    "To enforce the same city sort order as the Python script's TOP_CITIES list, "
    "create a CityOrder table in Power BI (Enter Data) with these values:"
))
city_order_data = (
    "City,sort_order\n"
    "Tehran(city),1\nKaraj,2\nIsfahan,3\nShiraz,4\nMashhad,5\n"
    "Qom,6\nTabriz,7\nAhwaz,8\nSari,9\nRasht,10\n"
    "Urumieh,11\nYazd,12\nKerman,13\nGorgan,14\nGhazvin,15\n"
    "Arak,16\nKermanshah,17\nHamedan,18\nArdebil,19\n"
    "Bojnurd,20\nKhorramabad,21\nZanjan,22\nKish,23"
)
add_code_block(doc, city_order_data)
add_para(doc, (
    "Create a relationship from CityOrder[city] to each view's city column. "
    "In each Matrix visual, sort the Rows field by CityOrder[sort_order] ascending."
))

# ── Part 4: N-Cutoff Parameter ─────────────────────────────────────────────────
add_heading(doc, "Part 4 — N-Cutoff What-If Parameter", 1)
add_para(doc, (
    "This replicates the SHEET_MIN_N dictionary in the Python script. "
    "A single global threshold blanks any city row where the sample count is too small."
))
add_heading(doc, "Creating the Parameter", 3)
steps_param = [
    "Go to Modeling → New Parameter (What-If Parameter).",
    "Name: Min N Cutoff   |   Data type: Whole Number   |   Min: 0   |   Max: 500   |   Default: 0   |   Increment: 1",
    "Check 'Add slicer to this page'. Power BI creates a measure [Min N Cutoff Value] automatically.",
    "Copy the slicer to every report page, or use View → Sync Slicers.",
]
for s in steps_param:
    add_numbered(doc, s)

add_heading(doc, "DAX Measure Pattern", 3)
add_para(doc, (
    "Create one DAX measure per metric column. The pattern is always the same — "
    "only the table name, n-column, and value column change:"
))
add_code_block(doc, (
    "[Measure Name] =\n"
    "VAR MinN = [Min N Cutoff Value]\n"
    "VAR RowN = SUM(TableName[n_column])\n"
    "RETURN IF( RowN >= MinN, AVERAGE(TableName[metric_column]), BLANK() )"
))

add_heading(doc, "Examples — vw_RA_SatReview", 3)
add_code_block(doc, DAX_SAT_REVIEW)

add_heading(doc, "Examples — vw_RA_CitiesOverview (three independent n columns)", 3)
add_code_block(doc, DAX_CITIES)

add_heading(doc, "Examples — Long-Format Views", 3)
add_code_block(doc, DAX_LONG)

add_note(doc, (
    "Percentages in the views are stored as 0-100. Always divide by 100 in the DAX measure "
    "so Power BI's percentage format displays correctly. "
    "Satisfaction scores (1-5) should NOT be divided — use AVERAGE() directly."
))

# ── Part 5: Matrix Visuals ─────────────────────────────────────────────────────
add_heading(doc, "Part 5 — Building Matrix Visuals", 1)
add_heading(doc, "General Steps (same for every report page)", 3)
matrix_steps = [
    "Add a Matrix visual to the canvas.",
    "Rows field: drag the city column from the relevant view.",
    "Values field: drag your DAX measures (one per metric). Do NOT use raw columns — always use measures for n-cutoff logic.",
    "Format pane → Values: set decimal places to 1 or 2.",
    "Format pane → Conditional Formatting → Background Color: enable for each measure. Choose a color scale.",
    "Add a weeknumber slicer (single select, default = latest week).",
]
for s in matrix_steps:
    add_numbered(doc, s)

add_heading(doc, "Page-by-Page Reference", 3)
page_headers = ["Report Page", "Python Sheet", "View", "Slicers", "Rows", "Values Summary"]
add_summary_table(doc, page_headers, PAGE_ROWS)
doc.add_paragraph()

add_heading(doc, "Conditional Formatting Color Scales", 3)
cf_rows = [
    ["Satisfaction (1-5)", "Red → Yellow (3) → Green", "Fare, Income, Request, Incentive sat"],
    ["% Participation / GotMsg %", "White → Green", "Part_pct, GotMsg_pct columns"],
    ["Dissatisfaction reasons", "White → Red", "#8_Dissat columns"],
    ["Jnt@Tapsi ride share", "White → Red", "Higher Tapsi share = concerning for Snapp"],
    ["LOC (months)", "White → Green", "AvgLOC columns"],
    ["% Joint, % Dual SU", "White → Green / White → Red", "% Joint = green, % Dual SU = red"],
]
add_summary_table(doc, ["Metric Type", "Color Scale", "Applies To"], cf_rows)

# ── Part 6: WoW ───────────────────────────────────────────────────────────────
add_heading(doc, "Part 6 — Week-Over-Week (WoW) in DAX", 1)
add_para(doc, (
    "The Python script computes WoW by comparing the current week to week-1. "
    "In Power BI, use a CALCULATE with a weeknumber offset. "
    "Apply this pattern to any numeric metric column in any view:"
))
add_code_block(doc, DAX_WOW)

# ── Part 7: Sheets not yet covered ────────────────────────────────────────────
add_heading(doc, "Part 7 — Sheets Using WideMain / LongMain", 1)
add_para(doc, (
    "Some Python sheets require binary one-hot columns from DriverSurvey_WideMain "
    "or long-format answer data from DriverSurvey_LongMain/LongRare. "
    "These are more complex to replicate and can be added incrementally. "
    "The following sheets fall into this category:"
))
complex_sheets = [
    "#5_6_IncType — Received Incentive Types (uses Snapp/Tapsi Incentive Type__ binary columns from WideMain)",
    "#8_Dissat — Incentive Dissatisfaction Reasons (uses Snapp/Tapsi Last Incentive Unsatisfaction__ binary columns)",
    "#9_Dissat_Sum — Dissatisfaction Summary with WoW (pivot of #8)",
    "#14_Nav — Navigation Usage (uses LongMain question='Navigation Familiarity' etc.)",
    "#16_Ref — Referral Plan (joining bonus distributions from ShortMain)",
    "#17_Inactivity — Inactivity Before Tapsi Incentive (crosstab from ShortMain)",
    "#19_LuckyWheel — Tapsi Lucky Wheel / Magical Window (ShortMain)",
    "#20_Refusal — Request Refusal Reasons (LongRare)",
    "#Carfix / #Garage — SnappCarFix and TapsiGarage (ShortRare)",
    "#GPS / #Unpaid / #Decline / #DistOrigin / #Speed (all from ShortRare)",
]
for s in complex_sheets:
    add_bullet(doc, s)

add_note(doc, (
    "Use the existing server views vw_LongSurveyAnswers, vw_LongRareSurveyAnswers, "
    "vw_NavigationUsage, vw_WideIncentiveTypes, and vw_WideUnsatisfactionReasons "
    "as starting points for these sheets."
))

# ── Summary Table ──────────────────────────────────────────────────────────────
add_heading(doc, "Summary — Views and Python Sheet Equivalents", 1)
add_summary_table(
    doc,
    ["SQL View", "Python Sheet(s)", "N-Cutoff Column(s)"],
    SUMMARY_ROWS
)
doc.add_paragraph()
add_note(doc, (
    "All vw_RA_* views are in the [Cab] schema on Cab_Studies (192.168.18.37). "
    "Run CREATE OR ALTER VIEW statements in SSMS. "
    "After adding views, refresh the Power BI data source to pick them up."
))

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}  ({os.path.getsize(OUT)//1024} KB)")
