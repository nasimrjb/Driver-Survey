import sys
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

doc = Document()

# ── helpers ──────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    return p

def body(text):
    return doc.add_paragraph(text)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x7A)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F8')
    p._p.get_or_add_pPr().append(shading)
    return p

def note(text):
    p = doc.add_paragraph()
    run = p.add_run('⚠️  ' + text)
    run.italic = True
    run.font.size = Pt(9)
    return p

def sep():
    doc.add_paragraph()

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
doc.add_heading('Power BI Routine Analysis Guide', 0)
p = doc.add_paragraph()
p.add_run('Version 5 — Driver Survey Pipeline').bold = True
body('Database: Cab_Studies | Schema: [Cab] | Server: 192.168.18.37')
body('Covers all Routine Analysis Excel pages: #1–#6, #8–#9, #12–#19 (Mystery sheet #7 excluded)')
body('Total views: 37  (20 dashboard + 17 routine analysis)')
sep()

# ── WHAT'S NEW IN V5 ─────────────────────────────────────────────────────────
h1('What\'s New in Version 5')
body('V5 extends the existing 10 routine analysis (RA) views with 7 additional views and updates one existing view:')
rows = [
    ('Updated', 'vw_RA_CommFree', 'Page #18', 'Added type-specific message counts, participation rate, avg CF rides, pct CF ride-share'),
    ('New RA-11', 'vw_RA_IncentiveTypes', 'Pages #5, #6', 'Incentive type distribution: Excl Snapp / Jnt Snapp / Jnt Tapsi, avg CF rides'),
    ('New RA-12', 'vw_RA_IncentiveUnsatCity', 'Page #8', 'Dissatisfaction reasons by city for low-sat drivers'),
    ('New RA-13', 'vw_RA_IncentiveUnsatNational', 'Page #9', 'Same reasons at national level by segment'),
    ('New RA-14', 'vw_RA_Navigation', 'Page #14', 'Navigation app usage by city (Snapp & Tapsi)'),
    ('New RA-15', 'vw_RA_Referral', 'Page #16', 'Referral / joining bonus uptake by city'),
    ('New RA-16', 'vw_RA_TapsiInactivity', 'Page #17', 'Tapsi inactivity before incentive, joint drivers only'),
    ('New RA-17', 'vw_RA_LuckyWheel', 'Page #19', 'Lucky wheel usage rate and average amount'),
]
tbl = doc.add_table(rows=1+len(rows), cols=4)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for i, h in enumerate(['Status', 'View Name', 'Excel Page', 'Description']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
for i, (st, vw, pg, desc) in enumerate(rows):
    r = tbl.rows[i+1].cells
    r[0].text = st; r[1].text = vw; r[2].text = pg; r[3].text = desc
sep()

# ── COMMON DAX PATTERNS ───────────────────────────────────────────────────────
h1('Common DAX Patterns (Reference)')
body('All RA measures follow these conventions:')
bullet('yearweek column is VARCHAR "YY-WW" (e.g. "26-01"); yearweek_sort is INT for ordering.')
bullet('Min N threshold: create a measure [Min N] = 5 (or per-page value) to suppress small cells.')
bullet('Platform filter: UNION ALL views have a [platform] column ("Snapp" / "Tapsi") — always filter explicitly.')
bullet('Segment filter: national views have a [segment] column ("All Snapp" / "Joint Snapp" / "Joint Tapsi").')
body('Base yearweek selector pattern used in every measure:')
code('''VAR yw = IF(HASONEVALUE(<view>[yearweek]),
            VALUES(<view>[yearweek]),
            MAX(<view>[yearweek]))''')
body('WoW previous-week lookup pattern:')
code('''VAR PrevYW = CALCULATE(MAX(<view>[yearweek]),
                <view>[yearweek] < yw,
                ALL(<view>))''')
sep()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION: UPDATED vw_RA_CommFree
# ══════════════════════════════════════════════════════════════════════════════
h1('Updated RA-8: vw_RA_CommFree (Page #18 — Additional Measures)')
body('The view now includes: type-specific GotMsg counts (PayRide, EarnCF, RideCF, IncGuar, PayInc, CFSome), participation counts, avg CF rides, avg total rides, and avg pct CF ride-share.')
body('Add the following DAX measures to the CommFree measure group (these supplement the existing v4 measures):')
sep()

h2('New CommFree Measures — Snapp')
code('''CF GotMsg PayRide Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(n>=[Min N],CALCULATE(SUM(vw_RA_CommFree[GotMsg_PayRide]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp"),BLANK())''')

code('''CF GotMsg PayInc Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(n>=[Min N],CALCULATE(SUM(vw_RA_CommFree[GotMsg_PayInc]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp"),BLANK())''')

code('''CF GotMsg EarnCF Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(n>=[Min N],CALCULATE(SUM(vw_RA_CommFree[GotMsg_EarnCF]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp"),BLANK())''')

code('''CF GotMsg RideCF Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(n>=[Min N],CALCULATE(SUM(vw_RA_CommFree[GotMsg_RideCF]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp"),BLANK())''')

code('''CF GotMsg IncGuar Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(n>=[Min N],CALCULATE(SUM(vw_RA_CommFree[GotMsg_IncGuar]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp"),BLANK())''')

code('''CF GotMsg CFSome Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(n>=[Min N],CALCULATE(SUM(vw_RA_CommFree[GotMsg_CFSome]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp"),BLANK())''')

code('''CF pct Participated Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR GotMsg = CALCULATE(SUM(vw_RA_CommFree[Who_Got_Message]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
VAR Part   = CALCULATE(SUM(vw_RA_CommFree[Participated]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(GotMsg>=[Min N],DIVIDE(Part,GotMsg)*100,BLANK())''')

code('''CF Avg CF Rides Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_CommFree[Avg_CF_Rides]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp"),BLANK())''')

code('''CF Avg Total Rides Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_CommFree[Avg_Total_Rides]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp"),BLANK())''')

code('''CF pct CF RideShare Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp")
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_CommFree[Avg_pct_CF_RideShare]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Snapp"),BLANK())''')

h2('New CommFree Measures — Tapsi')
body('Duplicate all Snapp measures above, replacing platform="Snapp" with platform="Tapsi" and renaming "Snapp" → "Tapsi" in measure names. Example:')
code('''CF GotMsg PayRide Tapsi =
VAR yw = IF(HASONEVALUE(vw_RA_CommFree[yearweek]),VALUES(vw_RA_CommFree[yearweek]),MAX(vw_RA_CommFree[yearweek]))
VAR n = CALCULATE(SUM(vw_RA_CommFree[n]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Tapsi")
RETURN IF(n>=[Min N],CALCULATE(SUM(vw_RA_CommFree[GotMsg_PayRide]),vw_RA_CommFree[yearweek]=yw,vw_RA_CommFree[platform]="Tapsi"),BLANK())''')
sep()

# ══════════════════════════════════════════════════════════════════════════════
# RA-11: IncentiveTypes
# ══════════════════════════════════════════════════════════════════════════════
h1('RA-11: vw_RA_IncentiveTypes (Pages #5 & #6)')
body('Wide-format view. One row per yearweek+city. Covers incentive type breakdown for Exclusives Snapp, Joint Snapp, and Joint Tapsi segments.')
body('Power BI setup: Import vw_RA_IncentiveTypes. Add yearweek to WeekList slicer. City slicer uses [city] column.')
sep()

h2('Base Counts')
code('''IT n = CALCULATE(SUM(vw_RA_IncentiveTypes[n]),vw_RA_IncentiveTypes[yearweek]=IF(HASONEVALUE(vw_RA_IncentiveTypes[yearweek]),VALUES(vw_RA_IncentiveTypes[yearweek]),MAX(vw_RA_IncentiveTypes[yearweek])))''')
code('''IT n Excl =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveTypes[yearweek]),VALUES(vw_RA_IncentiveTypes[yearweek]),MAX(vw_RA_IncentiveTypes[yearweek]))
RETURN CALCULATE(SUM(vw_RA_IncentiveTypes[n_excl]),vw_RA_IncentiveTypes[yearweek]=yw)''')
code('''IT n Joint =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveTypes[yearweek]),VALUES(vw_RA_IncentiveTypes[yearweek]),MAX(vw_RA_IncentiveTypes[yearweek]))
RETURN CALCULATE(SUM(vw_RA_IncentiveTypes[n_joint]),vw_RA_IncentiveTypes[yearweek]=yw)''')

h2('Got-Message % by Segment')
for col, name in [
    ('pct_GotMsg_Excl_Snapp', 'Excl Snapp'),
    ('pct_GotMsg_Jnt_Snapp',  'Jnt Snapp'),
    ('pct_GotMsg_Jnt_Tapsi',  'Jnt Tapsi'),
    ('pct_GotMsg_Both',       'Both'),
    ('pct_GotMsg_Diff',       'Diff Platform'),
]:
    code(f'''IT pct GotMsg {name} =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveTypes[yearweek]),VALUES(vw_RA_IncentiveTypes[yearweek]),MAX(vw_RA_IncentiveTypes[yearweek]))
VAR n = [IT n]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_IncentiveTypes[{col}]),vw_RA_IncentiveTypes[yearweek]=yw),BLANK())''')

h2('Type % — Snapp Exclusive (base = n_excl)')
for col, label in [
    ('pct_PayRide_Excl',  'Pay per Ride'),
    ('pct_EarnCF_Excl',   'Earning-Based CF'),
    ('pct_RideCF_Excl',   'Ride-Based CF'),
    ('pct_IncGuar_Excl',  'Income Guarantee'),
    ('pct_PayInc_Excl',   'Pay After Income'),
    ('pct_CFSome_Excl',   'CF Some Trips'),
]:
    code(f'''IT pct {label.replace(" ","")} Excl =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveTypes[yearweek]),VALUES(vw_RA_IncentiveTypes[yearweek]),MAX(vw_RA_IncentiveTypes[yearweek]))
VAR n = [IT n Excl]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_IncentiveTypes[{col}]),vw_RA_IncentiveTypes[yearweek]=yw),BLANK())''')

h2('Type % — Joint Snapp (base = n_joint)')
for col, label in [
    ('pct_PayRide_JntSn', 'Pay per Ride'),
    ('pct_EarnCF_JntSn',  'Earning-Based CF'),
    ('pct_RideCF_JntSn',  'Ride-Based CF'),
    ('pct_IncGuar_JntSn', 'Income Guarantee'),
    ('pct_PayInc_JntSn',  'Pay After Income'),
    ('pct_CFSome_JntSn',  'CF Some Trips'),
]:
    code(f'''IT pct {label.replace(" ","")} JntSn =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveTypes[yearweek]),VALUES(vw_RA_IncentiveTypes[yearweek]),MAX(vw_RA_IncentiveTypes[yearweek]))
VAR n = [IT n Joint]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_IncentiveTypes[{col}]),vw_RA_IncentiveTypes[yearweek]=yw),BLANK())''')

h2('Type % — Joint Tapsi (base = n_joint)')
for col, label in [
    ('pct_PayRide_JntTp', 'Pay per Ride'),
    ('pct_EarnCF_JntTp',  'Earning-Based CF'),
    ('pct_RideCF_JntTp',  'Ride-Based CF'),
    ('pct_IncGuar_JntTp', 'Income Guarantee'),
    ('pct_PayInc_JntTp',  'Pay After Income'),
    ('pct_CFSome_JntTp',  'CF Some Trips'),
]:
    code(f'''IT pct {label.replace(" ","")} JntTp =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveTypes[yearweek]),VALUES(vw_RA_IncentiveTypes[yearweek]),MAX(vw_RA_IncentiveTypes[yearweek]))
VAR n = [IT n Joint]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_IncentiveTypes[{col}]),vw_RA_IncentiveTypes[yearweek]=yw),BLANK())''')

h2('Average CF Rides')
code('''IT Avg CF Rides Snapp =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveTypes[yearweek]),VALUES(vw_RA_IncentiveTypes[yearweek]),MAX(vw_RA_IncentiveTypes[yearweek]))
VAR n = [IT n]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_IncentiveTypes[Avg_CF_Rides_Snapp]),vw_RA_IncentiveTypes[yearweek]=yw),BLANK())''')

code('''IT Avg CF Rides Tapsi =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveTypes[yearweek]),VALUES(vw_RA_IncentiveTypes[yearweek]),MAX(vw_RA_IncentiveTypes[yearweek]))
VAR n = [IT n Joint]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_IncentiveTypes[Avg_CF_Rides_Tapsi]),vw_RA_IncentiveTypes[yearweek]=yw),BLANK())''')
sep()

# ══════════════════════════════════════════════════════════════════════════════
# RA-12: IncentiveUnsatCity
# ══════════════════════════════════════════════════════════════════════════════
h1('RA-12: vw_RA_IncentiveUnsatCity (Page #8)')
body('Wide-format view. One row per yearweek+city. Shows dissatisfaction reasons for Snapp-low-sat drivers and Joint Tapsi-low-sat drivers separately.')
body('N-cutoff: use [n_sn_low_sat] for Snapp reason measures, [n_tp_low_sat] for Tapsi reason measures.')
sep()

h2('Count Measures')
for col, name in [('n_all','All'), ('n_joint','Joint'), ('n_sn_low_sat','Sn LowSat'), ('n_tp_low_sat','Tp LowSat')]:
    code(f'''UC n {name} =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveUnsatCity[yearweek]),VALUES(vw_RA_IncentiveUnsatCity[yearweek]),MAX(vw_RA_IncentiveUnsatCity[yearweek]))
RETURN CALCULATE(SUM(vw_RA_IncentiveUnsatCity[{col}]),vw_RA_IncentiveUnsatCity[yearweek]=yw)''')

h2('Snapp Dissatisfaction Reason % (base = n_sn_low_sat)')
for col, label in [
    ('pct_Sn_NoTime',   'No Time'),
    ('pct_Sn_ImpAmt',   'Improper Amount'),
    ('pct_Sn_LowTime',  'Low Time'),
    ('pct_Sn_HardToDo', 'Hard To Do'),
    ('pct_Sn_NonPay',   'Non-Payment'),
    ('pct_Sn_Other',    'Other'),
]:
    code(f'''UC pct Sn {label.replace(" ","")} =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveUnsatCity[yearweek]),VALUES(vw_RA_IncentiveUnsatCity[yearweek]),MAX(vw_RA_IncentiveUnsatCity[yearweek]))
VAR n = [UC n Sn LowSat]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_IncentiveUnsatCity[{col}]),vw_RA_IncentiveUnsatCity[yearweek]=yw),BLANK())''')

h2('Joint Tapsi Dissatisfaction Reason % (base = n_tp_low_sat)')
for col, label in [
    ('pct_Tp_NoTime',   'No Time'),
    ('pct_Tp_ImpAmt',   'Improper Amount'),
    ('pct_Tp_LowTime',  'Low Time'),
    ('pct_Tp_HardToDo', 'Hard To Do'),
    ('pct_Tp_NonPay',   'Non-Payment'),
    ('pct_Tp_Other',    'Other'),
]:
    code(f'''UC pct Tp {label.replace(" ","")} =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveUnsatCity[yearweek]),VALUES(vw_RA_IncentiveUnsatCity[yearweek]),MAX(vw_RA_IncentiveUnsatCity[yearweek]))
VAR n = [UC n Tp LowSat]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_IncentiveUnsatCity[{col}]),vw_RA_IncentiveUnsatCity[yearweek]=yw),BLANK())''')
sep()

# ══════════════════════════════════════════════════════════════════════════════
# RA-13: IncentiveUnsatNational
# ══════════════════════════════════════════════════════════════════════════════
h1('RA-13: vw_RA_IncentiveUnsatNational (Page #9)')
body('Long-format view. segment column = "All Snapp" / "Joint Snapp" / "Joint Tapsi". No city — national only.')
body('Power BI setup: Do NOT connect a city slicer to this view. Add segment as a page-level filter or use as chart series.')
note('Each measure filters to its segment. Use segment_sort to order series in charts (1=All Snapp, 2=Joint Snapp, 3=Joint Tapsi).')
sep()

for seg, short in [('All Snapp','AllSn'), ('Joint Snapp','JntSn'), ('Joint Tapsi','JntTp')]:
    h2(f'Segment: {seg}')
    code(f'''UN n {short} =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveUnsatNational[yearweek]),VALUES(vw_RA_IncentiveUnsatNational[yearweek]),MAX(vw_RA_IncentiveUnsatNational[yearweek]))
RETURN CALCULATE(SUM(vw_RA_IncentiveUnsatNational[n]),vw_RA_IncentiveUnsatNational[yearweek]=yw,vw_RA_IncentiveUnsatNational[segment]="{seg}")''')
    code(f'''UN n LowSat {short} =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveUnsatNational[yearweek]),VALUES(vw_RA_IncentiveUnsatNational[yearweek]),MAX(vw_RA_IncentiveUnsatNational[yearweek]))
RETURN CALCULATE(SUM(vw_RA_IncentiveUnsatNational[n_low_sat]),vw_RA_IncentiveUnsatNational[yearweek]=yw,vw_RA_IncentiveUnsatNational[segment]="{seg}")''')
    for col, label in [
        ('pct_NoTime',  'NoTime'), ('pct_ImpAmt',  'ImpAmt'), ('pct_LowTime', 'LowTime'),
        ('pct_HardToDo','HardToDo'),('pct_NonPay', 'NonPay'), ('pct_Other',   'Other'),
    ]:
        code(f'''UN pct {label} {short} =
VAR yw = IF(HASONEVALUE(vw_RA_IncentiveUnsatNational[yearweek]),VALUES(vw_RA_IncentiveUnsatNational[yearweek]),MAX(vw_RA_IncentiveUnsatNational[yearweek]))
VAR n = [UN n LowSat {short}]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_IncentiveUnsatNational[{col}]),vw_RA_IncentiveUnsatNational[yearweek]=yw,vw_RA_IncentiveUnsatNational[segment]="{seg}"),BLANK())''')
sep()

# ══════════════════════════════════════════════════════════════════════════════
# RA-14: Navigation
# ══════════════════════════════════════════════════════════════════════════════
h1('RA-14: vw_RA_Navigation (Page #14)')
body('Long-format UNION ALL view. platform = "Snapp" or "Tapsi". Tapsi n counts joint drivers only.')
body('Snapp apps: Neshan, Balad, No Navigation App, Google Map, Waze, Other.')
body('Tapsi apps: Neshan, Balad, No Navigation App, In-App Navigation, Other.')
note('pct_GoogleMap and pct_Waze are NULL for Tapsi rows; pct_InAppNav is NULL for Snapp rows.')
sep()

for plat, short, cols in [
    ('Snapp','Sn',[('pct_Neshan','Neshan'),('pct_Balad','Balad'),('pct_None','NoNav'),('pct_GoogleMap','GoogleMap'),('pct_Waze','Waze'),('pct_Other','Other')]),
    ('Tapsi','Tp',[('pct_Neshan','Neshan'),('pct_Balad','Balad'),('pct_None','NoNav'),('pct_InAppNav','InAppNav'),('pct_Other','Other')]),
]:
    h2(f'Platform: {plat}')
    code(f'''Nav n {short} =
VAR yw = IF(HASONEVALUE(vw_RA_Navigation[yearweek]),VALUES(vw_RA_Navigation[yearweek]),MAX(vw_RA_Navigation[yearweek]))
RETURN CALCULATE(SUM(vw_RA_Navigation[n]),vw_RA_Navigation[yearweek]=yw,vw_RA_Navigation[platform]="{plat}")''')
    for col, label in cols:
        code(f'''Nav pct {label} {short} =
VAR yw = IF(HASONEVALUE(vw_RA_Navigation[yearweek]),VALUES(vw_RA_Navigation[yearweek]),MAX(vw_RA_Navigation[yearweek]))
VAR n = [Nav n {short}]
RETURN IF(n>=[Min N],CALCULATE(AVERAGE(vw_RA_Navigation[{col}]),vw_RA_Navigation[yearweek]=yw,vw_RA_Navigation[platform]="{plat}"),BLANK())''')
sep()

# ══════════════════════════════════════════════════════════════════════════════
# RA-15: Referral
# ══════════════════════════════════════════════════════════════════════════════
h1('RA-15: vw_RA_Referral (Page #16)')
body('Wide-format view. One row per yearweek+city. Snapp referral uses all drivers; Tapsi referral uses joint drivers only.')
sep()

for plat in ['Snapp', 'Tapsi']:
    h2(f'Platform: {plat}')
    code(f'''Ref n {plat} =
VAR yw = IF(HASONEVALUE(vw_RA_Referral[yearweek]),VALUES(vw_RA_Referral[yearweek]),MAX(vw_RA_Referral[yearweek]))
RETURN CALCULATE(SUM(vw_RA_Referral[n_{plat}]),vw_RA_Referral[yearweek]=yw)''')
    code(f'''Ref Joining {plat} =
VAR yw = IF(HASONEVALUE(vw_RA_Referral[yearweek]),VALUES(vw_RA_Referral[yearweek]),MAX(vw_RA_Referral[yearweek]))
VAR n = [Ref n {plat}]
RETURN IF(n>=[Min N],CALCULATE(SUM(vw_RA_Referral[joining_{plat}]),vw_RA_Referral[yearweek]=yw),BLANK())''')
    code(f'''Ref pct Joining {plat} =
VAR yw = IF(HASONEVALUE(vw_RA_Referral[yearweek]),VALUES(vw_RA_Referral[yearweek]),MAX(vw_RA_Referral[yearweek]))
VAR n = [Ref n {plat}]
VAR join = [Ref Joining {plat}]
RETURN IF(n>=[Min N],DIVIDE(join,n)*100,BLANK())''')
sep()

# ══════════════════════════════════════════════════════════════════════════════
# RA-16: TapsiInactivity
# ══════════════════════════════════════════════════════════════════════════════
h1('RA-16: vw_RA_TapsiInactivity (Page #17)')
body('Long-format view. One row per yearweek+city+inactivity_bucket (joint drivers only). bucket_sort orders the distribution bars.')
body('Power BI setup: Import view. Use inactivity_bucket as chart axis, sort by bucket_sort. Sync yearweek slicer.')
note('n_total (window function sum) enables direct pct = n / n_total without DAX aggregation tricks.')
sep()

h2('Measures')
code('''TI n Bucket =
VAR yw = IF(HASONEVALUE(vw_RA_TapsiInactivity[yearweek]),VALUES(vw_RA_TapsiInactivity[yearweek]),MAX(vw_RA_TapsiInactivity[yearweek]))
RETURN CALCULATE(SUM(vw_RA_TapsiInactivity[n]),vw_RA_TapsiInactivity[yearweek]=yw)''')

code('''TI n Total =
VAR yw = IF(HASONEVALUE(vw_RA_TapsiInactivity[yearweek]),VALUES(vw_RA_TapsiInactivity[yearweek]),MAX(vw_RA_TapsiInactivity[yearweek]))
RETURN CALCULATE(MAX(vw_RA_TapsiInactivity[n_total]),vw_RA_TapsiInactivity[yearweek]=yw)''')

code('''TI pct Bucket =
VAR yw = IF(HASONEVALUE(vw_RA_TapsiInactivity[yearweek]),VALUES(vw_RA_TapsiInactivity[yearweek]),MAX(vw_RA_TapsiInactivity[yearweek]))
VAR n     = [TI n Bucket]
VAR total = [TI n Total]
RETURN IF(total>=[Min N],DIVIDE(n,total)*100,BLANK())''')

body('Typical visual: 100% stacked bar chart by city, with inactivity_bucket as legend. Sort inactivity_bucket "Sort by Column" → bucket_sort.')
sep()

# ══════════════════════════════════════════════════════════════════════════════
# RA-17: LuckyWheel
# ══════════════════════════════════════════════════════════════════════════════
h1('RA-17: vw_RA_LuckyWheel (Page #19)')
body('Wide-format view. One row per yearweek+city. wheel column is Rial amount; 0 = driver did not spin.')
body('avg_wheel_amount = average amount among drivers who spun (wheel > 0), stored in the view.')
sep()

h2('Measures')
code('''LW n =
VAR yw = IF(HASONEVALUE(vw_RA_LuckyWheel[yearweek]),VALUES(vw_RA_LuckyWheel[yearweek]),MAX(vw_RA_LuckyWheel[yearweek]))
RETURN CALCULATE(SUM(vw_RA_LuckyWheel[n]),vw_RA_LuckyWheel[yearweek]=yw)''')

code('''LW n Users =
VAR yw = IF(HASONEVALUE(vw_RA_LuckyWheel[yearweek]),VALUES(vw_RA_LuckyWheel[yearweek]),MAX(vw_RA_LuckyWheel[yearweek]))
VAR n = [LW n]
RETURN IF(n>=[Min N],CALCULATE(SUM(vw_RA_LuckyWheel[n_users]),vw_RA_LuckyWheel[yearweek]=yw),BLANK())''')

code('''LW pct Usage =
VAR yw = IF(HASONEVALUE(vw_RA_LuckyWheel[yearweek]),VALUES(vw_RA_LuckyWheel[yearweek]),MAX(vw_RA_LuckyWheel[yearweek]))
VAR n     = [LW n]
VAR users = [LW n Users]
RETURN IF(n>=[Min N],DIVIDE(users,n)*100,BLANK())''')

code('''LW Avg Wheel Amount =
VAR yw = IF(HASONEVALUE(vw_RA_LuckyWheel[yearweek]),VALUES(vw_RA_LuckyWheel[yearweek]),MAX(vw_RA_LuckyWheel[yearweek]))
VAR users = [LW n Users]
RETURN IF(users>=[Min N],CALCULATE(AVERAGE(vw_RA_LuckyWheel[avg_wheel_amount]),vw_RA_LuckyWheel[yearweek]=yw),BLANK())''')

code('''LW Avg Wheel Amount WoW =
VAR yw = IF(HASONEVALUE(vw_RA_LuckyWheel[yearweek]),VALUES(vw_RA_LuckyWheel[yearweek]),MAX(vw_RA_LuckyWheel[yearweek]))
VAR PrevYW = CALCULATE(MAX(vw_RA_LuckyWheel[yearweek]),vw_RA_LuckyWheel[yearweek]<yw,ALL(vw_RA_LuckyWheel))
VAR curr = [LW Avg Wheel Amount]
VAR prev = CALCULATE([LW Avg Wheel Amount],vw_RA_LuckyWheel[yearweek]=PrevYW)
RETURN curr - prev''')

code('''LW pct Usage WoW =
VAR yw = IF(HASONEVALUE(vw_RA_LuckyWheel[yearweek]),VALUES(vw_RA_LuckyWheel[yearweek]),MAX(vw_RA_LuckyWheel[yearweek]))
VAR PrevYW = CALCULATE(MAX(vw_RA_LuckyWheel[yearweek]),vw_RA_LuckyWheel[yearweek]<yw,ALL(vw_RA_LuckyWheel))
VAR curr = [LW pct Usage]
VAR prev = CALCULATE([LW pct Usage],vw_RA_LuckyWheel[yearweek]=PrevYW)
RETURN curr - prev''')
sep()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX: Page-View Mapping
# ══════════════════════════════════════════════════════════════════════════════
h1('Appendix: Page–View–Measure Reference')
rows_ref = [
    ('1', 'Dashboard overview', 'vw_SatOverview, vw_CitiesOverview', 'v4'),
    ('2', 'Ride-share metrics', 'vw_RideShare', 'v4'),
    ('3', 'Persona / Part-time', 'vw_PersonaPartTime', 'v4'),
    ('4', 'Incentive amounts', 'vw_IncentiveAmounts', 'v4'),
    ('5', 'Incentive types (Excl)', 'vw_RA_IncentiveTypes', 'v5'),
    ('6', 'Incentive types (Joint)', 'vw_RA_IncentiveTypes', 'v5'),
    ('7', 'Mystery Shopping', '(excluded — separate data source)', '—'),
    ('8', 'Incentive unsat by city', 'vw_RA_IncentiveUnsatCity', 'v5'),
    ('9', 'Incentive unsat national', 'vw_RA_IncentiveUnsatNational', 'v5'),
    ('12', 'Incentive duration', 'vw_RA_IncentiveDuration', 'v4'),
    ('13', 'Persona', 'vw_RA_Persona', 'v4'),
    ('14', 'Navigation app usage', 'vw_RA_Navigation', 'v5'),
    ('15', 'CS Satisfaction (Rare)', 'vw_RA_CSRare', 'v4'),
    ('16', 'Referral / Joining bonus', 'vw_RA_Referral', 'v5'),
    ('17', 'Tapsi inactivity', 'vw_RA_TapsiInactivity', 'v5'),
    ('18', 'CommFree incentive', 'vw_RA_CommFree', 'v4+v5'),
    ('19', 'Lucky wheel', 'vw_RA_LuckyWheel', 'v5'),
]
tbl2 = doc.add_table(rows=1+len(rows_ref), cols=4)
tbl2.style = 'Table Grid'
hdr2 = tbl2.rows[0].cells
for i, h in enumerate(['Page #', 'Description', 'SQL View(s)', 'Guide']):
    hdr2[i].text = h
    hdr2[i].paragraphs[0].runs[0].bold = True
for i, (pg, desc, view, ver) in enumerate(rows_ref):
    r = tbl2.rows[i+1].cells
    r[0].text = pg; r[1].text = desc; r[2].text = view; r[3].text = ver

sep()
body('Generated by generate_v5_guide.py — Driver Survey Power BI Pipeline v5')

out = r'D:\Work\Driver Survey\PowerBI\PowerBI_Routine_Analysis_Guide_v5.docx'
doc.save(out)
print(f'Saved: {out}')
